from docx import Document
from docx.shared import Pt, Inches
from pathlib import Path

def markdown_to_docx(md_path, docx_path):
    """Convert markdown file to Word document."""
    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    for line in lines:
        line = line.rstrip('\n')
        
        # Skip empty lines at document start
        if not line.strip() and len(doc.paragraphs) == 0:
            continue
        
        # Handle headings
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        
        # Handle bullet points
        elif line.strip().startswith('- '):
            doc.add_paragraph(line.strip()[2:], style='List Bullet')
        
        # Handle empty lines
        elif not line.strip():
            # Only add paragraph if previous wasn't empty
            if len(doc.paragraphs) > 0 and doc.paragraphs[-1].text.strip():
                doc.add_paragraph()
        
        # Handle regular text
        else:
            # Check if it's a code block indicator
            if line.strip() == '```':
                continue
            doc.add_paragraph(line)
    
    doc.save(docx_path)
    print(f"Created: {docx_path}")

# Export both documents
base_path = Path(r"c:\Users\pmend\project\Vetting app")
docs_path = base_path / "docs"

# Technical Overview
markdown_to_docx(
    docs_path / "TECHNICAL_OVERVIEW_LIVING.md",
    base_path / "Technical_Overview_Living.docx"
)

# Architecture Reference
markdown_to_docx(
    docs_path / "ARCHITECTURE_LIVING.md",
    base_path / "Architecture_Living.docx"
)

print("\nBoth documents exported successfully!")
