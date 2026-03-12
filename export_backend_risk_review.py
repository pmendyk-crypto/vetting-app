from docx import Document
from docx.shared import Pt

output_path = r"c:\Users\pmend\project\Vetting app\Backend_Risk_Review.docx"

doc = Document()

# Title
title = doc.add_paragraph()
run = title.add_run("Backend Engineering Risk Review")
run.bold = True
run.font.size = Pt(20)

doc.add_paragraph("Project: Vetting App (FastAPI)")
doc.add_paragraph("Generated: 2026-03-08")
doc.add_paragraph("Scope: architecture, security, database, maintainability, configuration, deployment assumptions, and testing posture.")


def add_section(title_text, bullets):
    doc.add_heading(title_text, level=1)
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")


add_section("ARCHITECTURE RISKS", [
    "Overly large central module: `app/main.py` is ~300KB and combines routing, auth, schema bootstrap/migration logic, storage, reporting, PDF generation, and tenant administration. This creates high regression risk and difficult ownership boundaries.",
    "Mixed responsibilities in route layer: route handlers include business rules, direct SQL, and response rendering in one place (for example large admin and superuser sections in `app/main.py`).",
    "Router split is incomplete: `app/routers/multitenant.py` exists, but router mounting is commented out in `app/main.py` (`app.include_router(multitenant.router)`), leaving a partially modularized architecture.",
    "Data-access abstraction duplication: DB connection and SQLAlchemy wrapper logic exist both in `app/main.py` and `app/db.py`, increasing drift risk.",
])

add_section("SECURITY RISKS", [
    "Secret handling fallback is weak: `APP_SECRET` defaults to `dev-secret-change-me` in `app/main.py`; app warns but still starts. In production, this should fail fast.",
    "Diagnostic endpoint exposure: `/diag/schema` is publicly reachable in `app/main.py` and discloses schema/table information useful for attackers.",
    "Attachment path trust risk: local attachment endpoints use `stored_filepath` from DB and serve via `FileResponse(stored_path)` (`app/main.py`) without strict path allowlist enforcement under `UPLOAD_DIR`.",
    "Session cookie hardening appears limited: `SessionMiddleware` is configured with `same_site=\"lax\"` and no explicit secure-cookie enforcement setting in code (`app/main.py`).",
    "Potential information leakage in logs: SMTP fallback prints email content/body to stdout when SMTP is not configured (`app/main.py`).",
])

add_section("DATABASE RISKS", [
    "Schema management is runtime-driven: extensive `CREATE TABLE IF NOT EXISTS` and `ALTER TABLE` execution inside app startup (`app/main.py`) can cause environment drift and non-repeatable behavior.",
    "Migration strategy is mixed: formal SQL migrations exist in `database/migrations/`, but app also mutates schema at runtime, reducing confidence in deterministic deployments.",
    "SQLite/PostgreSQL dual mode uses compatibility wrapper logic (`app/main.py` and `app/db.py`) that emulates sqlite-style behavior on SQLAlchemy; subtle SQL/transaction differences can cause hard-to-debug production-only issues.",
    "Column-conditional logic (`table_has_column`) is used widely across core workflows (`app/main.py`), indicating multiple schema states are expected and increasing path complexity.",
])

add_section("MAINTAINABILITY RISKS", [
    "Duplicated logic across modules: auth/org/dependency patterns are split among `app/main.py`, `app/dependencies.py`, and `app/models.py` with overlapping responsibilities.",
    "Complex route handlers: many endpoints in `app/main.py` include validation, SQL, file IO, and business events in a single function, making local changes risky.",
    "Inline SQL spread across application code: direct SQL statements are scattered through route handlers and helpers rather than centralized repositories/services.",
    "Template and backend coupling: many flows rely on hidden form fields (`org_id_hidden` in templates) and route-specific assumptions, increasing integration fragility.",
])

add_section("OBSERVABILITY RISKS", [
    "Logging is mostly `print()` statements (`app/main.py`) with no structured logging strategy, levels, correlation IDs, or centralized sink pattern.",
    "Health model is shallow: `/health` and `/healthz` return static healthy status and do not verify critical dependencies (DB, blob storage, SMTP).",
    "No explicit metrics/tracing instrumentation appears in code (no Prometheus/OpenTelemetry style hooks found).",
    "Operational failure paths often log and continue (for example startup warnings) rather than emitting machine-actionable alerts.",
])

add_section("CONFIGURATION RISKS", [
    "Environment variables are numerous and partly duplicated in naming conventions (for example `SMTP_PASS` and `SMTP_PASSWORD` both appear in code paths).",
    "Security-sensitive config may default to permissive behavior instead of blocking startup (for example default app secret).",
    "Feature toggles are implicit/partial (for example multi-tenant router disabled by comment rather than explicit runtime policy), increasing operator ambiguity.",
])

add_section("DEPLOYMENT ASSUMPTION RISKS", [
    "Deployment script is environment-specific and hardcoded (`deploy.ps1` has fixed ACR/app/resource-group names and live URL), reducing portability and increasing accidental-target risk.",
    "Docker image runs with shell-form CMD (`Dockerfile`), which can complicate signal handling and graceful shutdown behavior under orchestrators.",
    "Production path assumes Azure App Service + ACR operational model; alternative environments are not represented as first-class deployment workflows.",
])

add_section("TESTING RISKS", [
    "Automated test posture is limited: test files are mostly ad-hoc request scripts (`test_api.py`, `test_api_final.py`, `test_endpoints.py`, `test_e2e.py`) rather than structured pytest suites with assertions and CI gating.",
    "`tests/` directory currently contains helper scripts rather than comprehensive unit/integration coverage for critical auth/org isolation and data mutation logic.",
    "No clear evidence in repo of coverage thresholds, test matrix for SQLite vs PostgreSQL, or migration regression tests.",
])

add_section("PRIORITY NEXT STEPS", [
    "1) Split `app/main.py` into domain routers/services (auth, cases, admin, settings, superuser, storage).",
    "2) Enforce secure configuration at startup (fail if default secret in non-dev).",
    "3) Lock down `/diag/schema` and harden attachment file-path serving with strict allowlist checks.",
    "4) Move fully to migration-led schema changes and reduce runtime DDL.",
    "5) Introduce structured logging + readiness checks + basic metrics.",
    "6) Establish pytest-based test suite for authz/org isolation and DB backend parity.",
])

doc.save(output_path)
print(f"Created: {output_path}")
