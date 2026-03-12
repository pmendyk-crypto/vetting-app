import ast
from pathlib import Path
from docx import Document
from docx.shared import Pt

ROOT = Path(r"c:\Users\pmend\project\Vetting app")
OUT = ROOT / "FastAPI_Route_Inventory.docx"

HTTP_METHODS = {"get", "post", "put", "delete"}
EXCLUDE_DIRS = {".git", ".venv", "__pycache__", ".pytest_cache"}


def short_purpose(func_name: str, route_path: str, doc: str | None) -> str:
    if doc:
        first = doc.strip().split("\n")[0].strip()
        if first:
            return first[:180]

    name = func_name.lower()
    path = route_path.lower()

    if "login" in name or "/login" in path:
        return "Authenticate user and start session"
    if "logout" in name or "/logout" in path:
        return "Terminate session and sign out user"
    if "forgot" in name or "reset" in name:
        return "Password recovery and reset flow"
    if "health" in name or "health" in path:
        return "Service health check endpoint"
    if "diag" in name or "diag" in path:
        return "Diagnostic endpoint for runtime/schema inspection"
    if "admin" in path:
        return "Admin workflow endpoint"
    if "radiologist" in path or "vet" in path:
        return "Radiologist queue/review/vetting endpoint"
    if "submit" in path or "intake" in path:
        return "Case intake/submission endpoint"
    if "settings" in path:
        return "Configuration and master-data management"
    if "super" in path or "/mt" in path or "org" in path:
        return "Multi-tenant or organization management endpoint"
    return "Application endpoint"


def extract_route_decorator_info(decorator):
    if not isinstance(decorator, ast.Call):
        return None
    func = decorator.func
    if not isinstance(func, ast.Attribute):
        return None
    if func.attr not in HTTP_METHODS:
        return None
    if not isinstance(func.value, ast.Name):
        return None
    if func.value.id not in {"app", "router"}:
        return None

    path = "<unknown>"
    if decorator.args and isinstance(decorator.args[0], ast.Constant) and isinstance(decorator.args[0].value, str):
        path = decorator.args[0].value

    return {
        "method": func.attr.upper(),
        "owner": func.value.id,
        "path": path,
    }


def detect_auth_requirement(func_node: ast.FunctionDef | ast.AsyncFunctionDef, src_lines: list[str]) -> str:
    start = max(0, func_node.lineno - 1)
    end = min(len(src_lines), (func_node.end_lineno or func_node.lineno))
    body = "\n".join(src_lines[start:end])

    auth_flags = []
    if "require_superuser(" in body:
        auth_flags.append("superuser")
    if "require_admin(" in body:
        auth_flags.append("admin")
    if "require_radiologist(" in body:
        auth_flags.append("radiologist")
    if "require_login(" in body:
        auth_flags.append("authenticated")

    # Dependency-injected auth patterns
    if "Depends(require_superuser)" in body:
        auth_flags.append("superuser")
    if "Depends(require_org_admin)" in body:
        auth_flags.append("org_admin")
    if "Depends(require_org_context)" in body:
        auth_flags.append("org_context")
    if "Depends(require_login)" in body:
        auth_flags.append("authenticated")

    # Remove duplicates preserving order
    dedup = []
    for f in auth_flags:
        if f not in dedup:
            dedup.append(f)

    if not dedup:
        return "public/unspecified"
    return ", ".join(dedup)


def should_scan(path: Path) -> bool:
    if path.suffix != ".py":
        return False
    parts = set(path.parts)
    if parts & EXCLUDE_DIRS:
        return False
    return True


def collect_routes():
    routes = []
    for py_file in ROOT.rglob("*.py"):
        if not should_scan(py_file):
            continue

        try:
            src = py_file.read_text(encoding="utf-8")
        except Exception:
            continue

        try:
            tree = ast.parse(src)
        except Exception:
            continue

        src_lines = src.splitlines()

        for node in ast.walk(tree):
            if not isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)):
                continue

            decorators = []
            for dec in node.decorator_list:
                info = extract_route_decorator_info(dec)
                if info:
                    decorators.append(info)

            if not decorators:
                continue

            doc = ast.get_docstring(node)
            rel = py_file.relative_to(ROOT)
            auth = detect_auth_requirement(node, src_lines)

            for d in decorators:
                route_path = d["path"]
                routes.append({
                    "route": route_path,
                    "method": d["method"],
                    "function": node.name,
                    "file": str(rel).replace("\\", "/"),
                    "purpose": short_purpose(node.name, route_path, doc),
                    "auth": auth,
                })

    routes.sort(key=lambda r: (r["file"], r["route"], r["method"]))
    return routes


def classify_routes(routes):
    unused = []
    debug = []
    admin = []
    multitenant = []

    for r in routes:
        route = r["route"].lower()
        file = r["file"].lower()

        if file.startswith("app/routers/multitenant.py"):
            unused.append(r)
            multitenant.append(r)

        if any(k in route for k in ["/health", "/diag", "/test", "/robots"]):
            debug.append(r)

        if route.startswith("/admin") or route.startswith("/settings") or route.startswith("/super"):
            admin.append(r)

        if route.startswith("/mt") or route.startswith("/super") or "org" in route or file.endswith("multitenant.py"):
            multitenant.append(r)

    def unique(items):
        seen = set()
        out = []
        for i in items:
            key = (i["route"], i["method"], i["function"], i["file"])
            if key in seen:
                continue
            seen.add(key)
            out.append(i)
        return out

    return unique(unused), unique(debug), unique(admin), unique(multitenant)


def add_table(doc: Document, title: str, rows: list[dict]):
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Route"
    hdr[1].text = "Method"
    hdr[2].text = "Function"
    hdr[3].text = "File"
    hdr[4].text = "Purpose"
    hdr[5].text = "Auth"

    for r in rows:
        cells = table.add_row().cells
        cells[0].text = r["route"]
        cells[1].text = r["method"]
        cells[2].text = r["function"]
        cells[3].text = r["file"]
        cells[4].text = r["purpose"]
        cells[5].text = r["auth"]


def add_simple_list(doc: Document, title: str, rows: list[dict], note: str | None = None):
    doc.add_heading(title, level=2)
    if note:
        doc.add_paragraph(note)
    if not rows:
        doc.add_paragraph("None identified.")
        return
    for r in rows:
        doc.add_paragraph(
            f"{r['method']} {r['route']} ({r['function']}) - {r['file']}",
            style="List Bullet",
        )


def build_doc(routes, unused, debug, admin, multitenant):
    doc = Document()
    t = doc.add_paragraph()
    run = t.add_run("FastAPI Route Inventory")
    run.bold = True
    run.font.size = Pt(20)

    doc.add_paragraph("Generated: 2026-03-08")
    doc.add_paragraph(f"Total routes identified: {len(routes)}")

    add_table(doc, "All Routes", routes)

    doc.add_heading("Route Groups", level=1)
    add_simple_list(
        doc,
        "Routes That Appear Unused",
        unused,
        note="These are currently in app/routers/multitenant.py, and main app route mounting appears disabled/commented.",
    )
    add_simple_list(doc, "Debug/Diagnostic Routes", debug)
    add_simple_list(doc, "Admin Routes", admin)
    add_simple_list(doc, "Multi-tenant Related Routes", multitenant)

    doc.save(OUT)


if __name__ == "__main__":
    routes = collect_routes()
    unused, debug, admin, multitenant = classify_routes(routes)
    build_doc(routes, unused, debug, admin, multitenant)
    print(f"Created: {OUT}")
    print(f"Routes: {len(routes)}")
    print(f"Unused (appears): {len(unused)}")
    print(f"Debug: {len(debug)}")
    print(f"Admin: {len(admin)}")
    print(f"Multi-tenant: {len(multitenant)}")
