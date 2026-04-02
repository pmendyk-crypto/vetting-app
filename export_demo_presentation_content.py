from docx import Document
from docx.shared import Pt


OUTPUT_PATH = r"c:\Users\pmend\project\Vetting app\Demo_Presentation_Content_Draft.docx"


def add_title(doc: Document, text: str, size: int = 20) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_slide(
    doc: Document,
    slide_no: int,
    title: str,
    purpose: str,
    headline: str,
    body: str,
    bullets: list[str],
    presenter_script: str,
    visual: str,
) -> None:
    add_heading(doc, f"Slide {slide_no}: {title}", level=2)
    doc.add_paragraph(f"Purpose: {purpose}")

    p = doc.add_paragraph()
    p.add_run("Headline text").bold = True
    doc.add_paragraph(headline)

    p = doc.add_paragraph()
    p.add_run("Body text").bold = True
    doc.add_paragraph(body)

    p = doc.add_paragraph()
    p.add_run("Optional bullet text").bold = True
    add_bullets(doc, bullets)

    p = doc.add_paragraph()
    p.add_run("Presenter script").bold = True
    doc.add_paragraph(presenter_script)

    p = doc.add_paragraph()
    p.add_run("Suggested visual").bold = True
    doc.add_paragraph(visual)


doc = Document()

add_title(doc, "Lumos RadFlow - Demo Presentation Content", 20)
doc.add_paragraph(
    "Presentation working draft with fuller text that can be copied directly into PowerPoint slides or used as speaker notes."
)

add_heading(doc, "How To Use This Draft")
add_bullets(
    doc,
    [
        "Use the Headline text and Body text directly on the slide.",
        "Use Optional bullet text if you want a more feature-led slide layout.",
        "Use Presenter script as your speaking notes during the demo.",
        "Keep the tone outcome-led: speed, control, visibility, and governance.",
    ],
)

add_heading(doc, "Core Positioning")
add_bullets(
    doc,
    [
        "RadFlow should be presented as a workflow solution for day-to-day radiology operations, not just a simple vetting screen.",
        "It supports both acute and elective work.",
        "It improves turnaround, operational visibility, and governance.",
        "It supports auditable, role-based working with a clear path to stronger authentication from day one.",
    ],
)

add_slide(
    doc,
    1,
    "Introduction",
    "Open with a strong, simple positioning statement.",
    "Lumos RadFlow",
    (
        "A radiology workflow platform designed to support day-to-day vetting operations across acute and elective services, "
        "with faster decisions, clearer accountability, and stronger operational control."
    ),
    [
        "Acute and elective radiology workflow support",
        "Rapid, accurate vetting",
        "Visible turnaround times and audit-ready reporting",
    ],
    (
        "I would open by positioning RadFlow as more than a booking or triage tool. "
        "The value is that it supports how a radiology department actually works day to day, from referral intake through to justified outcome, "
        "while improving both speed and governance."
    ),
    "Product title with a clean workflow or dashboard image.",
)

add_slide(
    doc,
    2,
    "The Problem Today",
    "Show the audience that the current operating model is fragmented.",
    "There is still no true standalone workflow platform designed to support radiology departments in day-to-day vetting operations.",
    (
        "Today, departments usually rely either on vetting options built into RIS platforms such as Magentus or Soliton, "
        "or on email-based processes, usually through NHS.net. Both approaches create fragmentation, limited visibility, and extra manual follow-up."
    ),
    [
        "RIS vetting is often limited to the local workflow inside the system",
        "Email-based vetting is difficult to track, govern, and standardise",
        "Neither model gives strong operational visibility across the full case journey",
    ],
    (
        "This slide is important commercially. It creates the gap in the market. "
        "The message is that departments are still forced to work around systems rather than use a workflow designed specifically for radiology vetting operations."
    ),
    "A simple comparison graphic: RIS-only or email-based workflow versus RadFlow unified workflow.",
)

add_slide(
    doc,
    3,
    "Why This Matters",
    "Connect the problem to operational pain.",
    "Fragmented vetting creates delay, inconsistency, and poor visibility.",
    (
        "When vetting is split across system screens, inboxes, and manual follow-up, teams lose time, decisions are harder to track, "
        "and service leads have limited visibility over status, ownership, and turnaround."
    ),
    [
        "More admin burden",
        "Delayed decision-making",
        "Reduced consistency",
        "Limited traceability",
        "Harder KPI tracking",
    ],
    (
        "The audience should recognise that this is not only about convenience. "
        "It is about service performance, governance, and the ability to run a reliable radiology operation at scale."
    ),
    "Before-and-after process visual or an operations-focused infographic.",
)

add_slide(
    doc,
    4,
    "What RadFlow Delivers",
    "Present the platform as the answer to the operational gap.",
    "RadFlow brings radiology vetting into one controlled workflow.",
    (
        "RadFlow supports referral intake, assignment, practitioner review, status tracking, exports, and auditability in one place, "
        "helping teams move from fragmented administration to a more controlled and efficient operating model. "
        "It gives admins full visibility of what is in the system and gives practitioners a cleaner dedicated workspace to complete review quickly."
    ),
    [
        "Structured referral intake with attachment handling",
        "Admin worklist with full visibility of active cases in the system",
        "Fast case creation and assignment to the right practitioner queue",
        "Single workflow from submission to justification outcome",
        "Embedded iRefer guideline lookup",
        "Clear status tracking across pending, justified, rejected, and reopened cases",
        "PDF and CSV outputs for communication, reporting, and audit",
    ],
    (
        "This is the moment where you show RadFlow as a platform, not just a feature. "
        "It connects the full operational chain so the department can work faster and with more confidence."
    ),
    "Admin dashboard or simple end-to-end journey visual.",
)

add_slide(
    doc,
    5,
    "Service Scope",
    "Clarify how the product fits the clinical operating model.",
    "RadFlow is designed to support both acute and elective radiology work.",
    (
        "The platform is not limited to acute teleradiology. It also supports elective pathways, giving teams one place to manage workload, "
        "track status, and monitor turnaround time across different service types."
    ),
    [
        "Supports acute teleradiology workflows",
        "Supports elective radiology work",
        "Visible TAT can be tracked across the service",
        "Operational views can be aligned to service KPIs",
    ],
    (
        "This helps widen the commercial story. "
        "RadFlow is not just useful in one urgent pathway. It can support the broader operational reality of the department."
    ),
    "Use a split visual showing acute and elective flows feeding into one platform.",
)

add_slide(
    doc,
    6,
    "Clinical Rules and Standardisation",
    "Show that the workflow can support real service rules.",
    "RadFlow supports a more standardised and consistent vetting process.",
    (
        "The workflow can support service rules where vetting is required for all CT scans and for contrast examinations in MR, "
        "helping teams apply policy more consistently and reducing the risk of ad hoc process variation."
    ),
    [
        "Vetting required for all CT scans",
        "Vetting required for contrast MR examinations",
        "Protocol-led workflow supports consistency",
        "Guideline lookup supports clinical decision-making",
    ],
    (
        "This slide helps anchor the product in real clinical operations. "
        "It shows that RadFlow can support standardised policy-driven working rather than relying on informal practice."
    ),
    "A simple policy-to-workflow diagram or practitioner screen with protocol selection visible.",
)

add_slide(
    doc,
    7,
    "Core Functionality",
    "Show the practical features the team will see in the demo.",
    "The platform supports the real day-to-day steps of vetting, assignment, and review.",
    (
        "From the admin side, teams can create, assign, filter, and monitor cases. From the practitioner side, users can review case details, "
        "open attachments, check guidance, record decisions, and submit justification in a structured format. "
        "The design is focused on a clean, efficient workflow for both sides of the service."
    ),
    [
        "Admin worklist with filters, status views, and full visibility of what is in the system",
        "Admin can assign the right practitioner, edit a case whenever needed, and keep the workflow moving",
        "Admin notification option to chase practitioners when action is needed",
        "Practitioner worklist designed as a dedicated clean review workspace",
        "Practitioner review screen with case details, attachment preview, decision, protocol, and comments",
        "Mobile-compatible practitioner workflow for easier access, speed, and efficiency",
        "Protocol templates with instructions to support standardised justification",
        "Referral Scanner trial to pre-fill case details from referral documents",
        "Clear path to a draft review queue for automated intake in a later phase",
        "Case timeline and export capability for governance and traceability",
    ],
    (
        "This is the practical workflow slide. "
        "The audience should come away feeling that the system already supports the main daily activities required to run the service, "
        "and that both admin and practitioner users have a workflow that is built for speed rather than clutter."
    ),
    "Split screenshot of admin and practitioner views.",
)

add_slide(
    doc,
    8,
    "Operational Value",
    "Turn features into service-level benefits.",
    "RadFlow improves speed, visibility, and accountability across the workflow.",
    (
        "By bringing vetting into one structured platform, RadFlow reduces manual effort, supports more consistent decision-making, "
        "and gives operational leads clearer visibility of workload, status, and turnaround. "
        "It also makes it easier for practitioners to respond quickly through a dedicated, mobile-compatible worklist."
    ),
    [
        "Reduces manual admin effort and duplicate data entry",
        "Improves turnaround by replacing manual chasing with a guided workflow",
        "Supports the right decision for the right patient with the right context",
        "Improves operational visibility through dashboards and reporting",
        "Helps practitioners work faster through a clean and accessible review experience",
        "Creates a stronger foundation for quality improvement and service management",
    ],
    (
        "Keep this commercially focused. "
        "You are showing that the app helps the department run better, not just click through tasks faster."
    ),
    "Benefits layout with 4 to 5 strong value statements.",
)

add_slide(
    doc,
    9,
    "Turnaround Time and KPI Visibility",
    "Emphasise measurable operational impact.",
    "RadFlow makes turnaround time visible and trackable against service KPIs.",
    (
        "The platform introduces visible TAT across the workflow, allowing teams to monitor performance, identify delays, "
        "and track progress against operational KPIs in a way that is difficult to achieve through email-led processes."
    ),
    [
        "Visible TAT at case level",
        "Dashboard views of workload and status",
        "Better KPI tracking for operational leads",
        "Supports pilot measurement and service improvement",
    ],
    (
        "This is a strong buyer slide because it links the product to measurable impact. "
        "It gives the audience a reason to pilot the platform and test the effect on service performance."
    ),
    "Dashboard screenshot with counts, status mix, and average TAT highlighted.",
)

add_slide(
    doc,
    10,
    "Governance, Access and Data Security",
    "Build trust in access control and accountability.",
    "RadFlow is designed to support controlled access, traceability, and operational confidence.",
    (
        "Governance and data security are central to service confidence. Access can be restricted to the right users, important actions can be traced across the case lifecycle, "
        "and the platform has a clear path to stronger day-one authentication measures such as authenticator-based access for dedicated users."
    ),
    [
        "Role-based access",
        "Day-one roadmap for authenticator-based protection",
        "Only dedicated users can access relevant parts of the workflow",
        "Important actions can be traced and audited",
        "Supports stronger governance, accountability, and service confidence",
    ],
    (
        "This slide should reassure both operational and governance stakeholders. "
        "The message is that RadFlow supports controlled access and a traceable workflow, which is essential for service trust."
    ),
    "Security and audit-themed slide with access control and audit trail motifs.",
)

add_slide(
    doc,
    11,
    "Current Security Position",
    "Explain honestly where RadFlow stands today.",
    "Current position: RadFlow already has a strong baseline for access control and auditability, with a clear roadmap for further hardening.",
    (
        "Today, the platform already includes role-based access control, password hashing, session timeout management, login rate limiting, password reset token handling, "
        "and event history for case and notification actions. For production confidence, the next layer is stronger hardening around authenticator-based access, secure cookie policy, and production configuration controls."
    ),
    [
        "Role-based access for admin and practitioner workflows",
        "PBKDF2 password hashing for stored credentials",
        "Session-based authentication with idle timeout",
        "Login rate limiting to reduce brute-force risk",
        "Password reset token flow",
        "Case and notification event history to support traceability and audit",
        "Production path to PostgreSQL and managed storage",
    ],
    (
        "I would describe the current state as secure by design, but still on a hardening journey. "
        "That is a stronger and more credible message than claiming enterprise-grade completeness too early."
    ),
    "A simple maturity visual: baseline controls today, hardening roadmap next.",
)

add_slide(
    doc,
    12,
    "Moving from Proof of Concept to Pilot to Live",
    "Show that the platform has a credible route to operational rollout.",
    "RadFlow provides a practical path from controlled pilot to live service.",
    (
        "The application is already cloud-deployable, supports production-grade database persistence, and has a clear Azure deployment path. "
        "That gives the project a realistic route from proof of concept to pilot and then to live operational use."
    ),
    [
        "Cloud-deployable web application",
        "Supports production-grade PostgreSQL persistence",
        "Clear Azure deployment path",
        "Persistent attachment storage and exportable outputs",
        "Practical rollout path from pilot to live service",
    ],
    (
        "The aim here is confidence. "
        "You want the audience to feel this is not just an idea or design exercise, but something that can be trialled safely and matured."
    ),
    "Simple 3-step progression visual: Proof of Concept -> Pilot -> Live.",
)

add_slide(
    doc,
    13,
    "What the Future Brings",
    "Show future potential without overpromising.",
    "RadFlow creates a platform for further automation, integration, and service maturity.",
    (
        "As the platform develops, it can expand referral scanning and automation, connect more easily with external systems, "
        "and deepen reporting and decision-support capability as operational demand grows."
    ),
    [
        "Expand the Referral Scanner into a fuller OCR-enabled draft workflow",
        "Accept referrals from secure email inboxes, structured portal forms, and external systems",
        "Create draft cases for admin review before they enter the live workflow",
        "Add RIS/PACS or HL7-style intake where client environments support direct integration",
        "Deepen automation around justification workflows",
        "Extend reporting and integration capability",
        "Create a stronger foundation for broader rollout over time",
    ],
    (
        "Keep this realistic and credible. "
        "The message is that RadFlow already delivers value today, while also creating a strong base for future improvement."
    ),
    "Roadmap visual with near-term and future themes.",
)

add_slide(
    doc,
    14,
    "Closing Message",
    "End with a confident summary and clear next step.",
    "RadFlow is a better operating model for radiology vetting.",
    (
        "It supports acute and elective work, improves turnaround visibility, standardises decision workflows, "
        "and gives services a more controlled and auditable way to manage vetting. The next step is a focused pilot in a live environment."
    ),
    [
        "Supports day-to-day radiology operations",
        "Improves speed, visibility, and governance",
        "Creates a strong case for live pilot evaluation",
    ],
    (
        "Finish confidently and keep the ask simple. "
        "The purpose is to move the conversation towards a pilot that can prove value on turnaround, workload, and operational control."
    ),
    "Single bold closing statement with a clean product image.",
)

add_heading(doc, "Notes For Final Slide Design")
add_bullets(
    doc,
    [
        "Keep slides visually clean and do not overload them with too many bullets.",
        "Use screenshots on workflow slides and text-led slides on market and value slides.",
        "If challenged on MFA, describe it as a day-one implementation requirement rather than an already-live capability.",
        "If challenged on clinical rules, position them as service workflow rules that the platform can support and enforce.",
    ],
)

doc.save(OUTPUT_PATH)
print(f"Created: {OUTPUT_PATH}")
