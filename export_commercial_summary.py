from docx import Document
from docx.shared import Pt

output_path = r"c:\Users\pmend\project\Vetting app\Commercial_Summary_Client_Deck.docx"

doc = Document()

# Title
title = doc.add_paragraph()
run = title.add_run("Lumos RadFlow - Commercial Summary")
run.bold = True
run.font.size = Pt(20)

subtitle = doc.add_paragraph("Client Presentation Version")
subtitle.runs[0].italic = True


doc.add_heading("Executive Summary", level=1)
doc.add_paragraph(
    "Lumos RadFlow is a cloud-based radiology workflow platform that helps imaging teams "
    "receive, vet, assign, and track referrals faster, with stronger governance and auditability."
)

doc.add_heading("What The App Delivers", level=1)
for item in [
    "Faster referral-to-decision turnaround through structured intake and triage.",
    "Better coordination with role-based workflows for admins and radiologists.",
    "Stronger operational control with status tracking, dashboards, and exports.",
    "Safer scaling with multi-tenant architecture and organization-level access boundaries.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Core Features", level=1)
for item in [
    "Case intake and booking/case creation with required structured fields.",
    "Admin dashboard for pending, vetted, and rejected studies.",
    "Radiologist queue and review workflow.",
    "Institution and radiologist assignment controls.",
    "Case edit/reopen flows with event history.",
    "Role-based authentication and password reset.",
    "CSV export for operational reporting.",
    "Cloud deployment support on Azure with containerized delivery.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Differentiation", level=1)
doc.add_paragraph(
    "Unlike email-and-spreadsheet workflows, RadFlow connects the entire vetting lifecycle in one system: "
    "intake, triage, assignment, outcome tracking, and audit trail."
)

for item in [
    "Single accountable workflow instead of fragmented tools.",
    "Live case status visibility across the team.",
    "Clear ownership and traceable user actions.",
    "Operational standardization across multi-site organizations.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Why Clients Buy", level=1)
for item in [
    "Reduce admin effort and duplicate data entry.",
    "Improve consistency and speed of vetting decisions.",
    "Lower risk with structured governance and access controls.",
    "Scale operations without losing quality control.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Test-Environment Innovation", level=1)
doc.add_paragraph(
    "A referral parser trial flow is available in test environment to parse uploaded referral documents "
    "and pre-fill booking fields before final creation."
)
doc.add_paragraph(
    "Note: OCR for scanned image-heavy referrals remains controlled/test capability prior to full production rollout.",
    style="List Bullet",
)

doc.add_heading("30-Second Sales Pitch", level=1)
doc.add_paragraph(
    "Lumos RadFlow is a secure, cloud-native radiology vetting platform that streamlines referral intake, "
    "accelerates assignment workflows, and gives teams full case visibility from submission to outcome. "
    "It reduces admin burden, improves turnaround consistency, and supports multi-site growth with role-based governance."
)

doc.save(output_path)
print(f"Created: {output_path}")
