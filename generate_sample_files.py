#!/usr/bin/env python3
"""
Generate PDF and DOCX versions of the sample referral.
"""
from pathlib import Path
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

# Read the sample referral text
referral_path = Path(__file__).parent / "sample_referral.txt"
with open(referral_path, "r") as f:
    referral_text = f.read()

# Create PDF
pdf_path = Path(__file__).parent / "sample_referral.pdf"
doc = SimpleDocTemplate(str(pdf_path), pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)

styles = getSampleStyleSheet()
story = []

# Split by sections and add to story
for line in referral_text.split("\n"):
    line = line.rstrip()
    if not line:
        story.append(Spacer(1, 0.1*inch))
    elif line.startswith("---"):
        story.append(Spacer(1, 0.05*inch))
    else:
        # Use appropriate style based on content
        if any(line.startswith(x) for x in ["Patient", "Date", "Referral", "Urgency", "Clinical", "Requested", "Modality", "Contact"]):
            # Bold headers
            para_style = ParagraphStyle(
                'CustomHeader',
                parent=styles['Normal'],
                fontSize=11,
                textColor=colors.HexColor("#000000"),
                spaceAfter=0.1*inch,
                fontName='Helvetica-Bold'
            )
            story.append(Paragraph(line, para_style))
        else:
            story.append(Paragraph(line, styles['Normal']))

doc.build(story)
print(f"✅ PDF saved: {pdf_path}")

# Create a simple DOCX using python-docx if available, otherwise create RTF
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    docx_path = Path(__file__).parent / "sample_referral.docx"
    doc = Document()
    
    for line in referral_text.split("\n"):
        line = line.rstrip()
        if not line:
            doc.add_paragraph()
        elif line.startswith("---"):
            # Add a horizontal line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
        else:
            p = doc.add_paragraph(line)
            # Bold section headers
            if any(line.startswith(x) for x in ["Patient", "Date:", "Referral", "Urgency:", "REFERRAL", "CLINICAL", "Contact"]):
                p.runs[0].bold = True
    
    doc.save(docx_path)
    print(f"✅ Word (.docx) saved: {docx_path}")
    
except ImportError:
    print("⚠ python-docx not installed. Installing...")
    import subprocess
    subprocess.check_call(["pip", "install", "python-docx"])
    print("✅ python-docx installed. Re-run this script to generate DOCX.")
