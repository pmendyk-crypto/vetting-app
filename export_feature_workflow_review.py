from docx import Document
from docx.shared import Pt

output_path = r"c:\Users\pmend\project\Vetting app\Feature_Workflow_Review.docx"

doc = Document()

# Title
title = doc.add_paragraph()
run = title.add_run("Feature Workflow Review")
run.bold = True
run.font.size = Pt(20)

doc.add_paragraph("Project: Vetting App")
doc.add_paragraph("Generated: 2026-03-08")
doc.add_paragraph("Method: inferred from route behavior, templates, and code structure.")

# Main table
headers = ["Feature", "What it does", "Status guess", "Code location"]
rows = [
    (
        "User authentication and session login",
        "Handles sign-in/sign-out, session timeout, and role-based access entry into admin/radiologist workflows.",
        "Keep",
        "app/main.py (/login, /logout), app/security.py, templates/login.html",
    ),
    (
        "Password reset workflow",
        "Supports forgot-password and reset-password token flow with email delivery when SMTP is configured.",
        "Keep",
        "app/main.py (/forgot-password, /reset-password), templates/forgot_password.html, templates/reset_password.html",
    ),
    (
        "Admin dashboard and case monitoring",
        "Central admin workspace for filtering, viewing, and tracking case lifecycle states and workload.",
        "Keep",
        "app/main.py (/admin), templates/home.html",
    ),
    (
        "Case submission and intake",
        "Creates new cases from structured referral input including patient data, institution, modality, and notes.",
        "Keep",
        "app/main.py (/submit, /submitted/{case_id}, /intake/{org_id}), templates/submit.html, templates/submitted.html",
    ),
    (
        "File upload during case creation",
        "Uploads and stores referral attachments per case using local storage or blob backend.",
        "Keep",
        "app/main.py (submit/edit handlers, attachment endpoints), uploads/, Azure blob helpers",
    ),
    (
        "Radiologist queue and review workflow",
        "Presents assigned cases to radiologists and supports review/decision progression.",
        "Keep",
        "app/main.py (/radiologist, /vet/{case_id}), templates/radiologist_queue.html, templates/vet.html",
    ),
    (
        "Vetting decision capture",
        "Records approval/rejection decisions, comments, and vet timestamp updates on cases.",
        "Keep",
        "app/main.py (POST /vet/{case_id})",
    ),
    (
        "Case reassignment and reopen",
        "Allows admin to reassign radiologists and reopen completed/rejected workflows for correction.",
        "Keep",
        "app/main.py (/admin/case/{case_id}/assign-radiologist, /admin/case/{case_id}/reopen)",
    ),
    (
        "Settings master-data management",
        "Manages institutions, radiologists, protocols, users, and study-description presets.",
        "Keep",
        "app/main.py (/settings and /settings/*), templates/settings.html",
    ),
    (
        "Case attachment retrieval and inline preview",
        "Provides secure download/inline access to case attachments for authorized users.",
        "Keep",
        "app/main.py (/case/{case_id}/attachment, /case/{case_id}/attachment/inline)",
    ),
    (
        "PDF generation for case/timeline",
        "Builds downloadable PDF artifacts for case details and event timeline sharing.",
        "Keep",
        "app/main.py (/case/{case_id}/pdf, /admin/case/{case_id}/timeline.pdf), reportlab",
    ),
    (
        "CSV export and reporting",
        "Exports case and event datasets for operational and audit-style reporting.",
        "Keep",
        "app/main.py (/admin.csv, /admin.events.csv, /admin/case/{case_id}/timeline.csv)",
    ),
    (
        "Referral parser trial workflow",
        "Parses uploaded referral docs to prefill case fields, then allows admin confirmation/create.",
        "Review",
        "app/main.py (/submit/referral-trial*), app/referral_ingest.py, templates/referral_trial.html",
    ),
    (
        "iRefer guideline search",
        "Supports guideline lookup/search integration from app workflow context.",
        "Review",
        "app/main.py (/irefer/search)",
    ),
    (
        "Radiologist email notification flow",
        "Enables admin-initiated outbound notifications to radiologists via SMTP.",
        "Simplify",
        "app/main.py (/admin/notify-radiologist), templates/notify_radiologist.html",
    ),
    (
        "Superuser and organization governance",
        "Provides org-level administration, user management, billing views, and tenant governance actions.",
        "Review",
        "app/main.py (/mt/*, /super/*), templates/admin_orgs.html, templates/admin_users.html, templates/admin_billing.html",
    ),
    (
        "Alternate multitenant router module",
        "Contains APIRouter-based tenant login/org selection/admin routes that appear not mounted in active app path.",
        "Disable",
        "app/routers/multitenant.py (router routes), app/main.py include_router commented",
    ),
    (
        "Diagnostic and health endpoints",
        "Provides liveness and schema diagnostic visibility for operational troubleshooting.",
        "Simplify",
        "app/main.py (/health, /healthz, /diag/schema, /robots.txt)",
    ),
]

table = doc.add_table(rows=1, cols=4)
table.style = "Light Grid Accent 1"
for idx, h in enumerate(headers):
    table.rows[0].cells[idx].text = h

for feature, what_it_does, status, location in rows:
    cells = table.add_row().cells
    cells[0].text = feature
    cells[1].text = what_it_does
    cells[2].text = status
    cells[3].text = location

# Notes section

doc.add_heading("Status Guess Notes", level=1)
for note in [
    "Keep: appears core to product value and day-to-day user workflow.",
    "Simplify: useful feature, but implementation or surface area likely broader than necessary.",
    "Review: strategic feature that needs product/security/architecture decision before expansion.",
    "Disable: appears duplicated or currently non-primary path and may increase maintenance risk.",
]:
    doc.add_paragraph(note, style="List Bullet")

doc.save(output_path)
print(f"Created: {output_path}")
print(f"Features documented: {len(rows)}")
