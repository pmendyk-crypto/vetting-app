from docx import Document
from docx.shared import Pt

output_path = r"c:\Users\pmend\project\Vetting app\Technical_Overview_Project_Map.docx"

doc = Document()

title = doc.add_paragraph()
run = title.add_run("Technical Overview - Project Map")
run.bold = True
run.font.size = Pt(20)

doc.add_paragraph("Generated: 2026-03-08")

doc.add_heading("PROJECT STRUCTURE", level=1)
for item in [
    "app/: main application code (FastAPI app, routing, auth/session, DB helpers, parser, security, tenant modules).",
    "templates/: Jinja2 HTML templates for admin, submit, settings, radiologist, account, and superuser pages.",
    "static/: frontend assets (css, js, images).",
    "database/migrations/: SQL migration files and schema evolution assets.",
    "docs/: architecture and system documentation.",
    "tests/ and top-level test_*.py: automated and integration test scripts.",
    "scripts/: migration and utility scripts.",
    "uploads/: local attachment/PDF storage fallback.",
    "Top-level deployment/runtime files: Dockerfile, deploy.ps1, requirements.txt.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("CORE COMPONENTS", level=1)

doc.add_heading("API routes", level=2)
for item in [
    "Primary routes are defined in app/main.py via @app.get/post/put/delete.",
    "Coverage includes auth, admin/radiologist workflows, settings, exports, attachments, account, health, intake, referral trial, and multi-tenant/superuser routes.",
    "Secondary APIRouter exists in app/routers/multitenant.py but include_router is currently commented in main flow.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Authentication", level=2)
for item in [
    "Session-based auth through Starlette SessionMiddleware.",
    "Core auth functions and role guards are in app/main.py (verify_user, require_login, require_admin, require_radiologist, require_superuser).",
    "Password reset endpoints and token flow are in app/main.py.",
    "Additional dependency-based auth/org guards are in app/dependencies.py.",
    "Rate-limit and lockout helpers are in app/security.py.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Database layer", level=2)
for item in [
    "DB connection logic exists in app/main.py and app/db.py (SQLite default, SQLAlchemy wrapper for DATABASE_URL).",
    "Schema setup/compatibility DDL appears in app/main.py (CREATE TABLE / ALTER TABLE checks).",
    "Data models and CRUD for multi-tenant entities are defined in app/models.py dataclasses and helper functions.",
    "Most business queries are inline conn.execute(...) statements in app/main.py.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Services/business logic", level=2)
for item in [
    "Business logic is primarily embedded in app/main.py route handlers and helper functions.",
    "Referral parser extraction logic is isolated in app/referral_ingest.py.",
    "Multi-tenant service-like logic is split across app/models.py and app/dependencies.py.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Templates", level=2)
for item in [
    "Server-rendered UI templates in templates/ are used by main route handlers.",
    "Includes role-specific interfaces for admin, radiologist, and superuser views.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Storage/file handling", level=2)
for item in [
    "Local file handling via UPLOAD_DIR.",
    "Optional Azure Blob storage integration via upload/download helpers in app/main.py.",
    "Attachment and inline endpoints plus case PDF generation are implemented in app/main.py.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("TECH STACK", level=1)

doc.add_heading("Frameworks", level=2)
for item in [
    "FastAPI and Starlette middleware stack.",
    "Jinja2 server-side templating.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Libraries", level=2)
for item in [
    "uvicorn, python-multipart, itsdangerous.",
    "SQLAlchemy, psycopg2-binary.",
    "azure-storage-blob.",
    "reportlab.",
    "python-dotenv.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Database", level=2)
for item in [
    "SQLite local/default.",
    "PostgreSQL when DATABASE_URL is configured.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Hosting assumptions", level=2)
for item in [
    "Containerized deployment via Dockerfile.",
    "Azure deployment pipeline via deploy.ps1 and Azure Container Registry/App Service.",
    "Environment variable driven runtime configuration.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("RISK AREAS", level=1)
for item in [
    "Very large file risk: app/main.py is monolithic and high-change-risk.",
    "Monolithic module boundary: routing, auth, DB bootstrap, storage, email, and business logic are co-located.",
    "Duplicated logic: DB and auth/tenant logic spread across main.py, db.py, dependencies.py, and models.py.",
    "Mixed schema evolution strategy: runtime schema changes plus migration scripts can drift between environments.",
    "Refactor candidates: split routers/services, centralize DB layer, complete/clarify multi-tenant routing strategy.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.save(output_path)
print(f"Created: {output_path}")
