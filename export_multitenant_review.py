from docx import Document
from docx.shared import Pt

output_path = r"c:\Users\pmend\project\Vetting app\MultiTenant_Review.docx"

doc = Document()

title = doc.add_paragraph()
run = title.add_run("Multi-Tenant Logic Review")
run.bold = True
run.font.size = Pt(20)

doc.add_paragraph("Project: Vetting App")
doc.add_paragraph("Generated: 2026-03-08")
doc.add_paragraph("Scope: codebase analysis of organisation/tenant logic, org filtering, and org-selection behavior.")

# Section 1
doc.add_heading("MULTI-TENANT STRUCTURE", level=1)

doc.add_paragraph("How tenancy currently works:")
for item in [
    "The active runtime is hybrid: tenant-aware data model is embedded in core app logic, but the dedicated APIRouter tenant flow is commented out in main app bootstrapping.",
    "Primary org context is derived from session user plus memberships lookup (`get_current_org_context`) and then used in route-level query filters.",
    "Most admin/case/report routes in `app/main.py` apply `org_id` filtering for non-superusers (for example, case lists, case detail access, event exports, settings datasets).",
    "Superuser routes (`/mt/*`, `/super/*`) operate across organisations and manage orgs/users/memberships.",
    "Dedicated org selection routes (`/select-org`) exist in `app/routers/multitenant.py` but are likely inactive because `app.include_router(multitenant.router)` is commented out in `app/main.py`.",
    "UI forms in active templates (`submit.html`, `case_edit.html`) carry hidden `org_id` context values and use org-specific API calls for study descriptions.",
    "Database/migration layer assumes multi-org model with org-owned records and user-to-org memberships.",
]:
    doc.add_paragraph(item, style="List Bullet")


doc.add_paragraph("Database fields that reference organisation/tenant:")
for item in [
    "`organisations` table (tenant master).",
    "`memberships.org_id` and role mapping (`org_role`) for user-to-org access control.",
    "`cases.org_id` for case scoping.",
    "`institutions.org_id` for institution scoping.",
    "`protocols.org_id` for protocol scoping.",
    "`case_events.org_id` and `notify_events.org_id` for event scoping.",
    "`audit_logs.org_id` and `audit_logs.target_org_id` for admin auditing.",
    "`study_description_presets.organization_id` in preset migration/schema paths.",
    "Session-side org context keys: `user.org_id` and `current_org_id` (in multitenant dependency/router flows).",
]:
    doc.add_paragraph(item, style="List Bullet")


doc.add_paragraph("Routes that require org selection or org context:")
for item in [
    "Explicit org-selection routes: `GET /select-org`, `POST /select-org` in `app/routers/multitenant.py` (likely inactive).",
    "Superuser/tenant admin routes in active app: `/mt`, `/mt/dashboard`, `/mt/org/{org_id}/*`, `/super/orgs`, `/super/users`, `/super/billing`.",
    "Core admin/radiologist workflows in `app/main.py` rely on implicit org context from logged-in user and memberships (no separate active org switcher route in main app).",
]:
    doc.add_paragraph(item, style="List Bullet")


doc.add_paragraph("Logic that filters by organisation:")
for item in [
    "Common pattern in `app/main.py`: `if org_id and not user.get(\"is_superuser\"): ... AND ... org_id = ?`.",
    "Helper-level filtering in functions like `list_institutions(org_id)`, `list_radiologists(org_id)`, `list_protocols(..., org_id)`.",
    "Membership-aware role checks in `require_admin`/`require_radiologist` and `get_current_org_context`.",
    "Cross-org guard checks on case access/edit/reopen/assignment routes by validating case/org match.",
]:
    doc.add_paragraph(item, style="List Bullet")


doc.add_paragraph("UI components that allow org switching or org-aware interaction:")
for item in [
    "`templates/superuser_organisations.html` and `templates/superuser_org_members.html` provide organisation and membership management UI.",
    "`templates/admin_orgs.html` displays organisation administration views in active superuser routes.",
    "`templates/submit.html` and `templates/case_edit.html` store and update hidden `org_id` fields and use org-specific data fetches.",
    "No explicit end-user org switch dropdown/toggle was found in active mounted template flow; explicit selection UI is in dormant `app/routers/multitenant.py` HTML builder.",
]:
    doc.add_paragraph(item, style="List Bullet")


doc.add_paragraph("Code that assumes multiple organisations:")
for item in [
    "Schema and migration scripts create `organisations` and `memberships` as first-class entities.",
    "Superuser dashboards aggregate org counts, org member counts, and per-org billing/reporting.",
    "Many queries and helper methods are org-parameterized and include org-level constraints.",
    "Templates and form payloads include org-related fields/context.",
]:
    doc.add_paragraph(item, style="List Bullet")


# Section 2
doc.add_heading("FILES INVOLVED", level=1)
for item in [
    "`app/main.py` - primary active tenancy logic, org filtering, superuser/org routes, membership-aware access checks.",
    "`app/dependencies.py` - org-context dependencies (`require_org_context`, `require_org_admin`, `require_membership_role`, `OrgScope`).",
    "`app/models.py` - tenant models and CRUD (`Organisation`, `Membership`, org-aware queries).",
    "`app/routers/multitenant.py` - explicit org-selection and tenant route patterns (likely inactive mount).",
    "`app/db.py` - multi-tenant migration readiness helper references.",
    "`database/migrations/001_add_multi_tenant_schema.sql` - core multi-tenant schema migration.",
    "`database/migrations/002_password_reset_and_case_events.sql` - org_id in event table.",
    "`database/migrations/003_study_description_presets.sql` - `organization_id` scoped presets.",
    "`templates/admin_orgs.html` - superuser organisation listing.",
    "`templates/superuser_organisations.html` - organisation management UI.",
    "`templates/superuser_org_members.html` - membership management UI.",
    "`templates/submit.html` and `templates/case_edit.html` - org-aware hidden fields and org-scoped lookups.",
    "`docs/SYSTEM_OVERVIEW.md`, `docs/TECHNICAL_OVERVIEW_LIVING.md`, `docs/ARCHITECTURE_LIVING.md` - tenancy behavior notes.",
]:
    doc.add_paragraph(item, style="List Bullet")


# Section 3
doc.add_heading("SAFE TO DISABLE", level=1)
doc.add_paragraph("Items that can likely be removed or disabled for a single-client deployment with lower risk:")
for item in [
    "`app/routers/multitenant.py` and its `/select-org` flow, if you keep current monolithic `app/main.py` auth flow and do not mount this router.",
    "Superuser multi-tenant dashboard routes (`/mt`, `/mt/dashboard`) if platform-level org management is not required.",
    "Organisation management pages/routes (`/super/orgs`, `/super/users`, `/super/billing`) if there is only one client org and no cross-org admin use case.",
    "Template files dedicated to multi-org management (`templates/superuser_organisations.html`, `templates/superuser_org_members.html`, `templates/admin_orgs.html`) once related routes are removed.",
    "Multi-tenant reference docs (`MULTITENANT_*.md`) if you want leaner delivery docs for a single-client deployment.",
]:
    doc.add_paragraph(item, style="List Bullet")


# Section 4
doc.add_heading("REQUIRES CARE", level=1)
doc.add_paragraph("Deeply embedded parts that should remain but can be simplified carefully:")
for item in [
    "`org_id` columns on existing operational tables (`cases`, `institutions`, `protocols`, event tables). Removing these is high-risk due to broad query dependency.",
    "Membership-aware role resolution in `get_current_org_context` and `require_admin/require_radiologist` in `app/main.py`.",
    "Org filtering clauses across admin/case/report queries in `app/main.py` (these are intertwined with authorization and data isolation).",
    "Org-aware helper functions (`list_institutions`, `list_radiologists`, `list_protocols`, `get_institution`, etc.) used across active workflows.",
    "Form/template org context handling (`org_id_hidden` and org-scoped API calls) in `submit.html` and `case_edit.html`.",
    "Database migrations and indexes tied to org-aware performance and integrity; simplifying should be migration-driven, not ad-hoc code deletion.",
]:
    doc.add_paragraph(item, style="List Bullet")


doc.add_paragraph("Recommended simplification strategy for single-client mode:")
for item in [
    "Keep schema as-is initially; disable only superuser/multi-org UI and routes first.",
    "Introduce a single-client feature flag to bypass org selection while preserving existing `org_id` reads/writes.",
    "After stability period, collapse org filters in controlled migrations if truly needed.",
    "Retain tests around authorization and case visibility while simplifying to avoid accidental over-exposure.",
]:
    doc.add_paragraph(item, style="List Bullet")


doc.save(output_path)
print(f"Created: {output_path}")
