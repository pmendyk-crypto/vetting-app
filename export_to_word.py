from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# Read markdown
with open(r"c:\Users\pmend\project\Vetting app\APP_ARCHITECTURE.md", "r", encoding="utf-8") as f:
    content = f.read()

# Create Word doc
doc = Document()

# Process content line by line
lines = content.split('\n')
i = 0
in_code_block = False
code_lines = []

while i < len(lines):
    line = lines[i]
    
    # Handle code blocks
    if line.strip().startswith('```'):
        if not in_code_block:
            in_code_block = True
            code_lines = []
        else:
            in_code_block = False
            # Add code block to document
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph(code_text, style='Intense Quote')
            p.paragraph_format.left_indent = Inches(0.5)
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
        i += 1
        continue
    
    if in_code_block:
        code_lines.append(line)
        i += 1
        continue
    
    # Handle headings
    if line.startswith('# '):
        doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    elif line.startswith('#### '):
        doc.add_heading(line[5:], level=4)
    
    # Handle tables
    elif line.strip().startswith('|'):
        # Simple table parsing
        if '|' in line:
            rows = []
            j = i
            while j < len(lines) and '|' in lines[j]:
                cells = [cell.strip() for cell in lines[j].split('|')[1:-1]]
                rows.append(cells)
                j += 1
            
            if len(rows) > 0:
                # Skip separator row if present
                if len(rows) > 1 and all(c.replace('-', '').replace(':', '') == '' for c in rows[1]):
                    rows = [rows[0]] + rows[2:]
                
                # Create table
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Light Grid Accent 1'
                
                # Fill table
                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_data in enumerate(row_data):
                        table.rows[row_idx].cells[col_idx].text = cell_data
                        if row_idx == 0:
                            # Bold header
                            for paragraph in table.rows[row_idx].cells[col_idx].paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
            
            i = j
            continue
    
    # Handle empty lines
    elif line.strip() == '':
        if len(doc.paragraphs) > 0 and doc.paragraphs[-1].text.strip() != '':
            doc.add_paragraph()
    
    # Handle bullet points
    elif line.strip().startswith('- '):
        doc.add_paragraph(line.strip()[2:], style='List Bullet')
    
    # Handle regular text
    elif line.strip():
        doc.add_paragraph(line)
    
    i += 1

# Save
output_path = r"c:\Users\pmend\project\Vetting app\APP_ARCHITECTURE.docx"
doc.save(output_path)
print(f"✓ Successfully exported to: {output_path}")
