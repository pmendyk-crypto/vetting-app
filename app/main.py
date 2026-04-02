from fastapi import FastAPI, Request, UploadFile, File, Form, HTTPException
from fastapi.responses import RedirectResponse, FileResponse, HTMLResponse, StreamingResponse, JSONResponse
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles

from starlette.middleware.sessions import SessionMiddleware
from starlette.middleware.base import BaseHTTPMiddleware
from starlette.middleware.trustedhost import TrustedHostMiddleware

from pathlib import Path
from uuid import uuid4
from datetime import datetime, timezone, timedelta
import sqlite3
from sqlalchemy import create_engine, text
from sqlalchemy.exc import SQLAlchemyError
import os
import hashlib
import secrets
import mimetypes
import csv
import io
import html
import base64
import hmac
import struct
import time
import random
import string
import shutil
import re

# Security utilities
from app.security import (
    get_client_ip, check_rate_limit, reset_rate_limit,
    should_lock_account, get_lockout_until, is_account_locked, 
    get_lockout_remaining_minutes
)
from app.referral_ingest import parse_referral_attachment

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.pdfgen import canvas
import urllib.request
import urllib.parse
import json as _json

# Azure Blob Storage
try:
    from azure.storage.blob import BlobServiceClient, BlobSasPermissions, generate_blob_sas
except ImportError:
    BlobServiceClient = None
    print("[WARNING] azure-storage-blob not installed, blob storage disabled")

try:
    import qrcode
except ImportError:
    qrcode = None


# -------------------------
# Paths / App
# -------------------------
BASE_DIR = Path(__file__).resolve().parent.parent
DB_PATH = Path(os.environ.get("DB_PATH", str(BASE_DIR / "hub.db")))
UPLOAD_DIR = Path(os.environ.get("UPLOAD_DIR", str(BASE_DIR / "uploads")))
UPLOAD_DIR.mkdir(exist_ok=True, parents=True)

# Storage TTL constants (in days)
REFERRAL_FILE_TTL_DAYS = int(os.environ.get("REFERRAL_FILE_TTL_DAYS", "7"))   # delete uploaded file after 7 days
CASE_RECORD_TTL_DAYS   = int(os.environ.get("CASE_RECORD_TTL_DAYS",   "28"))  # keep case record/PDF for 28 days

# Azure Blob Storage config
AZURE_STORAGE_CONNECTION_STRING = os.environ.get("AZURE_STORAGE_CONNECTION_STRING", "")
REFERRAL_BLOB_CONTAINER = os.environ.get("REFERRAL_BLOB_CONTAINER", "referrals")
BLOB_STORAGE_ENABLED = bool(AZURE_STORAGE_CONNECTION_STRING and BlobServiceClient)

if BLOB_STORAGE_ENABLED:
    print(f"[startup] Azure Blob Storage ENABLED (container={REFERRAL_BLOB_CONTAINER})")
else:
    print("[startup] Azure Blob Storage DISABLED - using local filesystem")

# SMTP notification settings
SMTP_HOST = os.environ.get("SMTP_HOST", "")
SMTP_PORT = int(os.environ.get("SMTP_PORT", "587"))

# iRefer API settings
IREFER_API_KEY = os.environ.get("IREFER_API_KEY", "")
_irefer_guidelines_cache: list = []  # in-memory cache, populated on first request
SMTP_USER = os.environ.get("SMTP_USER", "")
SMTP_PASSWORD = os.environ.get("SMTP_PASSWORD", "") or os.environ.get("SMTP_PASS", "")
SMTP_PASS = SMTP_PASSWORD
SMTP_FROM = os.environ.get("SMTP_FROM", SMTP_USER)

# Log paths at startup for debugging persistence issues
print(f"[startup] BASE_DIR={BASE_DIR}, DB_PATH={DB_PATH}, UPLOAD_DIR={UPLOAD_DIR}")

# -------------------------
# Helper Functions
# -------------------------
def _institution_case_code(inst_name: str | None, fallback_id: int | None = None) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9 ]+", " ", (inst_name or "").strip()).upper()
    words = [w for w in cleaned.split() if w and w not in {"HOSPITAL", "CLINIC", "CENTRE", "CENTER", "TRUST", "THE"}]
    if words:
        primary = words[0]
        if len(primary) >= 3:
            return primary[:3]
        joined = "".join(words)
        if len(joined) >= 3:
            return joined[:3]
    if fallback_id is not None:
        return f"I{int(fallback_id):02d}"
    return "GEN"


def generate_case_id(institution_id: int | None = None) -> str:
    """Generate a unique readable case ID."""
    return generate_case_ids(1, institution_id=institution_id)[0]


def generate_case_ids(count: int, institution_id: int | None = None) -> list[str]:
    """Generate consecutive unique case IDs for a batch submission."""
    total = max(1, int(count or 1))
    date_prefix = datetime.now(timezone.utc).strftime("%Y%m%d")
    inst = get_institution(int(institution_id), None) if institution_id else None
    inst_code = _institution_case_code(inst.get("name") if inst else None, institution_id) if institution_id else None
    like_pattern = f"{date_prefix}-{inst_code}-%" if inst_code else f"{date_prefix}-%"

    conn = get_db()
    try:
        rows = conn.execute("SELECT id FROM cases WHERE id LIKE ?", (like_pattern,)).fetchall()
    finally:
        conn.close()

    max_seq = 0
    for row in rows or []:
        case_id = row.get("id") if isinstance(row, dict) else row[0]
        if not case_id or "-" not in case_id:
            continue
        suffix = str(case_id).rsplit("-", 1)[-1]
        if suffix.isdigit():
            max_seq = max(max_seq, int(suffix))

    if inst_code:
        return [f"{date_prefix}-{inst_code}-{max_seq + idx + 1:04d}" for idx in range(total)]
    return [f"{date_prefix}-{max_seq + idx + 1:04d}" for idx in range(total)]


def format_csv_timestamp(value: str | None) -> str:
    return format_display_datetime(value)


def clear_case_stored_filepath(case_id: str) -> None:
    if not case_id:
        return
    conn = get_db()
    conn.execute("UPDATE cases SET stored_filepath = NULL WHERE id = ?", (case_id,))
    conn.commit()
    conn.close()


def normalize_case_attachment(case_dict: dict) -> dict:
    path = case_dict.get("stored_filepath")
    if not path:
        return case_dict

    if BLOB_STORAGE_ENABLED and not str(path).startswith("/"):
        if not blob_exists(str(path)):
            clear_case_stored_filepath(case_dict.get("id"))
            case_dict["stored_filepath"] = None
        return case_dict

    if not Path(path).exists():
        clear_case_stored_filepath(case_dict.get("id"))
        case_dict["stored_filepath"] = None
    return case_dict


def is_inline_previewable(filename: str | None) -> bool:
    name = str(filename or "").strip().lower()
    if not name:
        return False
    previewable_exts = {
        ".pdf", ".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp",
        ".txt", ".csv", ".json", ".xml", ".html",
    }
    return any(name.endswith(ext) for ext in previewable_exts)


def load_case_attachment_bytes(stored_path: str | None) -> bytes | None:
    if not stored_path:
        return None
    if BLOB_STORAGE_ENABLED and stored_path and not str(stored_path).startswith("/"):
        return download_from_blob(stored_path)
    if os.path.exists(stored_path):
        with open(stored_path, "rb") as f:
            return f.read()
    return None


# -------------------------
# Azure Blob Storage Helpers
# -------------------------
def get_blob_service_client():
    """Get BlobServiceClient from connection string."""
    if not BLOB_STORAGE_ENABLED:
        return None
    try:
        return BlobServiceClient.from_connection_string(AZURE_STORAGE_CONNECTION_STRING)
    except Exception as e:
        print(f"[BLOB] Failed to create blob service client: {e}")
        return None


def upload_to_blob(case_id: str, file_bytes: bytes, original_filename: str) -> str | None:
    """
    Upload file to Azure Blob Storage.
    Returns the blob name (stored in DB as stored_filepath), or None if upload fails.
    """
    if not BLOB_STORAGE_ENABLED:
        return None
    
    try:
        # Use short blob name: case_id.ext (blob names have length limits)
        ext = Path(original_filename).suffix or ".bin"
        blob_name = f"{case_id}{ext}"
        client = get_blob_service_client()
        if not client:
            return None
        
        container_client = client.get_container_client(REFERRAL_BLOB_CONTAINER)
        container_client.upload_blob(blob_name, file_bytes, overwrite=True)
        print(f"[BLOB] Uploaded {blob_name} to {REFERRAL_BLOB_CONTAINER}")
        return blob_name  # Store blob name in DB
    except Exception as e:
        print(f"[BLOB] Upload failed for {case_id}: {e}")
        return None


def download_from_blob(blob_name: str) -> bytes | None:
    """
    Download file from Azure Blob Storage.
    Returns file bytes, or None if download fails.
    """
    if not BLOB_STORAGE_ENABLED or not blob_name:
        return None
    
    try:
        client = get_blob_service_client()
        if not client:
            return None
        
        container_client = client.get_container_client(REFERRAL_BLOB_CONTAINER)
        blob_client = container_client.get_blob_client(blob_name)
        
        download_stream = blob_client.download_blob()
        return download_stream.readall()
    except Exception as e:
        print(f"[BLOB] Download failed for {blob_name}: {e}")
        return None


def blob_exists(blob_name: str) -> bool:
    """Check if a blob exists in storage."""
    if not BLOB_STORAGE_ENABLED or not blob_name:
        return False
    
    try:
        client = get_blob_service_client()
        if not client:
            return False
        
        container_client = client.get_container_client(REFERRAL_BLOB_CONTAINER)
        blob_client = container_client.get_blob_client(blob_name)
        return blob_client.exists()
    except Exception as e:
        print(f"[BLOB] exists() check failed for {blob_name}: {e}")
        return False

app = FastAPI(title="RadFlow")

templates = Jinja2Templates(directory=str(BASE_DIR / "templates"))
app.mount("/static", StaticFiles(directory=str(BASE_DIR / "static")), name="static")

APP_ENV = (os.environ.get("APP_ENV") or os.environ.get("ENVIRONMENT") or "development").strip().lower()
IS_PRODUCTION = APP_ENV in {"production", "prod", "staging"}
DEFAULT_APP_SECRET = "dev-secret-change-me"
APP_SECRET = os.environ.get("APP_SECRET", DEFAULT_APP_SECRET)
SESSION_TIMEOUT_MINUTES = 20  # Session expires after 20 minutes of inactivity

APP_BASE_URL = os.environ.get("APP_BASE_URL", "http://localhost:8000")
DATABASE_URL = os.environ.get("DATABASE_URL", "")
SMTP_HOST = os.environ.get("SMTP_HOST")
SMTP_PORT = int(os.environ.get("SMTP_PORT", "587"))
SMTP_USER = os.environ.get("SMTP_USER")
SMTP_PASSWORD = os.environ.get("SMTP_PASSWORD") or os.environ.get("SMTP_PASS")
SMTP_PASS = SMTP_PASSWORD
SMTP_FROM = os.environ.get("SMTP_FROM") or SMTP_USER
LOGO_DARK_URL = os.environ.get("LOGO_DARK_URL", "/static/images/logo-light.png")
ALLOW_DIAGNOSTIC_ENDPOINT = (os.environ.get("ALLOW_DIAGNOSTIC_ENDPOINT") or "").strip().lower() in {"1", "true", "yes"}
COOKIE_HTTPS_ONLY = IS_PRODUCTION or APP_BASE_URL.startswith("https://")
TRUSTED_HOSTS = [host.strip() for host in (os.environ.get("TRUSTED_HOSTS") or "").split(",") if host.strip()]


def validate_security_configuration() -> None:
    issues: list[str] = []

    if IS_PRODUCTION and APP_SECRET == DEFAULT_APP_SECRET:
        issues.append("APP_SECRET must be set to a strong unique value in production.")

    if IS_PRODUCTION and not APP_BASE_URL.startswith("https://"):
        issues.append("APP_BASE_URL must use https:// in production.")

    if IS_PRODUCTION and not DATABASE_URL:
        issues.append("Production must use PostgreSQL instead of local SQLite.")

    if IS_PRODUCTION and not BLOB_STORAGE_ENABLED:
        issues.append("Production must use Azure Blob Storage for referral attachments.")

    if issues:
        raise RuntimeError("Security configuration error(s): " + " ".join(issues))

if APP_SECRET == DEFAULT_APP_SECRET:
    print("[WARNING] Using default APP_SECRET! Set APP_SECRET environment variable before any shared deployment.")

validate_security_configuration()

if TRUSTED_HOSTS:
    app.add_middleware(TrustedHostMiddleware, allowed_hosts=TRUSTED_HOSTS)

app.add_middleware(
    SessionMiddleware,
    secret_key=APP_SECRET,
    same_site="lax",
    https_only=COOKIE_HTTPS_ONLY,
    max_age=SESSION_TIMEOUT_MINUTES * 60,
)


class HTTPSRedirectMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        if IS_PRODUCTION:
            forwarded_proto = request.headers.get("x-forwarded-proto", "").split(",")[0].strip().lower()
            if request.url.scheme != "https" and forwarded_proto != "https":
                https_url = request.url.replace(scheme="https")
                return RedirectResponse(url=str(https_url), status_code=307)
        return await call_next(request)


app.add_middleware(HTTPSRedirectMiddleware)

# Middleware to add no-cache headers to authenticated pages
class NoCacheMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        response = await call_next(request)
        
        # Add no-cache headers to all responses to prevent browser caching of authenticated pages
        response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate, private"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"
        
        return response

app.add_middleware(NoCacheMiddleware)

# Middleware to add no-index headers for search engine discoverability
class NoIndexMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        try:
            response = await call_next(request)
            # Add no-index directive to prevent search engine indexing
            response.headers["X-Robots-Tag"] = "noindex, nofollow"
            return response
        except Exception as e:
            print(f"[ERROR] NoIndexMiddleware exception: {e}")
            raise

app.add_middleware(NoIndexMiddleware)


class SecurityHeadersMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        response = await call_next(request)
        response.headers["X-Content-Type-Options"] = "nosniff"
        response.headers["X-Frame-Options"] = "DENY"
        response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
        if IS_PRODUCTION:
            response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"
        return response


app.add_middleware(SecurityHeadersMiddleware)


# Global 401/403 handler — redirect to login instead of showing a raw error
@app.exception_handler(HTTPException)
async def http_exception_handler(request: Request, exc: HTTPException):
    if exc.status_code in (401, 403):
        # For AJAX / JSON requests keep the status code
        accept = request.headers.get("accept", "")
        if "application/json" in accept:
            return JSONResponse({"detail": exc.detail}, status_code=exc.status_code)
        if exc.detail == "MFA enrollment required" and get_session_user(request):
            return RedirectResponse(url="/account?msg=mfa_required", status_code=303)
        return RedirectResponse(url=f"/login?expired=1&next={request.url.path}", status_code=303)
    # For other HTTP errors return a simple styled error page
    return HTMLResponse(
        content=f"""<!DOCTYPE html><html><head><title>Error {exc.status_code}</title>
        <style>body{{font-family:sans-serif;background:#0f1724;color:#fff;display:flex;align-items:center;justify-content:center;height:100vh;margin:0;}}
        .box{{text-align:center;}}.btn{{margin-top:20px;padding:10px 20px;background:#1f6feb;color:#fff;border:none;border-radius:6px;text-decoration:none;cursor:pointer;font-size:14px;}}
        </style></head><body><div class="box"><h2>{exc.status_code}</h2><p>{exc.detail}</p>
        <a class="btn" href="/">Go Home</a>&nbsp;<a class="btn" href="/login">Login</a></div></body></html>""",
        status_code=exc.status_code,
    )


# -------------------------
# Health Check Endpoint
# -------------------------
@app.get("/health")
@app.get("/healthz")
async def health_check():
    """
    Lightweight health check endpoint.
    Returns 200 OK immediately without any external checks.
    Used for Azure App Service health monitoring and load balancer probes.
    """
    return JSONResponse(content={"status": "healthy"}, status_code=200)


@app.get("/diag/schema")
async def diagnostic_schema(request: Request):
    """
    Diagnostic endpoint to check database schema state.
    Shows which tables and key columns exist.
    """
    user = get_session_user(request)
    if not ALLOW_DIAGNOSTIC_ENDPOINT and not (user and user.get("is_superuser")):
        raise HTTPException(status_code=404, detail="Not found")
    try:
        conn = get_db()
        
        # Get all tables
        if using_postgres():
            result = conn.execute(text("""
                SELECT tablename 
                FROM pg_tables 
                WHERE schemaname = 'public' 
                ORDER BY tablename
            """))
            tables = [row[0] for row in result.fetchall()]
            
            # Check institutions columns
            result = conn.execute(text("""
                SELECT column_name 
                FROM information_schema.columns 
                WHERE table_name = 'institutions'
                ORDER BY ordinal_position
            """))
            institutions_columns = [row[0] for row in result.fetchall()]
        else:
            # SQLite
            result = conn.execute("SELECT name FROM sqlite_master WHERE type='table' ORDER BY name")
            tables = [row[0] for row in result.fetchall()]
            
            # Check institutions columns
            result = conn.execute("PRAGMA table_info(institutions)")
            institutions_columns = [row[1] for row in result.fetchall()]
        
        conn.close()
        
        return JSONResponse(content={
            "database_type": "PostgreSQL" if using_postgres() else "SQLite",
            "tables": tables,
            "institutions_columns": institutions_columns,
            "has_modified_at": "modified_at" in institutions_columns,
            "has_case_events": "case_events" in tables,
            "has_password_reset": "password_reset_tokens" in tables,
            "has_study_presets": "study_description_presets" in tables,
            "status": "ok"
        }, status_code=200)
    except Exception as e:
        return JSONResponse(content={
            "status": "error",
            "error": str(e)
        }, status_code=500)


# -------------------------
# Robots.txt Endpoint
# -------------------------
@app.get("/robots.txt", response_class=None, include_in_schema=False)
async def robots_txt():
    """
    Robots.txt endpoint to prevent search engine indexing.
    Returns plain text disallowing all crawlers.
    """
    content = """User-agent: *
Disallow: /"""
    return StreamingResponse(
        iter([content]),
        media_type="text/plain; charset=utf-8"
    )


DECISIONS = ["Approve", "Reject", "Approve with comment"]

STATUS_PENDING = "pending"
STATUS_VETTED = "vetted"
REOPEN_NOTE_MARKER = "[[REOPENED_NOTE]]"


# -------------------------
# Helpers
# -------------------------
def utc_now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _split_note_blocks(note_text: str) -> list[str]:
    text = str(note_text or "").strip()
    if not text:
        return []
    return [block.strip() for block in re.split(r"\n\s*\n", text) if block.strip()]


def partition_admin_notes(note_text: str | None) -> tuple[list[str], list[str]]:
    text = str(note_text or "").strip()
    if not text:
        return [], []

    normalized = text.replace("[REOPENED] ", REOPEN_NOTE_MARKER)
    parts = normalized.split(REOPEN_NOTE_MARKER)
    base_notes = _split_note_blocks(parts[0])
    reopened_notes: list[str] = []

    for chunk in parts[1:]:
        reopened_notes.extend(_split_note_blocks(chunk))

    return base_notes, reopened_notes


def parse_iso_dt(value: str | None) -> datetime | None:
    if not value:
        return None
    try:
        dt = datetime.fromisoformat(value.replace("Z", "+00:00"))
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        return dt.astimezone(timezone.utc)
    except Exception:
        return None


def format_display_datetime(value: str | None, fallback: str = "") -> str:
    if value is None:
        return fallback
    value_str = str(value).strip()
    if not value_str:
        return fallback
    dt = parse_iso_dt(value_str)
    if not dt:
        return value_str
    return dt.strftime("%d/%m/%Y %H:%M")


templates.env.filters["display_datetime"] = format_display_datetime


def tat_seconds(created_at: str | None, vetted_at: str | None) -> int:
    created_dt = parse_iso_dt(created_at)
    if not created_dt:
        return 0
    end_dt = parse_iso_dt(vetted_at) or datetime.now(timezone.utc)
    return max(0, int((end_dt - created_dt).total_seconds()))


def format_tat(seconds: int) -> str:
    minutes_total = seconds // 60
    hours_total = minutes_total // 60
    days = hours_total // 24
    minutes = minutes_total % 60
    hours = hours_total % 24

    if days > 0:
        return f"{days}d {hours:02}h {minutes:02}m"
    if hours_total > 0:
        return f"{hours_total:02}h {minutes:02}m"
    if minutes_total > 0:
        return f"{minutes_total}m"
    return "<1m"


# -------------------------
# DB
# -------------------------
def get_db() -> sqlite3.Connection:
    # If DATABASE_URL is set, return a SQLAlchemy-backed connection wrapper
    database_url = os.environ.get("DATABASE_URL")
    if database_url:
        # lazy create engine
        global SA_ENGINE
        if 'SA_ENGINE' not in globals():
            SA_ENGINE = create_engine(database_url)

        class SAResult:
            def __init__(self, result):
                self._result = result

            def fetchall(self):
                try:
                    return [dict(r) for r in self._result.mappings().all()]
                except Exception:
                    return []

            def fetchone(self):
                try:
                    row = self._result.mappings().first()
                    return dict(row) if row else None
                except Exception:
                    return None

        class SAConn:
            def __init__(self, engine):
                self._conn = engine.connect()
                self._trans = self._conn.begin()

            def execute(self, sql, params=None):
                # convert positional ? params to named parameters for SQLAlchemy
                if params is None:
                    params = []
                if isinstance(params, (list, tuple)) and "?" in sql:
                    # replace ? with :p0, :p1 ...
                    parts = sql.split("?")
                    named = []
                    param_map = {}
                    for i in range(len(parts) - 1):
                        name = f":p{i}"
                        named.append(parts[i] + name)
                        param_map[f"p{i}"] = params[i]
                    named.append(parts[-1])
                    sql_named = "".join(named)
                    res = self._conn.execute(text(sql_named), param_map)
                    return SAResult(res)
                else:
                    # assume dict or none
                    if isinstance(params, (list, tuple)):
                        # convert to positional mapping p0..pn
                        param_map = {f"p{i}": v for i, v in enumerate(params)}
                        return SAResult(self._conn.execute(text(sql), param_map))
                    else:
                        return SAResult(self._conn.execute(text(sql), params or {}))

            def commit(self):
                try:
                    self._trans.commit()
                except Exception:
                    pass

            def rollback(self):
                try:
                    self._trans.rollback()
                except Exception:
                    pass

            def close(self):
                try:
                    self._conn.close()
                except Exception:
                    pass

        return SAConn(SA_ENGINE)

    # default: sqlite3
    conn = sqlite3.connect(DB_PATH, timeout=30.0)
    conn.row_factory = sqlite3.Row
    try:
        conn.execute("PRAGMA journal_mode=WAL")
        conn.execute("PRAGMA busy_timeout = 30000")
    except Exception:
        pass
    return conn


def using_postgres() -> bool:
    return bool(os.environ.get("DATABASE_URL"))


def init_db() -> None:
    if using_postgres():
        conn = get_db()

        # Extended schema tables retained for current database compatibility
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS organisations (
                id SERIAL PRIMARY KEY,
                name TEXT NOT NULL,
                slug TEXT NOT NULL UNIQUE,
                is_active INTEGER NOT NULL DEFAULT 1,
                created_at TEXT NOT NULL,
                modified_at TEXT
            )
            """
        )

        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS users (
                id SERIAL PRIMARY KEY,
                username TEXT NOT NULL UNIQUE,
                email TEXT UNIQUE,
                password_hash TEXT NOT NULL,
                salt_hex TEXT NOT NULL,
                is_superuser INTEGER NOT NULL DEFAULT 0,
                is_active INTEGER NOT NULL DEFAULT 1,
                created_at TEXT NOT NULL,
                modified_at TEXT,
                first_name TEXT,
                surname TEXT
            )
            """
        )

        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS user_sessions (
                user_id INTEGER PRIMARY KEY,
                session_id TEXT NOT NULL,
                created_at TEXT NOT NULL,
                FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
            )
            """
        )

        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS memberships (
                id SERIAL PRIMARY KEY,
                org_id INTEGER NOT NULL,
                user_id INTEGER NOT NULL,
                org_role TEXT NOT NULL DEFAULT 'org_user',
                is_active INTEGER NOT NULL DEFAULT 1,
                created_at TEXT NOT NULL,
                modified_at TEXT,
                UNIQUE(org_id, user_id)
            )
            """
        )

        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS radiologist_profiles (
                id SERIAL PRIMARY KEY,
                user_id INTEGER NOT NULL UNIQUE,
                gmc TEXT,
                specialty TEXT,
                display_name TEXT,
                created_at TEXT NOT NULL,
                modified_at TEXT
            )
            """
        )

        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS audit_logs (
                id SERIAL PRIMARY KEY,
                org_id INTEGER,
                user_id INTEGER,
                action TEXT NOT NULL,
                target_user_id INTEGER,
                target_org_id INTEGER,
                details TEXT,
                created_at TEXT NOT NULL
            )
            """
        )

        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS password_reset_tokens (
                id SERIAL PRIMARY KEY,
                user_id INTEGER NOT NULL,
                token_hash TEXT NOT NULL,
                expires_at TEXT NOT NULL,
                used_at TEXT,
                created_at TEXT NOT NULL,
                requested_ip TEXT,
                requested_ua TEXT
            )
            """
        )

        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS case_events (
                id SERIAL PRIMARY KEY,
                case_id TEXT NOT NULL,
                org_id INTEGER,
                event_type TEXT NOT NULL,
                created_at TEXT NOT NULL,
                user_id INTEGER,
                username TEXT,
                org_role TEXT,
                decision TEXT,
                protocol TEXT,
                comment TEXT
            )
            """
        )

        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS notify_events (
                id SERIAL PRIMARY KEY,
                org_id INTEGER,
                radiologist_name TEXT NOT NULL,
                channel TEXT NOT NULL,
                recipient TEXT,
                message TEXT,
                created_at TEXT NOT NULL,
                created_by TEXT,
                created_by_id INTEGER
            )
            """
        )

        # Legacy operational tables (used by main app)
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS institutions (
                id SERIAL PRIMARY KEY,
                name TEXT NOT NULL UNIQUE,
                sla_hours INTEGER NOT NULL DEFAULT 48,
                created_at TEXT NOT NULL,
                modified_at TEXT,
                org_id INTEGER
            )
            """
        )

        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS cases (
                id TEXT PRIMARY KEY,
                created_at TEXT NOT NULL,
                patient_first_name TEXT NOT NULL,
                patient_surname TEXT NOT NULL,
                patient_referral_id TEXT,
                patient_dob TEXT,
                institution_id INTEGER,
                study_description TEXT NOT NULL,
                modality TEXT,
                admin_notes TEXT,
                radiologist TEXT NOT NULL,
                uploaded_filename TEXT,
                stored_filepath TEXT,
                status TEXT NOT NULL,
                protocol TEXT,
                decision TEXT,
                decision_comment TEXT,
                vetted_at TEXT,
                org_id INTEGER,
                contrast_required TEXT,
                contrast_details TEXT
            )
            """
        )

        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS protocols (
                id SERIAL PRIMARY KEY,
                name TEXT NOT NULL,
                institution_id INTEGER NOT NULL,
                study_description_preset_id INTEGER,
                instructions TEXT,
                last_modified TEXT,
                is_active INTEGER NOT NULL DEFAULT 1,
                org_id INTEGER,
                UNIQUE(name, institution_id)
            )
            """
        )

        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS radiologists (
                name TEXT PRIMARY KEY,
                first_name TEXT,
                email TEXT,
                surname TEXT,
                gmc TEXT,
                speciality TEXT
            )
            """
        )

        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS config (
                key TEXT PRIMARY KEY,
                value TEXT NOT NULL
            )
            """
        )

        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS study_description_presets (
                id SERIAL PRIMARY KEY,
                organization_id INTEGER NOT NULL,
                modality TEXT NOT NULL,
                description TEXT NOT NULL,
                is_active INTEGER NOT NULL DEFAULT 1,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL,
                created_by INTEGER NOT NULL,
                UNIQUE(organization_id, modality, description)
            )
            """
        )

        conn.execute(
            """
            CREATE INDEX IF NOT EXISTS idx_presets_org_modality
            ON study_description_presets(organization_id, modality)
            """
        )

        conn.commit()
        conn.close()
        return

    conn = get_db()

    # Institutions table
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS institutions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL UNIQUE,
            sla_hours INTEGER NOT NULL DEFAULT 48,
            created_at TEXT NOT NULL
        )
        """
    )

    # Cases table
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS cases (
            id TEXT PRIMARY KEY,
            created_at TEXT NOT NULL,
            patient_first_name TEXT NOT NULL,
            patient_surname TEXT NOT NULL,
            patient_referral_id TEXT,
            patient_dob TEXT,
            institution_id INTEGER,
            study_description TEXT NOT NULL,
            admin_notes TEXT,
            radiologist TEXT NOT NULL,
            uploaded_filename TEXT,
            stored_filepath TEXT,
            status TEXT NOT NULL,
            protocol TEXT,
            decision TEXT,
            decision_comment TEXT,
            vetted_at TEXT,
            FOREIGN KEY (institution_id) REFERENCES institutions(id)
        )
        """
    )

    # Config table
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS config (
            key TEXT PRIMARY KEY,
            value TEXT NOT NULL
        )
        """
    )

    # Radiologists table
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS radiologists (
            name TEXT PRIMARY KEY,
            first_name TEXT,
            email TEXT,
            surname TEXT,
            gmc TEXT,
            speciality TEXT
        )
        """
    )

    # Users table
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS users (
            username TEXT PRIMARY KEY,
            first_name TEXT,
            surname TEXT,
            email TEXT,
            role TEXT NOT NULL,
            radiologist_name TEXT,
            salt_hex TEXT NOT NULL,
            pw_hash_hex TEXT NOT NULL
        )
        """
    )

    # Protocols table
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS protocols (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            institution_id INTEGER NOT NULL,
            study_description_preset_id INTEGER,
            instructions TEXT,
            last_modified TEXT,
            is_active INTEGER NOT NULL DEFAULT 1,
            FOREIGN KEY (institution_id) REFERENCES institutions(id),
            UNIQUE(name, institution_id)
        )
        """
    )

    # Password reset tokens
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS password_reset_tokens (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            token_hash TEXT NOT NULL,
            expires_at TEXT NOT NULL,
            used_at TEXT,
            created_at TEXT NOT NULL,
            requested_ip TEXT,
            requested_ua TEXT
        )
        """
    )

    # Case event history
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS case_events (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            case_id TEXT NOT NULL,
            org_id INTEGER,
            event_type TEXT NOT NULL,
            created_at TEXT NOT NULL,
            user_id INTEGER,
            username TEXT,
            org_role TEXT,
            decision TEXT,
            protocol TEXT,
            comment TEXT
        )
        """
    )

    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS notify_events (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            org_id INTEGER,
            radiologist_name TEXT NOT NULL,
            channel TEXT NOT NULL,
            recipient TEXT,
            message TEXT,
            created_at TEXT NOT NULL,
            created_by TEXT,
            created_by_id INTEGER
        )
        """
    )

    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS study_description_presets (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            organization_id INTEGER NOT NULL,
            modality TEXT NOT NULL,
            description TEXT NOT NULL,
            is_active INTEGER NOT NULL DEFAULT 1,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL,
            created_by INTEGER NOT NULL,
            UNIQUE(organization_id, modality, description)
        )
        """
    )

    conn.execute(
        """
        CREATE INDEX IF NOT EXISTS idx_presets_org_modality
        ON study_description_presets(organization_id, modality)
        """
    )

    conn.commit()
    conn.close()


def ensure_cases_schema() -> None:
    """
    Safe schema upgrades for older hub.db files.
    """
    if using_postgres():
        conn = get_db()
        conn.execute("ALTER TABLE cases ADD COLUMN IF NOT EXISTS status TEXT NOT NULL DEFAULT 'pending'")
        conn.execute("ALTER TABLE cases ADD COLUMN IF NOT EXISTS vetted_at TEXT")
        conn.execute("ALTER TABLE cases ADD COLUMN IF NOT EXISTS protocol TEXT")
        conn.execute("ALTER TABLE cases ADD COLUMN IF NOT EXISTS decision TEXT")
        conn.execute("ALTER TABLE cases ADD COLUMN IF NOT EXISTS decision_comment TEXT")
        conn.execute("ALTER TABLE cases ADD COLUMN IF NOT EXISTS patient_first_name TEXT")
        conn.execute("ALTER TABLE cases ADD COLUMN IF NOT EXISTS patient_surname TEXT")
        conn.execute("ALTER TABLE cases ADD COLUMN IF NOT EXISTS patient_referral_id TEXT")
        conn.execute("ALTER TABLE cases ADD COLUMN IF NOT EXISTS patient_dob TEXT")
        conn.execute("ALTER TABLE cases ADD COLUMN IF NOT EXISTS institution_id INTEGER")
        conn.execute("ALTER TABLE cases ADD COLUMN IF NOT EXISTS modality TEXT")
        conn.execute("ALTER TABLE cases ADD COLUMN IF NOT EXISTS org_id INTEGER")
        conn.execute("ALTER TABLE cases ADD COLUMN IF NOT EXISTS contrast_required TEXT")
        conn.execute("ALTER TABLE cases ADD COLUMN IF NOT EXISTS contrast_details TEXT")
        conn.execute("ALTER TABLE cases ADD COLUMN IF NOT EXISTS uploaded_filename TEXT")
        conn.execute("ALTER TABLE cases ADD COLUMN IF NOT EXISTS stored_filepath TEXT")
        conn.execute("ALTER TABLE cases ADD COLUMN IF NOT EXISTS admin_notes TEXT")
        conn.execute("ALTER TABLE cases ADD COLUMN IF NOT EXISTS radiologist TEXT")
        conn.commit()
        conn.close()
        return
    conn = get_db()
    cur = conn.cursor()

    cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='cases'")
    if not cur.fetchone():
        conn.close()
        return

    cur.execute("PRAGMA table_info(cases)")
    cols = {row[1] for row in cur.fetchall()}

    if "status" not in cols:
        cur.execute("ALTER TABLE cases ADD COLUMN status TEXT NOT NULL DEFAULT 'pending'")
    if "vetted_at" not in cols:
        cur.execute("ALTER TABLE cases ADD COLUMN vetted_at TEXT")
    if "protocol" not in cols:
        cur.execute("ALTER TABLE cases ADD COLUMN protocol TEXT")
    if "decision" not in cols:
        cur.execute("ALTER TABLE cases ADD COLUMN decision TEXT")
    if "decision_comment" not in cols:
        cur.execute("ALTER TABLE cases ADD COLUMN decision_comment TEXT")
    if "patient_first_name" not in cols:
        cur.execute("ALTER TABLE cases ADD COLUMN patient_first_name TEXT")
    if "patient_surname" not in cols:
        cur.execute("ALTER TABLE cases ADD COLUMN patient_surname TEXT")
    if "patient_referral_id" not in cols:
        cur.execute("ALTER TABLE cases ADD COLUMN patient_referral_id TEXT")
    if "patient_dob" not in cols:
        cur.execute("ALTER TABLE cases ADD COLUMN patient_dob TEXT")
    if "institution_id" not in cols:
        cur.execute("ALTER TABLE cases ADD COLUMN institution_id INTEGER")
    if "modality" not in cols:
        cur.execute("ALTER TABLE cases ADD COLUMN modality TEXT")
    if "org_id" not in cols:
        cur.execute("ALTER TABLE cases ADD COLUMN org_id INTEGER")
    if "contrast_required" not in cols:
        cur.execute("ALTER TABLE cases ADD COLUMN contrast_required TEXT")
    if "contrast_details" not in cols:
        cur.execute("ALTER TABLE cases ADD COLUMN contrast_details TEXT")

    conn.commit()
    conn.close()


def ensure_institutions_schema() -> None:
    """
    Safe schema upgrades for the institutions table (Bug 2: Add modified_at column).
    """
    if using_postgres():
        conn = get_db()
        conn.execute("ALTER TABLE institutions ADD COLUMN IF NOT EXISTS modified_at TEXT")
        conn.execute("ALTER TABLE institutions ADD COLUMN IF NOT EXISTS org_id INTEGER")
        conn.commit()
        conn.close()
        return
    conn = get_db()
    cur = conn.cursor()

    cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='institutions'")
    if not cur.fetchone():
        conn.close()
        return

    cur.execute("PRAGMA table_info(institutions)")
    cols = {row[1] for row in cur.fetchall()}

    if "modified_at" not in cols:
        cur.execute("ALTER TABLE institutions ADD COLUMN modified_at TEXT")

    conn.commit()
    conn.close()


def ensure_radiologists_schema() -> None:
    """
    Safe schema upgrades for the radiologists table.
    """
    if using_postgres():
        conn = get_db()
        conn.execute("ALTER TABLE radiologists ADD COLUMN IF NOT EXISTS first_name TEXT")
        conn.execute("ALTER TABLE radiologists ADD COLUMN IF NOT EXISTS email TEXT")
        conn.execute("ALTER TABLE radiologists ADD COLUMN IF NOT EXISTS surname TEXT")
        conn.execute("ALTER TABLE radiologists ADD COLUMN IF NOT EXISTS gmc TEXT")
        conn.execute("ALTER TABLE radiologists ADD COLUMN IF NOT EXISTS speciality TEXT")
        conn.commit()
        conn.close()
        return
    conn = get_db()
    cur = conn.cursor()

    cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='radiologists'")
    if not cur.fetchone():
        conn.close()
        return

    cur.execute("PRAGMA table_info(radiologists)")
    cols = {row[1] for row in cur.fetchall()}

    if "first_name" not in cols:
        cur.execute("ALTER TABLE radiologists ADD COLUMN first_name TEXT")
    if "surname" not in cols:
        cur.execute("ALTER TABLE radiologists ADD COLUMN surname TEXT")
    if "gmc" not in cols:
        cur.execute("ALTER TABLE radiologists ADD COLUMN gmc TEXT")
    if "speciality" not in cols:
        cur.execute("ALTER TABLE radiologists ADD COLUMN speciality TEXT")

    conn.commit()
    conn.close()


def ensure_users_schema() -> None:
    """
    Safe schema upgrades for the users table.
    """
    if using_postgres():
        conn = get_db()
        conn.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS first_name TEXT")
        conn.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS surname TEXT")
        conn.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS email TEXT")
        conn.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS mfa_enabled INTEGER NOT NULL DEFAULT 0")
        conn.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS mfa_required INTEGER NOT NULL DEFAULT 0")
        conn.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS mfa_secret TEXT")
        conn.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS mfa_pending_secret TEXT")
        conn.commit()
        conn.close()
        return
    conn = get_db()
    cur = conn.cursor()

    cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='users'")
    if not cur.fetchone():
        conn.close()
        return

    cur.execute("PRAGMA table_info(users)")
    cols = {row[1] for row in cur.fetchall()}

    if "first_name" not in cols:
        cur.execute("ALTER TABLE users ADD COLUMN first_name TEXT")
    if "surname" not in cols:
        cur.execute("ALTER TABLE users ADD COLUMN surname TEXT")
    if "email" not in cols:
        cur.execute("ALTER TABLE users ADD COLUMN email TEXT")
    if "mfa_enabled" not in cols:
        cur.execute("ALTER TABLE users ADD COLUMN mfa_enabled INTEGER NOT NULL DEFAULT 0")
    if "mfa_required" not in cols:
        cur.execute("ALTER TABLE users ADD COLUMN mfa_required INTEGER NOT NULL DEFAULT 0")
    if "mfa_secret" not in cols:
        cur.execute("ALTER TABLE users ADD COLUMN mfa_secret TEXT")
    if "mfa_pending_secret" not in cols:
        cur.execute("ALTER TABLE users ADD COLUMN mfa_pending_secret TEXT")

    conn.commit()
    conn.close()


def ensure_protocols_schema() -> None:
    """
    Safe schema upgrades for the protocols table.
    """
    if using_postgres():
        conn = get_db()
        conn.execute("ALTER TABLE protocols ADD COLUMN IF NOT EXISTS institution_id INTEGER")
        conn.execute("ALTER TABLE protocols ADD COLUMN IF NOT EXISTS study_description_preset_id INTEGER")
        conn.execute("ALTER TABLE protocols ADD COLUMN IF NOT EXISTS instructions TEXT")
        conn.execute("ALTER TABLE protocols ADD COLUMN IF NOT EXISTS last_modified TEXT")
        conn.execute("ALTER TABLE protocols ADD COLUMN IF NOT EXISTS is_active INTEGER NOT NULL DEFAULT 1")
        conn.execute("ALTER TABLE protocols ADD COLUMN IF NOT EXISTS org_id INTEGER")
        conn.commit()
        conn.close()
        return
    conn = get_db()
    cur = conn.cursor()

    cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='protocols'")
    if not cur.fetchone():
        conn.close()
        return

    cur.execute("PRAGMA table_info(protocols)")
    cols = {row[1] for row in cur.fetchall()}

    if "institution_id" not in cols:
        cur.execute("ALTER TABLE protocols ADD COLUMN institution_id INTEGER")
    if "org_id" not in cols:
        cur.execute("ALTER TABLE protocols ADD COLUMN org_id INTEGER")
    if "study_description_preset_id" not in cols:
        cur.execute("ALTER TABLE protocols ADD COLUMN study_description_preset_id INTEGER")

    conn.commit()
    conn.close()


def ensure_study_description_presets_schema() -> None:
    if using_postgres():
        conn = get_db()
        conn.execute("ALTER TABLE study_description_presets ADD COLUMN IF NOT EXISTS is_active INTEGER NOT NULL DEFAULT 1")
        conn.commit()
        conn.close()
        return
    if not table_exists("study_description_presets"):
        return
    conn = get_db()
    cur = conn.cursor()
    cur.execute("PRAGMA table_info(study_description_presets)")
    cols = {row[1] for row in cur.fetchall()}
    if "is_active" not in cols:
        cur.execute("ALTER TABLE study_description_presets ADD COLUMN is_active INTEGER NOT NULL DEFAULT 1")
        cur.execute("UPDATE study_description_presets SET is_active = 1 WHERE is_active IS NULL")
    conn.commit()
    conn.close()


def ensure_notify_events_schema() -> None:
    if using_postgres():
        conn = get_db()
        conn.execute("ALTER TABLE notify_events ADD COLUMN IF NOT EXISTS org_id INTEGER")
        conn.execute("ALTER TABLE notify_events ADD COLUMN IF NOT EXISTS recipient TEXT")
        conn.execute("ALTER TABLE notify_events ADD COLUMN IF NOT EXISTS message TEXT")
        conn.execute("ALTER TABLE notify_events ADD COLUMN IF NOT EXISTS created_by TEXT")
        conn.execute("ALTER TABLE notify_events ADD COLUMN IF NOT EXISTS created_by_id INTEGER")
        conn.commit()
        conn.close()
        return
    conn = get_db()
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS notify_events (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            org_id INTEGER,
            radiologist_name TEXT NOT NULL,
            channel TEXT NOT NULL,
            recipient TEXT,
            message TEXT,
            created_at TEXT NOT NULL,
            created_by TEXT,
            created_by_id INTEGER
        )
        """
    )
    conn.commit()
    conn.close()


def ensure_case_events_schema() -> None:
    conn = get_db()
    if using_postgres():
        conn.execute("ALTER TABLE case_events ADD COLUMN IF NOT EXISTS org_id INTEGER")
        conn.execute("ALTER TABLE case_events ADD COLUMN IF NOT EXISTS user_id INTEGER")
        conn.execute("ALTER TABLE case_events ADD COLUMN IF NOT EXISTS username TEXT")
        conn.execute("ALTER TABLE case_events ADD COLUMN IF NOT EXISTS org_role TEXT")
        conn.execute("ALTER TABLE case_events ADD COLUMN IF NOT EXISTS decision TEXT")
        conn.execute("ALTER TABLE case_events ADD COLUMN IF NOT EXISTS protocol TEXT")
        conn.execute("ALTER TABLE case_events ADD COLUMN IF NOT EXISTS comment TEXT")
        conn.commit()
        conn.close()
        return

    if not table_exists("case_events"):
        conn.close()
        return

    cur = conn.cursor()
    cur.execute("PRAGMA table_info(case_events)")
    cols = {row[1] for row in cur.fetchall()}
    if "org_id" not in cols:
        cur.execute("ALTER TABLE case_events ADD COLUMN org_id INTEGER")
    if "user_id" not in cols:
        cur.execute("ALTER TABLE case_events ADD COLUMN user_id INTEGER")
    if "username" not in cols:
        cur.execute("ALTER TABLE case_events ADD COLUMN username TEXT")
    if "org_role" not in cols:
        cur.execute("ALTER TABLE case_events ADD COLUMN org_role TEXT")
    if "decision" not in cols:
        cur.execute("ALTER TABLE case_events ADD COLUMN decision TEXT")
    if "protocol" not in cols:
        cur.execute("ALTER TABLE case_events ADD COLUMN protocol TEXT")
    if "comment" not in cols:
        cur.execute("ALTER TABLE case_events ADD COLUMN comment TEXT")
    conn.commit()
    conn.close()


def cleanup_old_files() -> None:
    """
    Storage TTL enforcement:
    - Referral attachment files (stored_filepath) are deleted from disk after REFERRAL_FILE_TTL_DAYS days.
      The case record and DB row are kept; stored_filepath is set to NULL so the viewing link shows
      a "file no longer available" message instead of an error.
    - The case DB record itself is retained for CASE_RECORD_TTL_DAYS days before being archived/deleted.
      (Phase 1: only file deletion is implemented; DB record retention can be added later.)
    """
    if using_postgres():
        return  # Postgres / cloud deployments handle file lifecycle separately

    try:
        conn = get_db()
        cutoff_referral = (datetime.now() - timedelta(days=REFERRAL_FILE_TTL_DAYS)).isoformat()
        # Find cases whose upload file should be deleted (created more than TTL days ago and still have a stored path)
        rows = conn.execute(
            "SELECT id, stored_filepath FROM cases WHERE created_at < ? AND stored_filepath IS NOT NULL AND stored_filepath != ''",
            (cutoff_referral,)
        ).fetchall()

        deleted_count = 0
        for row in rows:
            filepath = row["stored_filepath"]
            try:
                p = Path(filepath)
                if p.exists():
                    p.unlink()
                    deleted_count += 1
            except Exception as e:
                print(f"[TTL] Could not delete file {filepath}: {e}")
            # Null out the stored path so UI shows 'unavailable' gracefully
            conn.execute(
                "UPDATE cases SET stored_filepath = NULL WHERE id = ?",
                (row["id"],)
            )

        if rows:
            conn.commit()
            print(f"[TTL] Referral file cleanup: deleted {deleted_count}/{len(rows)} files older than {REFERRAL_FILE_TTL_DAYS} days.")
        conn.close()
    except Exception as e:
        print(f"[TTL] cleanup_old_files error: {e}")


def get_setting(key: str, default: str) -> str:
    conn = get_db()
    row = conn.execute("SELECT value FROM config WHERE key = ?", (key,)).fetchone()
    conn.close()
    return row["value"] if row else default


def set_setting(key: str, value: str) -> None:
    conn = get_db()
    conn.execute(
        "INSERT INTO config(key, value) VALUES(?, ?) "
        "ON CONFLICT(key) DO UPDATE SET value=excluded.value",
        (key, value),
    )
    conn.commit()
    conn.close()


# -------------------------
# Radiologists
# -------------------------
def list_radiologists(org_id: int | None = None) -> list[dict]:
    if org_id and table_exists("memberships") and table_exists("users"):
        conn = get_db()
        rows = conn.execute(
            """
            SELECT u.username as name, u.email, u.surname, rp.gmc, rp.specialty as speciality, rp.display_name
            FROM memberships m
            JOIN users u ON m.user_id = u.id
            LEFT JOIN radiologist_profiles rp ON rp.user_id = u.id
            WHERE m.org_id = ? AND m.is_active = 1 AND m.org_role = 'radiologist'
            ORDER BY COALESCE(rp.display_name, u.username)
            """,
            (org_id,),
        ).fetchall()
        conn.close()
        result = []
        for r in rows:
            d = dict(r)
            if d.get("display_name"):
                d["name"] = d["display_name"]
            result.append(d)
        return result

    if table_exists("users"):
        conn = get_db()
        rows = conn.execute(
            """
            SELECT username, first_name, surname, email, radiologist_name
            FROM users
            WHERE role = 'radiologist'
            ORDER BY username
            """
        ).fetchall()
        conn.close()
        result = []
        for r in rows:
            d = dict(r)
            display_name = (
                d.get("radiologist_name")
                or " ".join(part for part in [d.get("first_name", "").strip(), d.get("surname", "").strip()] if part)
                or d.get("username")
            )
            result.append(
                {
                    "name": display_name,
                    "email": d.get("email") or "",
                    "surname": d.get("surname") or "",
                    "gmc": "",
                    "speciality": "",
                }
            )
        if result:
            return result

    conn = get_db()
    rows = conn.execute("SELECT name, email, surname, gmc, speciality FROM radiologists ORDER BY name").fetchall()
    conn.close()
    return [dict(r) for r in rows]


def upsert_radiologist(name: str, email: str, surname: str = "", gmc: str = "") -> None:
    conn = get_db()
    conn.execute(
        "INSERT INTO radiologists(name, email, surname, gmc) VALUES(?, ?, ?, ?) "
        "ON CONFLICT(name) DO UPDATE SET email=excluded.email, surname=excluded.surname, gmc=excluded.gmc",
        (name.strip(), email.strip(), surname.strip(), gmc.strip()),
    )
    conn.commit()
    conn.close()


def delete_radiologist(name: str) -> None:
    conn = get_db()
    conn.execute("DELETE FROM radiologists WHERE name = ?", (name.strip(),))
    conn.commit()
    conn.close()


def get_radiologist(name: str) -> dict | None:
    conn = get_db()
    row = conn.execute("SELECT name, email, surname, gmc, speciality FROM radiologists WHERE name = ?", (name,)).fetchone()
    conn.close()
    return dict(row) if row else None


# -------------------------
# Protocols
# -------------------------
DEFAULT_PROTOCOLS = [
    "CT Head (standard)",
    "CT Head (stroke)",
    "CT C-Spine",
    "CT Chest",
    "CT Abdomen/Pelvis",
    "CT KUB",
    "MRI Brain",
    "MRI Spine",
    "XR Chest",
]


def ensure_default_protocols() -> None:
    if using_postgres():
        return
    conn = get_db()
    row = conn.execute("SELECT COUNT(*) AS c FROM protocols").fetchone()
    if row and row["c"] == 0:
        for p in DEFAULT_PROTOCOLS:
            conn.execute("INSERT OR IGNORE INTO protocols(name, is_active) VALUES(?, 1)", (p,))
        conn.commit()
    conn.close()


def _load_study_presets_from_migration() -> list[tuple[str, str]]:
    full_csv_path = BASE_DIR / "database" / "migrations" / "004_study_description_presets_full.csv"
    if full_csv_path.exists():
        presets: list[tuple[str, str]] = []
        seen: set[tuple[str, str]] = set()
        try:
            with full_csv_path.open("r", encoding="utf-8", errors="ignore", newline="") as f:
                reader = csv.DictReader(f)
                for row in reader:
                    modality_clean = (row.get("modality") or "").strip().upper()
                    description_clean = (row.get("description") or "").strip()
                    if not modality_clean or not description_clean:
                        continue
                    key = (modality_clean, description_clean)
                    if key in seen:
                        continue
                    seen.add(key)
                    presets.append(key)
            if presets:
                return presets
        except Exception:
            pass

    migration_path = BASE_DIR / "database" / "migrations" / "003_study_description_presets.sql"
    if not migration_path.exists():
        return []

    try:
        sql_text = migration_path.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return []

    pattern = re.compile(r"\(\s*1\s*,\s*'([^']+)'\s*,\s*'((?:''|[^'])+)'\s*,")
    presets: list[tuple[str, str]] = []
    seen: set[tuple[str, str]] = set()

    for modality, description in pattern.findall(sql_text):
        modality_clean = modality.strip().upper()
        description_clean = description.replace("''", "'").strip()
        if not modality_clean or not description_clean:
            continue
        key = (modality_clean, description_clean)
        if key in seen:
            continue
        seen.add(key)
        presets.append(key)

    return presets


def ensure_default_study_description_presets() -> None:
    if not table_exists("study_description_presets"):
        return

    conn = get_db()
    row = conn.execute("SELECT COUNT(*) AS c FROM study_description_presets").fetchone()
    if row and row["c"]:
        conn.close()
        return

    presets = _load_study_presets_from_migration()
    if not presets:
        presets = [
            ("CT", "CT Head"),
            ("CT", "CT Thorax"),
            ("MRI", "MRI Brain"),
            ("MRI", "MRI Spine lumbar"),
            ("XR", "XR Chest"),
            ("PET", "PET FDG Whole body"),
            ("DEXA", "DXA Whole body"),
        ]

    now = utc_now_iso()
    for modality, description in presets:
        if using_postgres():
            conn.execute(
                """
                INSERT INTO study_description_presets (organization_id, modality, description, created_at, updated_at, created_by)
                VALUES (?, ?, ?, ?, ?, ?)
                ON CONFLICT (organization_id, modality, description) DO NOTHING
                """,
                (1, modality, description, now, now, 1),
            )
        else:
            conn.execute(
                """
                INSERT OR IGNORE INTO study_description_presets (organization_id, modality, description, created_at, updated_at, created_by)
                VALUES (?, ?, ?, ?, ?, ?)
                """,
                (1, modality, description, now, now, 1),
            )

    conn.commit()
    conn.close()


def list_protocols(active_only: bool = True, org_id: int | None = None) -> list[str]:
    conn = get_db()
    base_sql = "SELECT name FROM protocols"
    clauses = []
    params: list = []

    if active_only:
        clauses.append("is_active = 1")
    if org_id and table_has_column("protocols", "org_id"):
        clauses.append("org_id = ?")
        params.append(org_id)

    if clauses:
        base_sql += " WHERE " + " AND ".join(clauses)
    base_sql += " ORDER BY name"

    rows = conn.execute(base_sql, params).fetchall()
    conn.close()
    return [r["name"] for r in rows]


def list_protocol_rows(org_id: int | None = None) -> list[dict]:
    conn = get_db()
    if org_id and table_has_column("protocols", "org_id"):
        rows = conn.execute(
            "SELECT p.id, p.name, p.institution_id, p.study_description_preset_id, p.instructions, p.last_modified, p.is_active, "
            "i.name as institution_name, s.modality as study_modality, s.description as study_description "
            "FROM protocols p LEFT JOIN institutions i ON p.institution_id = i.id "
            "LEFT JOIN study_description_presets s ON p.study_description_preset_id = s.id "
            "WHERE p.org_id = ? ORDER BY p.name",
            (org_id,),
        ).fetchall()
    else:
        rows = conn.execute(
            "SELECT p.id, p.name, p.institution_id, p.study_description_preset_id, p.instructions, p.last_modified, p.is_active, "
            "i.name as institution_name, s.modality as study_modality, s.description as study_description "
            "FROM protocols p LEFT JOIN institutions i ON p.institution_id = i.id "
            "LEFT JOIN study_description_presets s ON p.study_description_preset_id = s.id "
            "ORDER BY p.name"
        ).fetchall()
    conn.close()
    result = []
    for r in rows:
        d = dict(r)
        d["last_modified"] = format_display_datetime(d.get("last_modified"), d.get("last_modified") or "")
        result.append(d)
    return result


def list_study_description_presets(org_id: int | None = None) -> list[dict]:
    conn = get_db()
    target_org_id = org_id or 1
    rows = conn.execute(
        """
        SELECT id, modality, description, is_active
        FROM study_description_presets
        WHERE organization_id = ? AND COALESCE(is_active, 1) = 1
        ORDER BY modality, description
        """,
        (target_org_id,),
    ).fetchall()
    if not rows and target_org_id != 1:
        rows = conn.execute(
            """
            SELECT id, modality, description, is_active
            FROM study_description_presets
            WHERE organization_id = 1 AND COALESCE(is_active, 1) = 1
            ORDER BY modality, description
            """
        ).fetchall()
    conn.close()
    return [dict(r) for r in rows]


def get_study_description_preset(preset_id: int, org_id: int | None = None) -> dict | None:
    conn = get_db()
    target_org_id = org_id or 1
    row = conn.execute(
        """
        SELECT id, modality, description
        FROM study_description_presets
        WHERE id = ? AND organization_id = ? AND COALESCE(is_active, 1) = 1
        LIMIT 1
        """,
        (preset_id, target_org_id),
    ).fetchone()
    if not row and target_org_id != 1:
        row = conn.execute(
            """
            SELECT id, modality, description
            FROM study_description_presets
            WHERE id = ? AND organization_id = 1 AND COALESCE(is_active, 1) = 1
            LIMIT 1
            """,
            (preset_id,),
        ).fetchone()
    conn.close()
    return dict(row) if row else None


def list_protocol_rows_for_study(
    study_description_preset_id: int,
    institution_id: int | None = None,
    org_id: int | None = None,
    active_only: bool = True,
) -> list[dict]:
    conn = get_db()
    clauses = ["p.study_description_preset_id = ?"]
    params: list = [study_description_preset_id]

    if active_only:
        clauses.append("p.is_active = 1")
    if institution_id:
        clauses.append("p.institution_id = ?")
        params.append(institution_id)
    if org_id and table_has_column("protocols", "org_id"):
        clauses.append("p.org_id = ?")
        params.append(org_id)

    rows = conn.execute(
        f"""
        SELECT p.id, p.name, p.instructions, p.institution_id, p.study_description_preset_id
        FROM protocols p
        WHERE {' AND '.join(clauses)}
        ORDER BY p.name
        """,
        params,
    ).fetchall()
    conn.close()
    return [dict(r) for r in rows]


def upsert_protocol(name: str) -> None:
    name = name.strip()
    if not name:
        return
    conn = get_db()
    if using_postgres():
        conn.execute("INSERT INTO protocols(name, is_active) VALUES(?, 1)", (name,))
    else:
        conn.execute("INSERT OR IGNORE INTO protocols(name, is_active) VALUES(?, 1)", (name,))
    conn.execute("UPDATE protocols SET is_active = 1 WHERE name = ?", (name,))
    conn.commit()
    conn.close()


def deactivate_protocol(name: str, org_id: int | None = None) -> None:
    conn = get_db()
    if org_id and table_has_column("protocols", "org_id"):
        conn.execute("UPDATE protocols SET is_active = 0 WHERE name = ? AND org_id = ?", (name.strip(), org_id))
    else:
        conn.execute("UPDATE protocols SET is_active = 0 WHERE name = ?", (name.strip(),))
    conn.commit()
    conn.close()


# -------------------------
# Users (PBKDF2)
# -------------------------
def hash_password(password: str, salt: bytes) -> bytes:
    return hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt, 200_000)


def generate_totp_secret(length: int = 20) -> str:
    return base64.b32encode(secrets.token_bytes(length)).decode("ascii").rstrip("=")


def _decode_totp_secret(secret: str) -> bytes:
    normalized = (secret or "").strip().replace(" ", "").upper()
    padding = "=" * ((8 - len(normalized) % 8) % 8)
    return base64.b32decode(normalized + padding, casefold=True)


def _totp_at(secret: str, timestamp: int | None = None, interval: int = 30, digits: int = 6) -> str:
    if timestamp is None:
        timestamp = int(time.time())
    counter = int(timestamp // interval)
    key = _decode_totp_secret(secret)
    counter_bytes = struct.pack(">Q", counter)
    digest = hmac.new(key, counter_bytes, hashlib.sha1).digest()
    offset = digest[-1] & 0x0F
    code_int = struct.unpack(">I", digest[offset:offset + 4])[0] & 0x7FFFFFFF
    return str(code_int % (10 ** digits)).zfill(digits)


def normalize_totp_code(code: str) -> str:
    return "".join(ch for ch in str(code or "") if ch.isdigit())


def verify_totp_code(secret: str | None, code: str, window: int = 1) -> bool:
    normalized_code = normalize_totp_code(code)
    if not secret or len(normalized_code) != 6:
        return False
    now = int(time.time())
    for offset in range(-window, window + 1):
        if _totp_at(secret, now + (offset * 30)) == normalized_code:
            return True
    return False


def build_totp_uri(secret: str, username: str) -> str:
    issuer = "RadFlow"
    label = urllib.parse.quote(f"{issuer}:{username}")
    issuer_param = urllib.parse.quote(issuer)
    secret_param = urllib.parse.quote(secret)
    return f"otpauth://totp/{label}?secret={secret_param}&issuer={issuer_param}&algorithm=SHA1&digits=6&period=30"


def build_totp_qr_data_uri(uri: str) -> str:
    if not uri or not qrcode:
        return ""

    qr = qrcode.QRCode(box_size=6, border=2)
    qr.add_data(uri)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    buffer = io.BytesIO()
    img.save(buffer, format="PNG")
    encoded = base64.b64encode(buffer.getvalue()).decode("ascii")
    return f"data:image/png;base64,{encoded}"


def verify_current_password(username: str, password: str) -> bool:
    if not username or not password:
        return False

    conn = get_db()
    db_user = conn.execute(
        "SELECT salt_hex, password_hash FROM users WHERE username = ?",
        (username,),
    ).fetchone()
    conn.close()

    if not db_user or not db_user["salt_hex"]:
        return False

    try:
        salt = bytes.fromhex(db_user["salt_hex"])
        expected_hash = db_user["password_hash"]
        actual_hash = hash_password(password, salt).hex()
        return secrets.compare_digest(actual_hash, expected_hash)
    except Exception:
        return False


def get_post_login_redirect_path(user: dict) -> str:
    if user.get("is_superuser"):
        return "/owner"
    if int(user.get("mfa_required") or 0) and not int(user.get("mfa_enabled") or 0):
        return "/account?msg=mfa_required"
    if user.get("role") == "admin":
        return "/admin"
    return "/radiologist"


def complete_login(request: Request, user: dict) -> RedirectResponse:
    import uuid

    session_id = str(uuid.uuid4())
    request.session.pop("pending_mfa_username", None)
    request.session["user"] = {
        "id": user.get("id"),
        "username": user["username"],
        "first_name": user.get("first_name"),
        "surname": user.get("surname"),
        "email": user.get("email"),
        "role": user["role"],
        "radiologist_name": user["radiologist_name"],
        "is_superuser": bool(user.get("is_superuser")),
        "mfa_enabled": int(user.get("mfa_enabled") or 0),
        "mfa_required": int(user.get("mfa_required") or 0),
    }
    request.session["login_time"] = time.time()
    request.session["session_id"] = session_id

    if user.get("id"):
        try:
            conn = get_db()
            try:
                conn.execute(
                    "INSERT INTO user_sessions(user_id, session_id, created_at) VALUES(?, ?, datetime('now')) ON CONFLICT(user_id) DO UPDATE SET session_id=excluded.session_id, created_at=excluded.created_at",
                    (user.get("id"), session_id)
                )
                conn.commit()
            except Exception:
                pass
            conn.close()
        except Exception:
            pass

    return RedirectResponse(url=get_post_login_redirect_path(user), status_code=303)


def create_user(username: str, password: str, role: str, radiologist_name: str | None = None, first_name: str = "", surname: str = "", email: str = "") -> None:
    username = username.strip()
    role = role.strip()
    if role not in ("admin", "radiologist", "user"):
        raise ValueError("Invalid role")
    if role == "radiologist" and not radiologist_name:
        raise ValueError("Radiologist name is required")

    salt = secrets.token_bytes(16)
    pw_hash = hash_password(password, salt)

    conn = get_db()
    conn.execute(
        """
        INSERT INTO users(username, first_name, surname, email, role, radiologist_name, salt_hex, pw_hash_hex)
        VALUES(?, ?, ?, ?, ?, ?, ?, ?)
        ON CONFLICT(username) DO UPDATE SET
          first_name=excluded.first_name,
          surname=excluded.surname,
          email=excluded.email,
          role=excluded.role,
          radiologist_name=excluded.radiologist_name,
          salt_hex=excluded.salt_hex,
          pw_hash_hex=excluded.pw_hash_hex
        """,
        (username, first_name.strip(), surname.strip(), email.strip(), role, radiologist_name, salt.hex(), pw_hash.hex()),
    )
    conn.commit()
    conn.close()


def verify_user(username: str, password: str) -> dict | None:
    normalized_username = username.strip()

    conn = get_db()
    row = conn.execute("SELECT * FROM users WHERE username = ?", (normalized_username,)).fetchone()
    conn.close()
    if not row:
        return None

    salt = bytes.fromhex(row["salt_hex"])
    
    # Support both old (pw_hash_hex) and new (password_hash) column names
    # Check which column exists in the row
    try:
        pw_hash_hex = row["password_hash"]
    except (KeyError, IndexError):
        try:
            pw_hash_hex = row["pw_hash_hex"]
        except (KeyError, IndexError):
            return None
    
    if not pw_hash_hex:
        return None
        
    expected = bytes.fromhex(pw_hash_hex)
    provided = hash_password(password, salt)
    
    if secrets.compare_digest(provided, expected):
        user_dict = dict(row)
        # Map new structure to old for backward compatibility
        row_keys = row.keys()
        stored_radiologist_name = (
            user_dict.get("radiologist_name")
            or " ".join(
                part
                for part in [
                    str(user_dict.get("first_name") or "").strip(),
                    str(user_dict.get("surname") or "").strip(),
                ]
                if part
            )
            or str(user_dict.get("username") or "").strip()
        )
        if "is_superuser" in row_keys:
            user_dict["is_superuser"] = bool(row["is_superuser"])

            if user_dict["is_superuser"]:
                user_dict["role"] = "admin"
            else:
                # Map role from active membership (only if user has id column)
                if "id" in row_keys:
                    conn = get_db()
                    membership = conn.execute(
                        "SELECT org_role FROM memberships WHERE user_id = ? AND is_active = 1 ORDER BY id LIMIT 1",
                        (row["id"],),
                    ).fetchone()
                    conn.close()

                    if membership and membership["org_role"] == "org_admin":
                        user_dict["role"] = "admin"
                    elif membership and membership["org_role"] == "radiologist":
                        user_dict["role"] = "radiologist"
                    else:
                        user_dict["role"] = "user"
                else:
                    # Fall back to role column for old schema
                    user_dict["role"] = row.get("role", "user")

            if table_exists("memberships"):
                user_dict["radiologist_name"] = None  # Will be looked up separately if needed
            else:
                user_dict["radiologist_name"] = stored_radiologist_name
        return user_dict
    return None


def list_users(org_id: int | None = None) -> list[dict]:
    conn = get_db()
    # Check which table structure we have (old vs new)
    if table_has_column("users", "is_superuser"):
        # Extended schema structure
        if org_id:
            rows = conn.execute(
                """
                SELECT 
                    u.id as user_id, u.username, u.email, u.is_superuser,
                    u.is_active, u.first_name, u.surname, COALESCE(u.mfa_enabled, 0) AS mfa_enabled,
                    COALESCE(u.mfa_required, 0) AS mfa_required,
                    m.org_role as org_role,
                    NULL as radiologist_name
                FROM users u
                INNER JOIN memberships m ON m.user_id = u.id AND m.org_id = ? AND m.is_active = 1
                WHERE u.is_superuser = 0
                ORDER BY u.username
                """,
                (org_id,),
            ).fetchall()
        else:
            rows = conn.execute(
                """
                SELECT 
                    u.id as user_id, u.username, u.email, u.is_superuser,
                    u.is_active, u.first_name, u.surname, COALESCE(u.mfa_enabled, 0) AS mfa_enabled,
                    COALESCE(u.mfa_required, 0) AS mfa_required,
                    NULL as org_role,
                    NULL as radiologist_name
                FROM users u
                ORDER BY u.username
                """
            ).fetchall()
    else:
        # Old structure
        rows = conn.execute("""
            SELECT username, first_name, surname, email, role, radiologist_name,
                   COALESCE(mfa_enabled, 0) AS mfa_enabled,
                   COALESCE(mfa_required, 0) AS mfa_required
            FROM users 
            ORDER BY username
        """).fetchall()
    
    conn.close()
    result = []
    for r in rows:
        d = dict(r)
        if "org_role" in d and d.get("org_role"):
            if d["org_role"] == "org_admin":
                d["role"] = "admin"
            elif d["org_role"] == "radiologist":
                d["role"] = "radiologist"
            else:
                d["role"] = "user"
        result.append(d)
    return result


def delete_user(username: str) -> None:
    conn = get_db()
    conn.execute("DELETE FROM users WHERE username = ?", (username.strip(),))
    conn.commit()
    conn.close()


# -------------------------
# Seed data
# -------------------------
def ensure_seed_data() -> None:
    if using_postgres():
        return
    if not get_setting("system_initialized", ""):
        set_setting("system_initialized", "true")


def ensure_local_owner_account() -> None:
    if using_postgres():
        return
    if not table_has_column("users", "is_superuser"):
        return

    owner_username = os.environ.get("OWNER_ADMIN_USERNAME", "P.Mendyk").strip() or "P.Mendyk"
    owner_password = (os.environ.get("OWNER_ADMIN_PASSWORD") or "").strip()
    owner_email = os.environ.get("OWNER_ADMIN_EMAIL", "").strip() or None
    if not owner_password:
        print("[WARNING] OWNER_ADMIN_PASSWORD is not set. Skipping local owner password bootstrap.")
        return
    now = utc_now_iso()
    salt = secrets.token_bytes(16)
    pw_hash = hash_password(owner_password, salt)

    conn = get_db()
    conn.execute(
        """
        UPDATE users
        SET is_superuser = 0,
            modified_at = ?
        WHERE COALESCE(username, '') != ?
        """,
        (now, owner_username),
    )
    existing = conn.execute("SELECT id FROM users WHERE username = ?", (owner_username,)).fetchone()
    if existing:
        conn.execute(
            """
            UPDATE users
            SET is_superuser = 1,
                is_active = 1,
                password_hash = ?,
                salt_hex = ?,
                email = ?,
                first_name = ?,
                surname = ?,
                role = COALESCE(role, 'admin'),
                modified_at = ?
            WHERE username = ?
            """,
            (pw_hash.hex(), salt.hex(), owner_email, "P", "Mendyk", now, owner_username),
        )
        conn.commit()
        conn.close()
        return

    conn.execute(
        """
        INSERT INTO users(username, email, password_hash, salt_hex, is_superuser, is_active, created_at, modified_at, first_name, surname, role, radiologist_name)
        VALUES (?, ?, ?, ?, 1, 1, ?, ?, ?, ?, 'admin', NULL)
        """,
        (
            owner_username,
            owner_email,
            pw_hash.hex(),
            salt.hex(),
            now,
            now,
            "P",
            "Mendyk",
        ),
    )
    conn.commit()
    conn.close()
def list_institutions(org_id: int | None = None) -> list[dict]:
    conn = get_db()
    if table_has_column("institutions", "org_id"):
        if org_id:
            rows = conn.execute(
                "SELECT id, name, sla_hours, created_at, modified_at, org_id FROM institutions WHERE org_id = ? ORDER BY name",
                (org_id,),
            ).fetchall()
        else:
            rows = conn.execute(
                "SELECT id, name, sla_hours, created_at, modified_at, org_id FROM institutions ORDER BY name"
            ).fetchall()
    else:
        rows = conn.execute(
            "SELECT id, name, sla_hours, created_at, modified_at FROM institutions ORDER BY name"
        ).fetchall()
    conn.close()
    result = []
    for r in rows:
        d = dict(r)
        d["created_at"] = format_display_datetime(d.get("created_at"), d.get("created_at") or "")
        d["modified_at"] = format_display_datetime(d.get("modified_at"), d.get("modified_at") or "")
        if not d.get("modified_at"):
            d["modified_at"] = d.get("created_at")  # Use created_at if modified_at not set
        result.append(d)
    return result


def get_institution(inst_id: int, org_id: int | None = None) -> dict | None:
    conn = get_db()
    if org_id and table_has_column("institutions", "org_id"):
        row = conn.execute(
            "SELECT id, name, sla_hours FROM institutions WHERE id = ? AND org_id = ?",
            (inst_id, org_id),
        ).fetchone()
    else:
        row = conn.execute("SELECT id, name, sla_hours FROM institutions WHERE id = ?", (inst_id,)).fetchone()
    conn.close()
    return dict(row) if row else None


def upsert_institution(name: str, sla_hours: int, org_id: int | None = None) -> int:
    conn = get_db()
    if org_id and table_has_column("institutions", "org_id"):
        conn.execute(
            "INSERT INTO institutions(name, sla_hours, created_at, modified_at, org_id) VALUES(?, ?, ?, ?, ?) "
            "ON CONFLICT(name) DO UPDATE SET sla_hours=excluded.sla_hours, modified_at=excluded.modified_at",
            (name.strip(), sla_hours, utc_now_iso(), utc_now_iso(), org_id),
        )
    else:
        conn.execute(
            "INSERT INTO institutions(name, sla_hours, created_at, modified_at) VALUES(?, ?, ?, ?) "
            "ON CONFLICT(name) DO UPDATE SET sla_hours=excluded.sla_hours, modified_at=excluded.modified_at",
            (name.strip(), sla_hours, utc_now_iso(), utc_now_iso()),
        )
    conn.commit()
    if org_id and table_has_column("institutions", "org_id"):
        row = conn.execute("SELECT id FROM institutions WHERE name = ? AND org_id = ?", (name.strip(), org_id)).fetchone()
    else:
        row = conn.execute("SELECT id FROM institutions WHERE name = ?", (name.strip(),)).fetchone()
    inst_id = row["id"] if row else None
    conn.close()
    return inst_id


def delete_institution(inst_id: int, org_id: int | None = None) -> None:
    conn = get_db()
    if org_id and table_has_column("institutions", "org_id"):
        conn.execute("DELETE FROM institutions WHERE id = ? AND org_id = ?", (inst_id, org_id))
    else:
        conn.execute("DELETE FROM institutions WHERE id = ?", (inst_id,))
    conn.commit()
    conn.close()


# -------------------------
# Auth helpers
# -------------------------
def get_session_user(request: Request) -> dict | None:
    """Get current user from session with expiration check and multi-window detection"""
    user = request.session.get("user")
    if not user:
        return None
    
    # Check if session has login timestamp and if it's expired
    login_time = request.session.get("login_time")
    session_id = request.session.get("session_id")
    
    if login_time:
        try:
            import time
            current_time = time.time()
            # Session timeout = 20 minutes of inactivity
            if current_time - login_time > SESSION_TIMEOUT_MINUTES * 60:
                request.session.clear()
                return None
            
            # Validate session_id for multi-window logout (detect new login from another window)
            if session_id and user.get("id"):
                try:
                    conn = get_db()
                    cur = conn.cursor()
                    cur.execute("SELECT session_id FROM user_sessions WHERE user_id = ? LIMIT 1", (user.get("id"),))
                    row = cur.fetchone()
                    conn.close()
                    if row:
                        stored_session_id = row[0] if isinstance(row, tuple) else row.get("session_id")
                        if stored_session_id and stored_session_id != session_id:
                            # User logged in from another window/browser - invalidate this session
                            request.session.clear()
                            return None
                except Exception:
                    pass  # Table might not exist for old schema, continue
            
            # Update login_time on each request (sliding window)
            request.session["login_time"] = current_time
        except Exception:
            return None
    
    return user


def require_login(request: Request) -> dict:
    user = get_session_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Not logged in")
    return user


def get_current_org_context(request: Request) -> tuple[int | None, bool, int | None, str | None]:
    user = get_session_user(request) or {}
    user_id = user.get("id")
    is_superuser = bool(user.get("is_superuser"))
    org_id = user.get("org_id")
    org_role = user.get("org_role") or user.get("role")

    if user_id and table_exists("memberships"):
        membership = get_user_primary_membership(user_id)
        if membership:
            org_id = membership.get("org_id") or org_id
            org_role = membership.get("org_role") or org_role
            if not user.get("org_id"):
                user["org_id"] = org_id
            if not user.get("org_role") and org_role:
                user["org_role"] = org_role
            request.session["user"] = user

    return user_id, is_superuser, org_id, org_role


def require_admin(request: Request) -> dict:
    user = require_login(request)
    _user_id, is_superuser, _org_id, org_role = get_current_org_context(request)
    if is_superuser:
        return user
    if table_exists("memberships"):
        if org_role in ("org_admin", "radiology_admin"):
            if int(user.get("mfa_required") or 0) and not int(user.get("mfa_enabled") or 0):
                raise HTTPException(status_code=403, detail="MFA enrollment required")
            return user
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin only")
    if int(user.get("mfa_required") or 0) and not int(user.get("mfa_enabled") or 0):
        raise HTTPException(status_code=403, detail="MFA enrollment required")
    return user


def require_superuser(request: Request) -> dict:
    user = require_admin(request)
    if not user.get("is_superuser"):
        raise HTTPException(status_code=403, detail="Superuser only")
    return user


def require_radiologist(request: Request) -> dict:
    user = require_login(request)
    _user_id, is_superuser, _org_id, org_role = get_current_org_context(request)
    if is_superuser:
        raise HTTPException(status_code=403, detail="Practitioner only")
    if table_exists("memberships"):
        if org_role != "radiologist":
            raise HTTPException(status_code=403, detail="Practitioner only")
    elif user.get("role") != "radiologist":
        raise HTTPException(status_code=403, detail="Practitioner only")

    if not user.get("radiologist_name"):
        fallback_name = (
            " ".join(
                part
                for part in [
                    str(user.get("first_name") or "").strip(),
                    str(user.get("surname") or "").strip(),
                ]
                if part
            )
            or str(user.get("username") or "").strip()
        )
        if fallback_name:
            user["radiologist_name"] = fallback_name
            request.session["user"] = user
    
    # Populate radiologist_name from radiologist profile if not already set
    if not user.get("radiologist_name") and user.get("id"):
        try:
            conn = get_db()
            rad_profile = conn.execute(
                "SELECT display_name FROM radiologist_profiles WHERE user_id = ? LIMIT 1",
                (user.get("id"),)
            ).fetchone()
            if rad_profile:
                radiologist_name = rad_profile.get("display_name") if isinstance(rad_profile, dict) else rad_profile[0]
                user["radiologist_name"] = radiologist_name
                # Also update the session
                request.session["user"] = user
            conn.close()
        except Exception:
            pass
    
    # Also try to get from radiologists table using first/last name
    if not user.get("radiologist_name") and user.get("first_name"):
        try:
            conn = get_db()
            first_name = user.get("first_name", "")
            surname = user.get("surname", "")
            full_name = f"{first_name} {surname}".strip() if surname else first_name
            
            rad_row = conn.execute(
                "SELECT name FROM radiologists WHERE name = ? OR first_name = ? LIMIT 1",
                (full_name, first_name)
            ).fetchone()
            if rad_row:
                radiologist_name = rad_row.get("name") if isinstance(rad_row, dict) else rad_row[0]
                user["radiologist_name"] = radiologist_name
                # Also update the session
                request.session["user"] = user
            conn.close()
        except Exception:
            pass
    
    return user


def generate_token() -> str:
    return secrets.token_urlsafe(32)


def hash_token(token: str) -> str:
    import hmac
    digest = hmac.new(APP_SECRET.encode("utf-8"), token.encode("utf-8"), hashlib.sha256)
    return digest.hexdigest()


def send_email(to_address: str, subject: str, body: str) -> bool:
    if not SMTP_HOST or not SMTP_FROM:
        print("[email] SMTP not configured. Message suppressed to avoid leaking email content into logs.")
        return False

    import smtplib
    import ssl
    from email.message import EmailMessage

    msg = EmailMessage()
    msg["From"] = SMTP_FROM
    msg["To"] = to_address
    msg["Subject"] = subject
    msg.set_content(body)

    server = smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=10)
    try:
        server.starttls(context=ssl.create_default_context())
        if SMTP_USER and SMTP_PASSWORD:
            server.login(SMTP_USER, SMTP_PASSWORD)
        server.send_message(msg)
        return True
    finally:
        try:
            server.quit()
        except Exception:
            pass


def get_user_by_email(email: str) -> dict | None:
    if not email:
        return None
    conn = get_db()
    if table_has_column("users", "is_active"):
        row = conn.execute(
            "SELECT * FROM users WHERE email = ? AND is_active = 1",
            (email.strip().lower(),),
        ).fetchone()
    else:
        row = conn.execute(
            "SELECT * FROM users WHERE email = ?",
            (email.strip().lower(),),
        ).fetchone()
    conn.close()
    return dict(row) if row else None


def table_exists(table_name: str) -> bool:
    conn = get_db()
    if using_postgres():
        row = conn.execute(
            "SELECT 1 FROM information_schema.tables WHERE table_schema = 'public' AND table_name = ?",
            (table_name,),
        ).fetchone()
        conn.close()
        return bool(row)

    row = conn.execute(
        "SELECT name FROM sqlite_master WHERE type='table' AND name = ?",
        (table_name,),
    ).fetchone()
    conn.close()
    return bool(row)


def table_has_column(table_name: str, column_name: str) -> bool:
    conn = get_db()
    if using_postgres():
        row = conn.execute(
            "SELECT 1 FROM information_schema.columns WHERE table_schema = 'public' AND table_name = ? AND column_name = ?",
            (table_name, column_name),
        ).fetchone()
        conn.close()
        return bool(row)

    cur = conn.cursor()
    cur.execute(f"PRAGMA table_info({table_name})")
    cols = {row[1] for row in cur.fetchall()}
    conn.close()
    return column_name in cols


def get_user_primary_membership(user_id: int) -> dict | None:
    if not table_exists("memberships"):
        return None
    conn = get_db()
    row = conn.execute(
        "SELECT org_id, org_role FROM memberships WHERE user_id = ? AND is_active = 1 ORDER BY id LIMIT 1",
        (user_id,),
    ).fetchone()
    conn.close()
    return dict(row) if row else None


def get_request_org_id(request: Request) -> int | None:
    user = get_session_user(request) or {}
    return user.get("org_id")


def redirect_to_login(role: str, next_path: str):
    return RedirectResponse(url=f"/login?role={role}&next={next_path}", status_code=303)


def slugify_org_name(name: str) -> str:
    base = re.sub(r"[^a-z0-9]+", "-", str(name or "").strip().lower()).strip("-")
    return base or "organisation"


def ensure_extended_identity_schema() -> None:
    if using_postgres():
        return

    conn = get_db()
    cur = conn.cursor()

    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS organisations (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            slug TEXT NOT NULL UNIQUE,
            is_active INTEGER NOT NULL DEFAULT 1,
            created_at TEXT NOT NULL,
            modified_at TEXT
        )
        """
    )
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS memberships (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            org_id INTEGER NOT NULL,
            user_id INTEGER NOT NULL,
            org_role TEXT NOT NULL DEFAULT 'org_user',
            is_active INTEGER NOT NULL DEFAULT 1,
            created_at TEXT NOT NULL,
            modified_at TEXT,
            UNIQUE(org_id, user_id)
        )
        """
    )
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS user_sessions (
            user_id INTEGER PRIMARY KEY,
            session_id TEXT NOT NULL,
            created_at TEXT NOT NULL
        )
        """
    )
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS radiologist_profiles (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL UNIQUE,
            gmc TEXT,
            specialty TEXT,
            display_name TEXT,
            created_at TEXT NOT NULL,
            modified_at TEXT
        )
        """
    )

    cur.execute("PRAGMA table_info(users)")
    user_cols = {row[1] for row in cur.fetchall()}
    needs_user_upgrade = {"id", "password_hash", "is_superuser", "is_active", "created_at", "modified_at"} - user_cols

    if needs_user_upgrade:
        now = utc_now_iso()
        cur.execute("DROP TABLE IF EXISTS users_extended_new")
        cur.execute(
            """
            CREATE TABLE users_extended_new (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT NOT NULL UNIQUE,
                email TEXT UNIQUE,
                password_hash TEXT NOT NULL,
                salt_hex TEXT NOT NULL,
                is_superuser INTEGER NOT NULL DEFAULT 0,
                is_active INTEGER NOT NULL DEFAULT 1,
                created_at TEXT NOT NULL,
                modified_at TEXT,
                first_name TEXT,
                surname TEXT,
                role TEXT,
                radiologist_name TEXT
            )
            """
        )
        old_rows = cur.execute(
            "SELECT username, first_name, surname, email, role, radiologist_name, salt_hex, pw_hash_hex FROM users ORDER BY username"
        ).fetchall()
        promote_username = None
        for old in old_rows:
            if str(old[4] or "").strip().lower() == "admin" and promote_username is None:
                promote_username = old[0]
        for old in old_rows:
            username = old[0]
            email = (old[3] or "").strip() or None
            role = (old[4] or "user").strip()
            is_superuser = 1 if username == promote_username else 0
            cur.execute(
                """
                INSERT INTO users_extended_new(
                    username, email, password_hash, salt_hex, is_superuser, is_active, created_at, modified_at,
                    first_name, surname, role, radiologist_name
                )
                VALUES (?, ?, ?, ?, ?, 1, ?, ?, ?, ?, ?, ?)
                """,
                (
                    username,
                    email,
                    old[7],
                    old[6],
                    is_superuser,
                    now,
                    now,
                    old[1],
                    old[2],
                    role,
                    old[5],
                ),
            )
        cur.execute("ALTER TABLE users RENAME TO users_legacy_backup")
        cur.execute("ALTER TABLE users_extended_new RENAME TO users")

    cur.execute("PRAGMA table_info(institutions)")
    institution_cols = {row[1] for row in cur.fetchall()}
    if "org_id" not in institution_cols:
        cur.execute("ALTER TABLE institutions ADD COLUMN org_id INTEGER")

    cur.execute("PRAGMA table_info(cases)")
    case_cols = {row[1] for row in cur.fetchall()}
    if "org_id" not in case_cols:
        cur.execute("ALTER TABLE cases ADD COLUMN org_id INTEGER")

    cur.execute("PRAGMA table_info(protocols)")
    protocol_cols = {row[1] for row in cur.fetchall()}
    if "org_id" not in protocol_cols:
        cur.execute("ALTER TABLE protocols ADD COLUMN org_id INTEGER")

    conn.commit()

    org_row = cur.execute("SELECT id FROM organisations ORDER BY id LIMIT 1").fetchone()
    if org_row:
        default_org_id = org_row[0]
    else:
        default_org_name = "Default Organisation"
        default_slug = slugify_org_name(default_org_name)
        suffix = 1
        while cur.execute("SELECT 1 FROM organisations WHERE slug = ?", (default_slug,)).fetchone():
            suffix += 1
            default_slug = f"{slugify_org_name(default_org_name)}-{suffix}"
        cur.execute(
            "INSERT INTO organisations(name, slug, is_active, created_at, modified_at) VALUES (?, ?, 1, ?, ?)",
            (default_org_name, default_slug, utc_now_iso(), utc_now_iso()),
        )
        default_org_id = cur.lastrowid

    cur.execute("UPDATE institutions SET org_id = ? WHERE org_id IS NULL", (default_org_id,))
    cur.execute("UPDATE cases SET org_id = ? WHERE org_id IS NULL", (default_org_id,))
    cur.execute("UPDATE protocols SET org_id = ? WHERE org_id IS NULL", (default_org_id,))

    if table_exists("study_description_presets"):
        cur.execute("UPDATE study_description_presets SET organization_id = ? WHERE organization_id IS NULL OR organization_id = 0", (default_org_id,))

    cur.execute("PRAGMA table_info(users)")
    user_cols_after_upgrade = {row[1] for row in cur.fetchall()}
    if "role" not in user_cols_after_upgrade:
        cur.execute("ALTER TABLE users ADD COLUMN role TEXT")
        user_cols_after_upgrade.add("role")
    if "radiologist_name" not in user_cols_after_upgrade:
        cur.execute("ALTER TABLE users ADD COLUMN radiologist_name TEXT")
        user_cols_after_upgrade.add("radiologist_name")

    cur.execute("PRAGMA table_info(users)")
    user_cols_after_upgrade = {row[1] for row in cur.fetchall()}
    if "role" not in user_cols_after_upgrade:
        cur.execute("ALTER TABLE users ADD COLUMN role TEXT")
        user_cols_after_upgrade.add("role")
    if "radiologist_name" not in user_cols_after_upgrade:
        cur.execute("ALTER TABLE users ADD COLUMN radiologist_name TEXT")
        user_cols_after_upgrade.add("radiologist_name")

    user_rows = cur.execute("SELECT id, username, role, first_name, surname, radiologist_name FROM users").fetchall()
    for user_row in user_rows:
        user_id = user_row[0]
        username = user_row[1]
        role = str(user_row[2] or "user").strip().lower()
        org_role = "org_admin" if role == "admin" else "radiologist" if role == "radiologist" else "org_user"
        if not cur.execute("SELECT 1 FROM memberships WHERE org_id = ? AND user_id = ?", (default_org_id, user_id)).fetchone():
            cur.execute(
                "INSERT INTO memberships(org_id, user_id, org_role, is_active, created_at, modified_at) VALUES (?, ?, ?, 1, ?, ?)",
                (default_org_id, user_id, org_role, utc_now_iso(), utc_now_iso()),
            )
        if role == "radiologist":
            display_name = (
                str(user_row[5] or "").strip()
                or " ".join(part for part in [str(user_row[3] or "").strip(), str(user_row[4] or "").strip()] if part)
                or username
            )
            if not cur.execute("SELECT 1 FROM radiologist_profiles WHERE user_id = ?", (user_id,)).fetchone():
                cur.execute(
                    "INSERT INTO radiologist_profiles(user_id, gmc, specialty, display_name, created_at, modified_at) VALUES (?, ?, ?, ?, ?, ?)",
                    (user_id, None, None, display_name, utc_now_iso(), utc_now_iso()),
                )

    conn.commit()
    conn.close()


def insert_case_event(
    case_id: str,
    org_id: int | None,
    event_type: str,
    user: dict | None = None,
    decision: str | None = None,
    protocol: str | None = None,
    comment: str | None = None,
) -> None:
    if not table_exists("case_events"):
        return

    now = utc_now_iso()
    user_id = None
    username = None
    org_role = None

    if user:
        user_id = user.get("id")
        username = user.get("username")
        org_role = user.get("org_role") or user.get("role")

    conn = get_db()
    conn.execute(
        """
        INSERT INTO case_events (case_id, org_id, event_type, created_at, user_id, username, org_role, decision, protocol, comment)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (case_id, org_id, event_type, now, user_id, username, org_role, decision, protocol, comment),
    )
    conn.commit()
    conn.close()


# -------------------------
# Init DB on startup
# -------------------------
try:
    print("[startup] Initializing database...")
    init_db()
    ensure_extended_identity_schema()
    ensure_cases_schema()
    ensure_institutions_schema()
    ensure_radiologists_schema()
    ensure_users_schema()
    ensure_protocols_schema()
    ensure_study_description_presets_schema()
    ensure_notify_events_schema()
    ensure_case_events_schema()
    ensure_seed_data()
    ensure_local_owner_account()
    ensure_default_protocols()
    ensure_default_study_description_presets()
    cleanup_old_files()
    print("[startup] Database initialization complete")
except Exception as e:
    print(f"[ERROR] Database initialization failed: {e}")
    import traceback
    traceback.print_exc()
    print("[ERROR] Application may not function correctly. Check DATABASE_URL environment variable.")


# -------------------------
# Landing + login/logout
# -------------------------
@app.get("/", response_class=HTMLResponse)
def landing(request: Request, expired: str = ""):
    user = get_session_user(request)
    if user:
        return RedirectResponse(url=get_post_login_redirect_path(user), status_code=303)
    return templates.TemplateResponse("index.html", {"request": request, "expired": expired})


@app.get("/login", response_class=HTMLResponse)
def login_page(request: Request, expired: str = ""):
    # Redirect to home if already logged in
    user = get_session_user(request)
    if user:
        return RedirectResponse(url=get_post_login_redirect_path(user), status_code=303)
    if request.session.get("pending_mfa_username"):
        return RedirectResponse(url="/login/mfa", status_code=303)
    return templates.TemplateResponse("index.html", {"request": request, "expired": expired})


@app.post("/login")
def login_submit(
    request: Request,
    username: str = Form(...),
    password: str = Form(...),
):
    # Rate limiting: max 5 login attempts per 60 seconds per IP
    client_ip = get_client_ip(request)
    is_allowed, remaining = check_rate_limit(client_ip, max_attempts=5, window_seconds=60)
    
    if not is_allowed:
        # Rate limited - return 429 Too Many Requests
        return templates.TemplateResponse(
            "index.html",
            {"request": request, "error": "Too many login attempts. Please try again in a few moments."},
            status_code=429,
        )
    
    try:
        user = verify_user(username, password)
        if not user:
            return templates.TemplateResponse(
                "index.html",
                {"request": request, "error": "Invalid username or password"},
                status_code=401,
            )
    except Exception as e:
        print(f"[ERROR] Login failed with exception: {e}")
        import traceback
        traceback.print_exc()
        return templates.TemplateResponse(
            "index.html",
            {"request": request, "error": "System error during login. Please check server logs."},
            status_code=500,
        )

    if int(user.get("mfa_enabled") or 0) and user.get("mfa_secret"):
        request.session.clear()
        request.session["pending_mfa_username"] = user["username"]
        return RedirectResponse(url="/login/mfa", status_code=303)

    return complete_login(request, user)


@app.get("/login/mfa", response_class=HTMLResponse)
def login_mfa_page(request: Request, error: str = ""):
    pending_username = request.session.get("pending_mfa_username")
    if not pending_username:
        return RedirectResponse(url="/login", status_code=303)

    return templates.TemplateResponse(
        "mfa_verify.html",
        {
            "request": request,
            "username": pending_username,
            "error": error,
        },
    )


@app.post("/login/mfa", response_class=HTMLResponse)
def login_mfa_submit(request: Request, code: str = Form(...)):
    pending_username = request.session.get("pending_mfa_username")
    if not pending_username:
        return RedirectResponse(url="/login", status_code=303)

    client_ip = get_client_ip(request)
    is_allowed, _remaining = check_rate_limit(f"mfa:{client_ip}", max_attempts=10, window_seconds=300)
    if not is_allowed:
        return templates.TemplateResponse(
            "mfa_verify.html",
            {
                "request": request,
                "username": pending_username,
                "error": "Too many authentication attempts. Please wait a few minutes and try again.",
            },
            status_code=429,
        )

    conn = get_db()
    user = conn.execute(
        "SELECT * FROM users WHERE username = ? AND COALESCE(is_active, 1) = 1",
        (pending_username,),
    ).fetchone()
    conn.close()

    if not user:
        request.session.pop("pending_mfa_username", None)
        return RedirectResponse(url="/login", status_code=303)

    user_dict = dict(user)
    if not verify_totp_code(user_dict.get("mfa_secret"), code):
        return templates.TemplateResponse(
            "mfa_verify.html",
            {
                "request": request,
                "username": pending_username,
                "error": "Invalid authentication code.",
            },
            status_code=401,
        )

    reset_rate_limit(f"mfa:{client_ip}")
    return complete_login(request, user_dict)


@app.get("/forgot-password", response_class=HTMLResponse)
def forgot_password_page(request: Request, role: str = "admin"):
    user = get_session_user(request)
    if user:
        if user.get("is_superuser"):
            return RedirectResponse(url="/owner", status_code=303)
        if user.get("role") == "admin":
            return RedirectResponse(url="/admin", status_code=303)
        return RedirectResponse(url="/radiologist", status_code=303)

    role = (role or "admin").strip().lower()
    if role not in ("admin", "radiologist"):
        role = "admin"

    return templates.TemplateResponse(
        "forgot_password.html",
        {"request": request, "role": role, "submitted": False, "email_failed": False},
    )


@app.post("/forgot-password", response_class=HTMLResponse)
def forgot_password_submit(request: Request, role: str = Form("admin"), email: str = Form(...)):
    role = (role or "admin").strip().lower()
    if role not in ("admin", "radiologist"):
        role = "admin"

    email = (email or "").strip().lower()
    user = get_user_by_email(email)

    email_sent = False

    if user and user.get("id"):
        token = generate_token()
        token_hash = hash_token(token)
        now = utc_now_iso()
        expires_at = (datetime.now(timezone.utc) + timedelta(minutes=60)).isoformat()

        requested_ip = request.client.host if request.client else None
        requested_ua = request.headers.get("user-agent")

        conn = get_db()
        conn.execute(
            """
            INSERT INTO password_reset_tokens (user_id, token_hash, expires_at, used_at, created_at, requested_ip, requested_ua)
            VALUES (?, ?, ?, NULL, ?, ?, ?)
            """,
            (user.get("id"), token_hash, expires_at, now, requested_ip, requested_ua),
        )
        conn.commit()
        conn.close()

        reset_link = f"{APP_BASE_URL}/reset-password?token={token}"
        email_sent = send_email(
            to_address=email,
            subject="RadFlow password reset",
            body=(
                "You requested a password reset for your RadFlow account.\n\n"
                f"Reset your password using this link (valid for 60 minutes):\n{reset_link}\n\n"
                "If you did not request this, you can ignore this email."
            ),
        )

    return templates.TemplateResponse(
        "forgot_password.html",
        {
            "request": request,
            "role": role,
            "submitted": True,
            "email_failed": bool(user and user.get("id") and not email_sent),
        },
    )


@app.get("/reset-password", response_class=HTMLResponse)
def reset_password_page(request: Request, token: str = ""):
    token = (token or "").strip()
    if not token:
        return templates.TemplateResponse(
            "reset_password.html",
            {"request": request, "token": "", "error": "Invalid or expired reset link."},
            status_code=400,
        )

    token_hash = hash_token(token)
    conn = get_db()
    row = conn.execute(
        "SELECT * FROM password_reset_tokens WHERE token_hash = ?",
        (token_hash,),
    ).fetchone()
    conn.close()

    if not row:
        return templates.TemplateResponse(
            "reset_password.html",
            {"request": request, "token": "", "error": "Invalid or expired reset link."},
            status_code=400,
        )

    if row.get("used_at"):
        return templates.TemplateResponse(
            "reset_password.html",
            {"request": request, "token": "", "error": "Reset link already used."},
            status_code=400,
        )

    expires_at = parse_iso_dt(row.get("expires_at"))
    if not expires_at or expires_at < datetime.now(timezone.utc):
        return templates.TemplateResponse(
            "reset_password.html",
            {"request": request, "token": "", "error": "Reset link expired."},
            status_code=400,
        )

    return templates.TemplateResponse(
        "reset_password.html",
        {"request": request, "token": token},
    )


@app.post("/reset-password", response_class=HTMLResponse)
def reset_password_submit(
    request: Request,
    token: str = Form(""),
    password: str = Form(""),
    confirm_password: str = Form(""),
):
    token = (token or "").strip()
    if not token or not password or not confirm_password:
        return templates.TemplateResponse(
            "reset_password.html",
            {"request": request, "token": token, "error": "Token, password, and confirmation are required."},
            status_code=400,
        )

    if password != confirm_password:
        return templates.TemplateResponse(
            "reset_password.html",
            {"request": request, "token": token, "error": "Passwords do not match."},
            status_code=400,
        )

    if len(password) < 8:
        return templates.TemplateResponse(
            "reset_password.html",
            {"request": request, "token": token, "error": "Password must be at least 8 characters long."},
            status_code=400,
        )

    token_hash = hash_token(token)
    conn = get_db()
    row = conn.execute(
        "SELECT * FROM password_reset_tokens WHERE token_hash = ?",
        (token_hash,),
    ).fetchone()

    if not row or row.get("used_at"):
        conn.close()
        return templates.TemplateResponse(
            "reset_password.html",
            {"request": request, "token": "", "error": "Invalid or expired reset link."},
            status_code=400,
        )

    expires_at = parse_iso_dt(row.get("expires_at"))
    if not expires_at or expires_at < datetime.now(timezone.utc):
        conn.close()
        return templates.TemplateResponse(
            "reset_password.html",
            {"request": request, "token": "", "error": "Reset link expired."},
            status_code=400,
        )

    user_id = row.get("user_id")
    if not user_id:
        conn.close()
        return templates.TemplateResponse(
            "reset_password.html",
            {"request": request, "token": "", "error": "Invalid or expired reset link."},
            status_code=400,
        )

    salt = secrets.token_bytes(16)
    pw_hash = hash_password(password, salt)
    now = utc_now_iso()

    if table_has_column("users", "password_hash"):
        conn.execute(
            "UPDATE users SET password_hash = ?, salt_hex = ?, modified_at = ? WHERE id = ?",
            (pw_hash.hex(), salt.hex(), now, user_id),
        )
    else:
        conn.close()
        return templates.TemplateResponse(
            "reset_password.html",
            {"request": request, "token": "", "error": "Password reset not supported for this user schema."},
            status_code=400,
        )

    conn.execute(
        "UPDATE password_reset_tokens SET used_at = ? WHERE id = ?",
        (now, row.get("id")),
    )
    conn.commit()
    conn.close()

    return RedirectResponse(url="/?reset=success", status_code=303)


@app.get("/logout")
def logout(request: Request):
    request.session.clear()
    response = RedirectResponse(url="/", status_code=303)
    # Add headers to prevent caching of authenticated pages
    response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate, private"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    return response


@app.get("/owner", response_class=HTMLResponse)
def owner_dashboard(request: Request, created: str = "", error: str = ""):
    user = require_superuser(request)
    organisations = list_organisations_summary()
    return templates.TemplateResponse(
        "owner_dashboard.html",
        {
            "request": request,
            "user": user,
            "organisations": organisations,
            "created": created,
            "error": error,
            "form_data": {
                "org_name": "",
                "org_slug": "",
                "admin_first_name": "",
                "admin_surname": "",
                "admin_email": "",
                "admin_username": "",
                "admin_mfa_required": 0,
            },
        },
    )


@app.post("/owner/organisations")
def owner_create_organisation(
    request: Request,
    org_name: str = Form(...),
    org_slug: str = Form(""),
    admin_first_name: str = Form(...),
    admin_surname: str = Form(""),
    admin_email: str = Form(""),
    admin_username: str = Form(...),
    admin_password: str = Form(...),
    admin_mfa_required: str = Form("0"),
):
    user = require_superuser(request)

    org_name = org_name.strip()
    requested_slug = (org_slug or "").strip()
    admin_first_name = admin_first_name.strip()
    admin_surname = admin_surname.strip()
    admin_email = admin_email.strip()
    admin_username = admin_username.strip()
    admin_password = admin_password.strip()
    admin_mfa_required_value = 1 if str(admin_mfa_required).strip().lower() in {"1", "true", "on", "yes"} else 0

    form_data = {
        "org_name": org_name,
        "org_slug": requested_slug,
        "admin_first_name": admin_first_name,
        "admin_surname": admin_surname,
        "admin_email": admin_email,
        "admin_username": admin_username,
        "admin_mfa_required": admin_mfa_required_value,
    }

    if not org_name or not admin_first_name or not admin_username or not admin_password:
        return templates.TemplateResponse(
            "owner_dashboard.html",
            {
                "request": request,
                "user": user,
                "organisations": list_organisations_summary(),
                "created": "",
                "error": "Organisation name, first admin first name, username, and password are required.",
                "form_data": form_data,
            },
            status_code=400,
        )

    slug = slugify_org_name(requested_slug or org_name)
    now = utc_now_iso()
    salt = secrets.token_bytes(16)
    pw_hash = hash_password(admin_password, salt)

    conn = get_db()
    try:
        existing_slug = conn.execute("SELECT id FROM organisations WHERE slug = ?", (slug,)).fetchone()
        if existing_slug:
            return templates.TemplateResponse(
                "owner_dashboard.html",
                {
                    "request": request,
                    "user": user,
                    "organisations": list_organisations_summary(),
                    "created": "",
                    "error": "That organisation code is already in use.",
                    "form_data": form_data,
                },
                status_code=400,
            )

        existing_user = conn.execute("SELECT 1 FROM users WHERE username = ?", (admin_username,)).fetchone()
        if existing_user:
            return templates.TemplateResponse(
                "owner_dashboard.html",
                {
                    "request": request,
                    "user": user,
                    "organisations": list_organisations_summary(),
                    "created": "",
                    "error": "That admin username is already in use.",
                    "form_data": form_data,
                },
                status_code=400,
            )

        if admin_email:
            existing_email = conn.execute("SELECT 1 FROM users WHERE email = ?", (admin_email,)).fetchone()
            if existing_email:
                return templates.TemplateResponse(
                    "owner_dashboard.html",
                    {
                        "request": request,
                        "user": user,
                        "organisations": list_organisations_summary(),
                        "created": "",
                        "error": "That admin email is already in use.",
                        "form_data": form_data,
                    },
                    status_code=400,
                )

        conn.execute(
            "INSERT INTO organisations(name, slug, is_active, created_at, modified_at) VALUES (?, ?, 1, ?, ?)",
            (org_name, slug, now, now),
        )
        org_row = conn.execute("SELECT id FROM organisations WHERE slug = ?", (slug,)).fetchone()
        org_id = org_row["id"] if isinstance(org_row, dict) else org_row[0]

        if table_has_column("users", "role"):
            conn.execute(
                """
                INSERT INTO users(username, email, password_hash, salt_hex, is_superuser, is_active, created_at, modified_at, first_name, surname, role, radiologist_name)
                VALUES (?, ?, ?, ?, 0, 1, ?, ?, ?, ?, ?, NULL)
                """,
                (
                    admin_username,
                    admin_email or None,
                    pw_hash.hex(),
                    salt.hex(),
                    now,
                    now,
                    admin_first_name,
                    admin_surname,
                    "admin",
                ),
            )
        else:
            conn.execute(
                """
                INSERT INTO users(username, email, password_hash, salt_hex, is_superuser, is_active, created_at, modified_at, first_name, surname)
                VALUES (?, ?, ?, ?, 0, 1, ?, ?, ?, ?)
                """,
                (
                    admin_username,
                    admin_email or None,
                    pw_hash.hex(),
                    salt.hex(),
                    now,
                    now,
                    admin_first_name,
                    admin_surname,
                ),
            )
        conn.execute(
            "UPDATE users SET mfa_required = ? WHERE username = ?",
            (admin_mfa_required_value, admin_username),
        )
        user_row = conn.execute("SELECT id FROM users WHERE username = ?", (admin_username,)).fetchone()
        user_id = user_row["id"] if isinstance(user_row, dict) else user_row[0]

        conn.execute(
            """
            INSERT INTO memberships(org_id, user_id, org_role, is_active, created_at, modified_at)
            VALUES (?, ?, 'org_admin', 1, ?, ?)
            """,
            (org_id, user_id, now, now),
        )

        if table_exists("study_description_presets"):
            seed_rows = conn.execute(
                """
                SELECT modality, description
                FROM study_description_presets
                WHERE organization_id = 1 AND COALESCE(is_active, 1) = 1
                ORDER BY modality, description
                """
            ).fetchall()
            for seed in seed_rows:
                if using_postgres():
                    conn.execute(
                        """
                        INSERT INTO study_description_presets(
                            organization_id, modality, description, is_active, created_at, updated_at, created_by
                        )
                        VALUES (?, ?, ?, 1, ?, ?, ?)
                        ON CONFLICT (organization_id, modality, description) DO NOTHING
                        """,
                        (
                            org_id,
                            seed["modality"] if isinstance(seed, dict) else seed[0],
                            seed["description"] if isinstance(seed, dict) else seed[1],
                            now,
                            now,
                            user.get("id") or user_id,
                        ),
                    )
                else:
                    conn.execute(
                        """
                        INSERT OR IGNORE INTO study_description_presets(
                            organization_id, modality, description, is_active, created_at, updated_at, created_by
                        )
                        VALUES (?, ?, ?, 1, ?, ?, ?)
                        """,
                        (
                            org_id,
                            seed["modality"] if isinstance(seed, dict) else seed[0],
                            seed["description"] if isinstance(seed, dict) else seed[1],
                            now,
                            now,
                            user.get("id") or user_id,
                        ),
                    )

        conn.commit()
    except Exception as exc:
        conn.rollback()
        return templates.TemplateResponse(
            "owner_dashboard.html",
            {
                "request": request,
                "user": user,
                "organisations": list_organisations_summary(),
                "created": "",
                "error": f"Unable to create organisation: {exc}",
                "form_data": form_data,
            },
            status_code=400,
        )
    finally:
        conn.close()

    return RedirectResponse(url="/owner?created=1", status_code=303)


@app.get("/owner/organisations/{org_id}", response_class=HTMLResponse)
def owner_edit_organisation_page(request: Request, org_id: int, saved: str = "", error: str = ""):
    user = require_superuser(request)
    organisation = get_organisation_summary(org_id)
    if not organisation:
        raise HTTPException(status_code=404, detail="Organisation not found")
    return templates.TemplateResponse(
        "owner_organisation_edit.html",
        {
            "request": request,
            "user": user,
            "organisation": organisation,
            "org_users": list_organisation_users(org_id),
            "org_institutions": list_organisation_institutions(org_id),
            "saved": saved,
            "notice": request.query_params.get("notice", ""),
            "error": error,
        },
    )


@app.post("/owner/organisations/{org_id}")
def owner_edit_organisation_submit(
    request: Request,
    org_id: int,
    name: str = Form(...),
    slug: str = Form(...),
    is_active: str = Form("1"),
):
    user = require_superuser(request)
    organisation = get_organisation_summary(org_id)
    if not organisation:
        raise HTTPException(status_code=404, detail="Organisation not found")

    clean_name = name.strip()
    clean_slug = slugify_org_name(slug.strip() or clean_name)
    active_value = 1 if str(is_active).strip() == "1" else 0

    if not clean_name:
        return templates.TemplateResponse(
            "owner_organisation_edit.html",
            {
                "request": request,
                "user": user,
                "organisation": {
                    **organisation,
                    "name": clean_name,
                    "slug": clean_slug,
                    "is_active": active_value,
                },
                "saved": "",
                "error": "Organisation name is required.",
            },
            status_code=400,
        )

    conn = get_db()
    try:
        existing = conn.execute(
            "SELECT id FROM organisations WHERE slug = ? AND id != ?",
            (clean_slug, org_id),
        ).fetchone()
        if existing:
            return templates.TemplateResponse(
                "owner_organisation_edit.html",
                {
                    "request": request,
                    "user": user,
                    "organisation": {
                        **organisation,
                        "name": clean_name,
                        "slug": clean_slug,
                        "is_active": active_value,
                    },
                    "saved": "",
                    "error": "That organisation code is already in use.",
                },
                status_code=400,
            )

        conn.execute(
            """
            UPDATE organisations
            SET name = ?, slug = ?, is_active = ?, modified_at = ?
            WHERE id = ?
            """,
            (clean_name, clean_slug, active_value, utc_now_iso(), org_id),
        )
        conn.commit()
    finally:
        conn.close()

    return RedirectResponse(url=f"/owner/organisations/{org_id}?saved=1", status_code=303)


@app.post("/owner/organisations/{org_id}/users/add")
def owner_add_organisation_user(
    request: Request,
    org_id: int,
    username: str = Form(...),
    password: str = Form(...),
    role: str = Form(...),
    first_name: str = Form(""),
    surname: str = Form(""),
    email: str = Form(""),
    gmc: str = Form(""),
    speciality: str = Form(""),
    mfa_required: str = Form("0"),
):
    require_superuser(request)
    organisation = get_organisation_summary(org_id)
    if not organisation:
        raise HTTPException(status_code=404, detail="Organisation not found")

    username = username.strip()
    password = password.strip()
    role = role.strip()
    first_name = first_name.strip()
    surname = surname.strip()
    email = email.strip()
    gmc = gmc.strip()
    speciality = speciality.strip()
    mfa_required_value = 1 if str(mfa_required).strip().lower() in {"1", "true", "on", "yes"} else 0

    if not username or not password or role not in {"admin", "radiologist", "user"}:
        return templates.TemplateResponse(
            "owner_organisation_edit.html",
            {
                "request": request,
                "user": get_session_user(request),
                "organisation": organisation,
                "org_users": list_organisation_users(org_id),
                "saved": "",
                "notice": "",
                "error": "Username, password, and a valid role are required.",
            },
            status_code=400,
        )

    if role == "radiologist" and not first_name:
        return templates.TemplateResponse(
            "owner_organisation_edit.html",
            {
                "request": request,
                "user": get_session_user(request),
                "organisation": organisation,
                "org_users": list_organisation_users(org_id),
                "saved": "",
                "notice": "",
                "error": "Practitioner accounts need at least a first name.",
            },
            status_code=400,
        )

    salt = secrets.token_bytes(16)
    pw_hash = hash_password(password, salt)
    now = utc_now_iso()
    email_val = email or None
    org_role = "org_admin" if role == "admin" else "radiologist" if role == "radiologist" else "org_user"

    conn = get_db()
    try:
        if conn.execute("SELECT 1 FROM users WHERE username = ?", (username,)).fetchone():
            raise HTTPException(status_code=400, detail="That username is already in use.")
        if email_val and conn.execute("SELECT 1 FROM users WHERE email = ?", (email_val,)).fetchone():
            raise HTTPException(status_code=400, detail="That email address is already in use.")

        conn.execute(
            """
            INSERT INTO users(username, email, password_hash, salt_hex, is_superuser, is_active, created_at, modified_at, first_name, surname, role, radiologist_name)
            VALUES (?, ?, ?, ?, 0, 1, ?, ?, ?, ?, ?, ?)
            """,
            (
                username,
                email_val,
                pw_hash.hex(),
                salt.hex(),
                now,
                now,
                first_name,
                surname,
                "admin" if role == "admin" else role,
                None if role != "radiologist" else (f"{first_name} {surname}".strip() or username),
            ),
        )
        conn.execute(
            "UPDATE users SET mfa_required = ? WHERE username = ?",
            (mfa_required_value, username),
        )
        user_row = conn.execute("SELECT id FROM users WHERE username = ?", (username,)).fetchone()
        user_id = user_row["id"] if isinstance(user_row, dict) else user_row[0]
        conn.execute(
            """
            INSERT INTO memberships(org_id, user_id, org_role, is_active, created_at, modified_at)
            VALUES (?, ?, ?, 1, ?, ?)
            """,
            (org_id, user_id, org_role, now, now),
        )

        if role == "radiologist" and table_exists("radiologist_profiles"):
            display_name = f"{first_name} {surname}".strip() or username
            existing_profile = conn.execute(
                "SELECT id FROM radiologist_profiles WHERE user_id = ?",
                (user_id,),
            ).fetchone()
            if existing_profile:
                conn.execute(
                    """
                    UPDATE radiologist_profiles
                    SET gmc = ?, specialty = ?, display_name = ?, modified_at = ?
                    WHERE user_id = ?
                    """,
                    (gmc or None, speciality or None, display_name, now, user_id),
                )
            else:
                conn.execute(
                    """
                    INSERT INTO radiologist_profiles(user_id, gmc, specialty, display_name, created_at, modified_at)
                    VALUES(?, ?, ?, ?, ?, ?)
                    """,
                    (user_id, gmc or None, speciality or None, display_name, now, now),
                )

        conn.commit()
    except HTTPException as exc:
        conn.rollback()
        conn.close()
        return templates.TemplateResponse(
            "owner_organisation_edit.html",
            {
                "request": request,
                "user": get_session_user(request),
                "organisation": organisation,
                "org_users": list_organisation_users(org_id),
                "saved": "",
                "notice": "",
                "error": exc.detail,
            },
            status_code=400,
        )
    except Exception as exc:
        conn.rollback()
        conn.close()
        return templates.TemplateResponse(
            "owner_organisation_edit.html",
            {
                "request": request,
                "user": get_session_user(request),
                "organisation": organisation,
                "org_users": list_organisation_users(org_id),
                "saved": "",
                "notice": "",
                "error": f"Unable to add user: {exc}",
            },
            status_code=400,
        )
    conn.close()

    return RedirectResponse(url=f"/owner/organisations/{org_id}?notice=user_created", status_code=303)


@app.post("/owner/organisations/{org_id}/users/{user_id}/edit")
def owner_edit_organisation_user(
    request: Request,
    org_id: int,
    user_id: int,
    first_name: str = Form(""),
    surname: str = Form(""),
    email: str = Form(""),
    role: str = Form(...),
    is_active: str = Form("1"),
    gmc: str = Form(""),
    speciality: str = Form(""),
    mfa_required: str = Form("0"),
):
    require_superuser(request)

    role = role.strip()
    if role not in {"admin", "radiologist", "user"}:
        raise HTTPException(status_code=400, detail="Invalid role")

    conn = get_db()
    row = conn.execute(
        """
        SELECT u.id, u.username, u.email
        FROM memberships m
        INNER JOIN users u ON u.id = m.user_id
        WHERE m.org_id = ? AND m.user_id = ? AND m.is_active = 1
        """,
        (org_id, user_id),
    ).fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="User not found in that organisation")

    clean_email = email.strip()
    current_email = row["email"] if isinstance(row, dict) else row[2]
    if clean_email and clean_email != current_email:
        conflict = conn.execute("SELECT 1 FROM users WHERE email = ? AND id != ?", (clean_email, user_id)).fetchone()
        if conflict:
            conn.close()
            raise HTTPException(status_code=400, detail="That email address is already in use.")

    now = utc_now_iso()
    active_value = 1 if str(is_active).strip() == "1" else 0
    mfa_required_value = 1 if str(mfa_required).strip().lower() in {"1", "true", "on", "yes"} else 0
    org_role = "org_admin" if role == "admin" else "radiologist" if role == "radiologist" else "org_user"
    display_name = f"{first_name.strip()} {surname.strip()}".strip() or (row["username"] if isinstance(row, dict) else row[1])

    conn.execute(
        """
        UPDATE users
        SET first_name = ?, surname = ?, email = ?, role = ?, is_active = ?, radiologist_name = ?, modified_at = ?
        WHERE id = ?
        """,
        (
            first_name.strip(),
            surname.strip(),
            clean_email or None,
            "admin" if role == "admin" else role,
            active_value,
            display_name if role == "radiologist" else None,
            now,
            user_id,
        ),
    )
    if mfa_required_value:
        conn.execute(
            "UPDATE users SET mfa_required = ?, modified_at = ? WHERE id = ?",
            (1, now, user_id),
        )
    else:
        conn.execute(
            """
            UPDATE users
            SET mfa_required = 0, mfa_enabled = 0, mfa_secret = NULL, mfa_pending_secret = NULL, modified_at = ?
            WHERE id = ?
            """,
            (now, user_id),
        )
    conn.execute(
        """
        UPDATE memberships
        SET org_role = ?, modified_at = ?
        WHERE org_id = ? AND user_id = ? AND is_active = 1
        """,
        (org_role, now, org_id, user_id),
    )

    if table_exists("radiologist_profiles"):
        if role == "radiologist":
            profile_row = conn.execute("SELECT id FROM radiologist_profiles WHERE user_id = ?", (user_id,)).fetchone()
            if profile_row:
                conn.execute(
                    """
                    UPDATE radiologist_profiles
                    SET gmc = ?, specialty = ?, display_name = ?, modified_at = ?
                    WHERE user_id = ?
                    """,
                    (gmc.strip() or None, speciality.strip() or None, display_name, now, user_id),
                )
            else:
                conn.execute(
                    """
                    INSERT INTO radiologist_profiles(user_id, gmc, specialty, display_name, created_at, modified_at)
                    VALUES(?, ?, ?, ?, ?, ?)
                    """,
                    (user_id, gmc.strip() or None, speciality.strip() or None, display_name, now, now),
                )
        else:
            conn.execute("DELETE FROM radiologist_profiles WHERE user_id = ?", (user_id,))

    conn.commit()
    conn.close()
    return RedirectResponse(url=f"/owner/organisations/{org_id}?notice=user_updated", status_code=303)


@app.post("/owner/organisations/{org_id}/users/{user_id}/reset-password")
def owner_reset_organisation_user_password(
    request: Request,
    org_id: int,
    user_id: int,
    password: str = Form(...),
):
    require_superuser(request)
    password = password.strip()
    if not password:
        raise HTTPException(status_code=400, detail="Password is required")

    conn = get_db()
    membership = conn.execute(
        "SELECT 1 FROM memberships WHERE org_id = ? AND user_id = ? AND is_active = 1",
        (org_id, user_id),
    ).fetchone()
    if not membership:
        conn.close()
        raise HTTPException(status_code=404, detail="User not found in that organisation")

    salt = secrets.token_bytes(16)
    pw_hash = hash_password(password, salt)
    conn.execute(
        "UPDATE users SET password_hash = ?, salt_hex = ?, modified_at = ? WHERE id = ?",
        (pw_hash.hex(), salt.hex(), utc_now_iso(), user_id),
    )
    conn.commit()
    conn.close()
    return RedirectResponse(url=f"/owner/organisations/{org_id}?notice=password_reset", status_code=303)


@app.post("/owner/organisations/{org_id}/users/{user_id}/delete")
def owner_delete_organisation_user(request: Request, org_id: int, user_id: int):
    current_user = require_superuser(request)
    conn = get_db()
    row = conn.execute(
        """
        SELECT u.id, u.username
        FROM memberships m
        INNER JOIN users u ON u.id = m.user_id
        WHERE m.org_id = ? AND m.user_id = ? AND m.is_active = 1
        """,
        (org_id, user_id),
    ).fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="User not found in that organisation")

    username = row["username"] if isinstance(row, dict) else row[1]
    if username == current_user.get("username"):
        conn.close()
        raise HTTPException(status_code=400, detail="You cannot delete your own owner account from here.")

    conn.execute("DELETE FROM memberships WHERE org_id = ? AND user_id = ?", (org_id, user_id))
    remaining = conn.execute(
        "SELECT COUNT(*) AS c FROM memberships WHERE user_id = ? AND is_active = 1",
        (user_id,),
    ).fetchone()
    remaining_count = remaining["c"] if isinstance(remaining, dict) else remaining[0]
    if remaining_count == 0:
        if table_exists("radiologist_profiles"):
            conn.execute("DELETE FROM radiologist_profiles WHERE user_id = ?", (user_id,))
        conn.execute("DELETE FROM users WHERE id = ?", (user_id,))
    conn.commit()
    conn.close()
    return RedirectResponse(url=f"/owner/organisations/{org_id}?notice=user_deleted", status_code=303)


@app.post("/owner/organisations/{org_id}/institutions/{inst_id}/delete")
def owner_delete_organisation_institution(request: Request, org_id: int, inst_id: int):
    require_superuser(request)
    inst = get_institution(inst_id, org_id)
    if not inst:
        raise HTTPException(status_code=404, detail="Institution not found")

    conn = get_db()
    if table_exists("cases"):
        conn.execute(
            "UPDATE cases SET institution_id = NULL WHERE institution_id = ? AND org_id = ?",
            (inst_id, org_id),
        )
    if table_exists("protocols"):
        conn.execute(
            "DELETE FROM protocols WHERE institution_id = ? AND org_id = ?",
            (inst_id, org_id),
        )
    conn.commit()
    conn.close()

    delete_institution(inst_id, org_id)
    return RedirectResponse(url=f"/owner/organisations/{org_id}?notice=institution_deleted", status_code=303)


@app.post("/owner/organisations/{org_id}/delete")
def owner_delete_organisation(request: Request, org_id: int):
    current_user = require_superuser(request)
    organisation = get_organisation_summary(org_id)
    if not organisation:
        raise HTTPException(status_code=404, detail="Organisation not found")

    conn = get_db()
    member_rows = conn.execute(
        "SELECT DISTINCT user_id FROM memberships WHERE org_id = ?",
        (org_id,),
    ).fetchall()
    member_ids = [row["user_id"] if isinstance(row, dict) else row[0] for row in member_rows]

    if table_exists("cases"):
        conn.execute("DELETE FROM cases WHERE org_id = ?", (org_id,))
    if table_exists("protocols"):
        conn.execute("DELETE FROM protocols WHERE org_id = ?", (org_id,))
    if table_exists("institutions"):
        conn.execute("DELETE FROM institutions WHERE org_id = ?", (org_id,))
    if table_exists("study_description_presets"):
        conn.execute("DELETE FROM study_description_presets WHERE organization_id = ?", (org_id,))
    if table_exists("notify_events"):
        conn.execute("DELETE FROM notify_events WHERE org_id = ?", (org_id,))
    if table_exists("case_events"):
        conn.execute("DELETE FROM case_events WHERE org_id = ?", (org_id,))
    if table_exists("memberships"):
        conn.execute("DELETE FROM memberships WHERE org_id = ?", (org_id,))

    for member_id in member_ids:
        if member_id == current_user.get("id"):
            continue
        remaining = conn.execute(
            "SELECT COUNT(*) AS c FROM memberships WHERE user_id = ? AND is_active = 1",
            (member_id,),
        ).fetchone()
        remaining_count = remaining["c"] if isinstance(remaining, dict) else remaining[0]
        if remaining_count == 0:
            if table_exists("radiologist_profiles"):
                conn.execute("DELETE FROM radiologist_profiles WHERE user_id = ?", (member_id,))
            conn.execute("DELETE FROM users WHERE id = ?", (member_id,))

    conn.execute("DELETE FROM organisations WHERE id = ?", (org_id,))
    conn.commit()
    conn.close()
    return RedirectResponse(url="/owner?created=deleted", status_code=303)


# -------------------------
# Admin dashboard (tabs + filters + TAT)
# -------------------------
def build_admin_case_filters(
    org_id,
    is_superuser: bool,
    institution: str | None = None,
    radiologist: str | None = None,
    modality: str | None = None,
    q: str | None = None,
    status: str | None = None,
    created_since: str | None = None,
    created_on: str | None = None,
    created_from: str | None = None,
    created_to: str | None = None,
):
    clauses = ["1=1"]
    params: list = []
    has_modality_column = table_has_column("cases", "modality")

    if org_id and not is_superuser:
        clauses.append("c.org_id = ?")
        params.append(org_id)

    if status:
        clauses.append("LOWER(c.status) = ?")
        params.append(status.strip().lower())

    if institution and institution.strip():
        clauses.append("c.institution_id = ?")
        params.append(int(institution))

    if radiologist and radiologist.strip():
        clauses.append("c.radiologist = ?")
        params.append(radiologist.strip())

    if has_modality_column and modality and modality.strip():
        clauses.append("UPPER(COALESCE(c.modality, '')) = ?")
        params.append(modality.strip().upper())

    if q and q.strip():
        like = f"%{q.strip()}%"
        clauses.append(
            "(c.id LIKE ? OR c.patient_first_name LIKE ? OR c.patient_surname LIKE ? "
            "OR c.patient_referral_id LIKE ? OR c.study_description LIKE ?)"
        )
        params.extend([like, like, like, like, like])

    if created_since:
        clauses.append("c.created_at >= ?")
        params.append(created_since)

    if created_on:
        clauses.append("SUBSTR(c.created_at, 1, 10) = ?")
        params.append(created_on)

    if created_from:
        clauses.append("SUBSTR(c.created_at, 1, 10) >= ?")
        params.append(created_from)

    if created_to:
        clauses.append("SUBSTR(c.created_at, 1, 10) <= ?")
        params.append(created_to)

    return clauses, params


def get_admin_org_name(org_id):
    org_name = "Organisation Name"
    if org_id:
        conn = get_db()
        org_row = conn.execute("SELECT name FROM organisations WHERE id = ?", (org_id,)).fetchone()
        if org_row:
            org_name = org_row["name"]
        conn.close()
    return org_name


def build_dashboard_filter_summary(
    dashboard_range: str,
    dashboard_date_from: str | None,
    dashboard_date_to: str | None,
    dashboard_institution: str | None,
    dashboard_radiologist: str | None,
    institutions: list[dict],
) -> list[str]:
    quick_range_map = {
        "7d": "Last 7 days",
        "30d": "Last 30 days",
        "90d": "Last 90 days",
        "365d": "Last 12 months",
        "all": "All time",
    }
    institution_name = "All institutions"
    if dashboard_institution:
        for inst in institutions:
            if str(inst.get("id")) == str(dashboard_institution):
                institution_name = str(inst.get("name") or institution_name)
                break

    if dashboard_date_from and dashboard_date_to:
        if dashboard_date_from == dashboard_date_to:
            date_label = dashboard_date_from
        else:
            date_label = f"{dashboard_date_from} to {dashboard_date_to}"
    else:
        date_label = quick_range_map.get(dashboard_range, "Last 30 days")

    return [
        f"Date scope: {date_label}",
        f"Institution: {institution_name}",
        f"Practitioner: {dashboard_radiologist or 'All practitioners'}",
    ]


def list_case_modalities(org_id, is_superuser: bool):
    if not table_has_column("cases", "modality"):
        return []

    sql = "SELECT DISTINCT UPPER(TRIM(modality)) AS modality FROM cases c WHERE modality IS NOT NULL AND TRIM(modality) != ''"
    params: list = []
    if org_id and not is_superuser:
        sql += " AND c.org_id = ?"
        params.append(org_id)
    sql += " ORDER BY modality"
    conn = get_db()
    rows = conn.execute(sql, params).fetchall()
    conn.close()
    return [r["modality"] for r in rows if r["modality"]]


def build_dashboard_series(rows: list[dict]):
    status_counts = {"pending": 0, "vetted": 0, "rejected": 0, "reopened": 0}
    over_time_counts: dict[str, dict] = {}
    institution_counts: dict[str, int] = {}
    radiologist_counts: dict[str, int] = {}
    avg_tat_values: list[int] = []
    completed_tat_values: list[int] = []
    unassigned_cases = 0
    sla_breaches = 0

    for row in rows:
        status_key = str(row.get("status") or "").strip().lower()
        if status_key in status_counts:
            status_counts[status_key] += 1

        created_dt = parse_iso_dt(row.get("created_at"))
        if created_dt:
            bucket_key = created_dt.strftime("%Y-%m-%d")
            label = created_dt.strftime("%d %b")
            if bucket_key not in over_time_counts:
                over_time_counts[bucket_key] = {"label": label, "value": 0}
            over_time_counts[bucket_key]["value"] += 1

        institution_name = row.get("institution_name") or "Unassigned"
        institution_counts[institution_name] = institution_counts.get(institution_name, 0) + 1

        radiologist_name = row.get("radiologist") or "Unassigned"
        radiologist_counts[radiologist_name] = radiologist_counts.get(radiologist_name, 0) + 1
        if not str(row.get("radiologist") or "").strip():
            unassigned_cases += 1

        tat_value = tat_seconds(row.get("created_at"), row.get("vetted_at"))
        avg_tat_values.append(tat_value)
        if status_key == "vetted":
            completed_tat_values.append(tat_value)

        try:
            sla_hours = int(row.get("institution_sla_hours") or row.get("sla_hours") or 48)
        except Exception:
            sla_hours = 48
        if status_key == "pending" and tat_value > sla_hours * 3600:
            sla_breaches += 1

    status_chart = [
        {"label": "Pending", "value": status_counts["pending"], "tone": "pending"},
        {"label": "Vetted", "value": status_counts["vetted"], "tone": "vetted"},
        {"label": "Rejected", "value": status_counts["rejected"], "tone": "rejected"},
        {"label": "Reopened", "value": status_counts["reopened"], "tone": "reopened"},
    ]

    total_cases = len(rows)
    status_max = max([item["value"] for item in status_chart] + [1])

    over_time = [over_time_counts[key] for key in sorted(over_time_counts.keys())]
    if len(over_time) > 12:
        over_time = over_time[-12:]
    over_time_max = max([item["value"] for item in over_time] + [1])

    top_institutions = sorted(institution_counts.items(), key=lambda item: (-item[1], item[0]))[:6]
    institution_chart = [
        {"label": label, "value": value}
        for label, value in top_institutions
    ]
    institution_max = max([item["value"] for item in institution_chart] + [1])

    top_radiologists = sorted(radiologist_counts.items(), key=lambda item: (-item[1], item[0]))[:6]
    radiologist_chart = [
        {"label": label, "value": value}
        for label, value in top_radiologists
    ]
    radiologist_max = max([item["value"] for item in radiologist_chart] + [1])

    avg_tat_seconds = int(sum(avg_tat_values) / len(avg_tat_values)) if avg_tat_values else 0
    completed_avg_tat_seconds = int(sum(completed_tat_values) / len(completed_tat_values)) if completed_tat_values else 0
    completion_rate = round((status_counts["vetted"] / total_cases) * 100, 1) if total_cases else 0.0
    top_institution_label = top_institutions[0][0] if top_institutions else "No data"
    top_practitioner_label = top_radiologists[0][0] if top_radiologists else "No data"

    return {
        "status_chart": status_chart,
        "status_max": status_max,
        "over_time_chart": over_time,
        "over_time_max": over_time_max,
        "institution_chart": institution_chart,
        "institution_max": institution_max,
        "radiologist_chart": radiologist_chart,
        "radiologist_max": radiologist_max,
        "kpis": {
            "total": total_cases,
            "pending": status_counts["pending"],
            "vetted": status_counts["vetted"],
            "rejected": status_counts["rejected"],
            "reopened": status_counts["reopened"],
            "avg_tat": format_tat(avg_tat_seconds),
        },
        "operational_kpis": {
            "completed_avg_tat": format_tat(completed_avg_tat_seconds),
            "unassigned": unassigned_cases,
            "sla_breaches": sla_breaches,
            "completion_rate": f"{completion_rate:.1f}%",
        },
        "insights": {
            "top_institution": top_institution_label,
            "top_practitioner": top_practitioner_label,
            "largest_status": max(status_counts.items(), key=lambda item: item[1])[0].capitalize() if total_cases else "No data",
        },
    }


@app.get("/admin", response_class=HTMLResponse)
def admin_dashboard(
    request: Request,
    view: str = "worklist",
    tab: str = "pending",
    institution: str | None = None,
    radiologist: str | None = None,
    modality: str | None = None,
    q: str | None = None,
    sort_by: str = "created_at",
    sort_dir: str = "desc",
    dashboard_range: str = "30d",
    dashboard_institution: str | None = None,
    dashboard_radiologist: str | None = None,
    dashboard_date_from: str | None = None,
    dashboard_date_to: str | None = None,
):
    try:
        user = require_admin(request)
    except HTTPException:
        return redirect_to_login("admin", "/admin")

    org_id = user.get("org_id")
    is_superuser = bool(user.get("is_superuser"))
    org_name = get_admin_org_name(org_id)
    view = (view or "worklist").strip().lower()
    if view not in ("worklist", "dashboard"):
        view = "worklist"

    institutions = list_institutions(org_id)
    radiologists = [r["name"] for r in list_radiologists(org_id)]
    modalities = list_case_modalities(org_id, is_superuser)

    if table_exists("memberships") and not is_superuser and not org_id:
        return templates.TemplateResponse(
            "home.html",
            {
                "request": request,
                "view": view,
                "view_title": "Dashboard" if view == "dashboard" else "Worklist",
                "tab": tab,
                "cases": [],
                "institutions": institutions,
                "selected_institution": institution or "",
                "radiologists": radiologists,
                "selected_radiologist": radiologist or "",
                "modalities": modalities,
                "selected_modality": modality or "",
                "q": q or "",
                "sort_by": sort_by,
                "sort_dir": sort_dir,
                "pending_count": 0,
                "vetted_count": 0,
                "rejected_count": 0,
                "reopened_count": 0,
                "total_count": 0,
                "dashboard_range": dashboard_range,
                "dashboard_institution": dashboard_institution or "",
                "dashboard_radiologist": dashboard_radiologist or "",
                "dashboard_date_from": dashboard_date_from or "",
                "dashboard_date_to": dashboard_date_to or "",
                "dashboard": build_dashboard_series([]),
                "dashboard_filter_summary": build_dashboard_filter_summary(
                    dashboard_range,
                    dashboard_date_from,
                    dashboard_date_to,
                    dashboard_institution,
                    dashboard_radiologist,
                    institutions,
                ),
                "org_name": org_name,
                "current_user": get_session_user(request),
            },
        )

    tab = (tab or "pending").strip().lower()
    if tab not in ("all", "pending", "vetted", "rejected", "reopened"):
        tab = "pending"

    # Validate sort parameters
    has_modality_column = table_has_column("cases", "modality")
    valid_sorts = ["created_at", "patient_first_name", "patient_surname", "patient_referral_id", "institution_id", "tat", "status", "study_description", "radiologist"]
    if has_modality_column:
        valid_sorts.append("modality")
    if sort_by not in valid_sorts:
        sort_by = "created_at"
    if sort_dir not in ("asc", "desc"):
        sort_dir = "desc"

    worklist_clauses, worklist_params = build_admin_case_filters(
        org_id,
        is_superuser,
        institution=institution,
        radiologist=radiologist,
        modality=modality,
        q=q,
    )

    status_value = None if tab == "all" else tab
    row_clauses, row_params = build_admin_case_filters(
        org_id,
        is_superuser,
        institution=institution,
        radiologist=radiologist,
        modality=modality,
        q=q,
        status=status_value,
    )

    sql = (
        "SELECT c.*, i.name as institution_name "
        "FROM cases c LEFT JOIN institutions i ON c.institution_id = i.id "
        f"WHERE {' AND '.join(row_clauses)}"
    )
    if sort_by == "tat":
        sql += " ORDER BY (JULIANDAY(c.vetted_at) - JULIANDAY(c.created_at)) " + sort_dir
    else:
        sort_col = f"c.{sort_by}" if sort_by != "institution_name" else "i.name"
        sql += f" ORDER BY {sort_col} {sort_dir.upper()}"

    conn = get_db()
    rows = conn.execute(sql, row_params).fetchall()

    counts_sql = (
        "SELECT LOWER(c.status) AS status, COUNT(*) AS c "
        "FROM cases c "
        f"WHERE {' AND '.join(worklist_clauses)} "
        "GROUP BY LOWER(c.status)"
    )
    counts_rows = conn.execute(counts_sql, worklist_params).fetchall()

    dashboard_range = (dashboard_range or "30d").strip().lower()
    if dashboard_range not in ("7d", "30d", "90d", "365d", "all"):
        dashboard_range = "30d"
    dashboard_date_from = (dashboard_date_from or "").strip()
    dashboard_date_to = (dashboard_date_to or "").strip()
    if dashboard_date_from and not dashboard_date_to:
        dashboard_date_to = dashboard_date_from
    if dashboard_date_to and not dashboard_date_from:
        dashboard_date_from = dashboard_date_to
    if dashboard_date_from and dashboard_date_to and dashboard_date_from > dashboard_date_to:
        dashboard_date_from, dashboard_date_to = dashboard_date_to, dashboard_date_from
    created_since = None
    if dashboard_range != "all" and not dashboard_date_from and not dashboard_date_to:
        days = int(dashboard_range.replace("d", ""))
        created_since = (datetime.now(timezone.utc) - timedelta(days=days)).isoformat()

    dashboard_clauses, dashboard_params = build_admin_case_filters(
        org_id,
        is_superuser,
        institution=dashboard_institution,
        radiologist=dashboard_radiologist,
        created_since=created_since,
        created_from=dashboard_date_from or None,
        created_to=dashboard_date_to or None,
    )
    dashboard_sql = (
        "SELECT c.*, i.name as institution_name, i.sla_hours as institution_sla_hours "
        "FROM cases c LEFT JOIN institutions i ON c.institution_id = i.id "
        f"WHERE {' AND '.join(dashboard_clauses)} "
        "ORDER BY c.created_at DESC"
    )
    dashboard_rows = [dict(r) for r in conn.execute(dashboard_sql, dashboard_params).fetchall()]
    conn.close()

    counts = {r["status"]: r["c"] for r in counts_rows}
    pending_count = counts.get("pending", 0)
    vetted_count = counts.get("vetted", 0)
    rejected_count = counts.get("rejected", 0)
    reopened_count = counts.get("reopened", 0)
    total_count = pending_count + vetted_count + rejected_count + reopened_count

    cases: list[dict] = []
    for r in rows:
        d = dict(r)
        d["created_display"] = format_display_datetime(d.get("created_at"), "")
        secs = tat_seconds(d.get("created_at"), d.get("vetted_at"))
        d["tat_display"] = format_tat(secs)
        d["tat_seconds"] = secs
        inst = get_institution(d.get("institution_id")) if d.get("institution_id") else None
        sla_hours = inst["sla_hours"] if inst else 48
        sla_seconds = sla_hours * 3600
        d["sla_breached"] = (d.get("status") == "pending") and (secs > sla_seconds)
        d["display_case_id"] = d.get("id") or "-"
        cases.append(d)

    dashboard = build_dashboard_series(dashboard_rows)
    dashboard_filter_summary = build_dashboard_filter_summary(
        dashboard_range,
        dashboard_date_from,
        dashboard_date_to,
        dashboard_institution,
        dashboard_radiologist,
        institutions,
    )

    return templates.TemplateResponse(
        "home.html",
        {
            "request": request,
            "view": view,
            "view_title": "Dashboard" if view == "dashboard" else "Worklist",
            "tab": tab,
            "cases": cases,
            "institutions": institutions,
            "selected_institution": institution or "",
            "radiologists": radiologists,
            "selected_radiologist": radiologist or "",
            "modalities": modalities,
            "selected_modality": modality or "",
            "q": q or "",
            "sort_by": sort_by,
            "sort_dir": sort_dir,
            "pending_count": pending_count,
            "vetted_count": vetted_count,
            "rejected_count": rejected_count,
            "reopened_count": reopened_count,
            "total_count": total_count,
            "dashboard_range": dashboard_range,
            "dashboard_institution": dashboard_institution or "",
            "dashboard_radiologist": dashboard_radiologist or "",
            "dashboard_date_from": dashboard_date_from or "",
            "dashboard_date_to": dashboard_date_to or "",
            "dashboard": dashboard,
            "dashboard_filter_summary": dashboard_filter_summary,
            "org_name": org_name,
            "current_user": get_session_user(request),
        },
    )


@app.get("/admin.csv")
def admin_dashboard_csv(
    request: Request,
    view: str = "worklist",
    tab: str = "pending",
    institution: str | None = None,
    radiologist: str | None = None,
    modality: str | None = None,
    q: str | None = None,
    dashboard_range: str = "30d",
    dashboard_institution: str | None = None,
    dashboard_radiologist: str | None = None,
    dashboard_date_from: str | None = None,
    dashboard_date_to: str | None = None,
):
    user = require_admin(request)
    org_id = user.get("org_id")
    is_superuser = bool(user.get("is_superuser"))
    view = (view or "worklist").strip().lower()
    if view not in ("worklist", "dashboard"):
        view = "worklist"

    created_since = None
    if view == "dashboard":
        dashboard_range = (dashboard_range or "30d").strip().lower()
        if dashboard_range not in ("7d", "30d", "90d", "365d", "all"):
            dashboard_range = "30d"
        dashboard_date_from = (dashboard_date_from or "").strip()
        dashboard_date_to = (dashboard_date_to or "").strip()
        if dashboard_date_from and not dashboard_date_to:
            dashboard_date_to = dashboard_date_from
        if dashboard_date_to and not dashboard_date_from:
            dashboard_date_from = dashboard_date_to
        if dashboard_date_from and dashboard_date_to and dashboard_date_from > dashboard_date_to:
            dashboard_date_from, dashboard_date_to = dashboard_date_to, dashboard_date_from
        if dashboard_range != "all" and not dashboard_date_from and not dashboard_date_to:
            days = int(dashboard_range.replace("d", ""))
            created_since = (datetime.now(timezone.utc) - timedelta(days=days)).isoformat()
        clauses, params = build_admin_case_filters(
            org_id,
            is_superuser,
            institution=dashboard_institution,
            radiologist=dashboard_radiologist,
            created_since=created_since,
            created_from=dashboard_date_from or None,
            created_to=dashboard_date_to or None,
        )
    else:
        tab = (tab or "pending").strip().lower()
        if tab not in ("all", "pending", "vetted", "rejected", "reopened"):
            tab = "pending"
        clauses, params = build_admin_case_filters(
            org_id,
            is_superuser,
            institution=institution,
            radiologist=radiologist,
            modality=modality,
            q=q,
            status=None if tab == "all" else tab,
        )

    sql = (
        "SELECT c.*, i.name as institution_name "
        "FROM cases c LEFT JOIN institutions i ON c.institution_id = i.id "
        f"WHERE {' AND '.join(clauses)} "
        "ORDER BY c.created_at DESC"
    )

    conn = get_db()
    rows = conn.execute(sql, params).fetchall()
    conn.close()

    case_ids = [r["id"] for r in rows]
    events_map: dict[str, list[dict]] = {}
    if case_ids and table_exists("case_events"):
        placeholders = ",".join(["?"] * len(case_ids))
        conn = get_db()
        ev_rows = conn.execute(
            f"SELECT * FROM case_events WHERE case_id IN ({placeholders}) ORDER BY created_at",
            case_ids,
        ).fetchall()
        conn.close()
        for e in ev_rows:
            d = dict(e)
            events_map.setdefault(d["case_id"], []).append(d)

    org_names: dict[int, str] = {}
    if table_exists("organisations"):
        conn = get_db()
        org_rows = conn.execute("SELECT id, name FROM organisations").fetchall()
        conn.close()
        org_names = {r["id"]: r["name"] for r in org_rows}

    def iter_csv():
        buf = io.StringIO()
        w = csv.writer(buf)

        w.writerow([
            "Case ID",
            "Org ID",
            "Org Name",
            "Submitted At",
            "Reopened",
            "Reopened At",
            "Reopened By",
            "Latest Decision",
            "Latest Decision At",
            "Latest Vetted By",
            "Latest Protocol",
            "Latest Protocol At",
            "Current Status",
            "Radiologist",
            "Patient ID",
            "Patient DOB",
            "Study",
        ])
        yield buf.getvalue()
        buf.seek(0)
        buf.truncate(0)

        for r in rows:
            d = dict(r)
            events = events_map.get(d.get("id"), [])

            submitted_at = ""
            reopened = "N"
            reopened_at = ""
            reopened_by = ""
            latest_decision = ""
            latest_decision_at = ""
            latest_vetted_by = ""
            latest_protocol = ""
            latest_protocol_at = ""

            if events:
                submitted = next((e for e in events if e.get("event_type") == "SUBMITTED"), None)
                if submitted:
                    submitted_at = submitted.get("created_at") or ""

                reopened_events = [e for e in events if e.get("event_type") == "REOPENED"]
                if reopened_events:
                    reopened = "Y"
                    last_reopen = reopened_events[-1]
                    reopened_at = last_reopen.get("created_at") or ""
                    reopened_by = last_reopen.get("username") or ""

                vetted_events = [e for e in events if e.get("event_type") == "VETTED"]
                if vetted_events:
                    last_vet = vetted_events[-1]
                    latest_decision = last_vet.get("decision") or ""
                    latest_decision_at = last_vet.get("created_at") or ""
                    latest_vetted_by = last_vet.get("username") or ""
                    latest_protocol = last_vet.get("protocol") or ""
                    latest_protocol_at = last_vet.get("created_at") or ""

            w.writerow([
                d.get("id", ""),
                d.get("org_id", ""),
                org_names.get(d.get("org_id"), ""),
                format_csv_timestamp(submitted_at),
                reopened,
                format_csv_timestamp(reopened_at),
                reopened_by,
                latest_decision,
                format_csv_timestamp(latest_decision_at),
                latest_vetted_by,
                latest_protocol,
                format_csv_timestamp(latest_protocol_at),
                d.get("status", ""),
                d.get("radiologist", ""),
                d.get("patient_referral_id", ""),
                d.get("patient_dob", "") or "",
                d.get("study_description", ""),
            ])
            yield buf.getvalue()
            buf.seek(0)
            buf.truncate(0)

    filename = f"cases_{tab}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
    return StreamingResponse(iter_csv(), media_type="text/csv", headers={"Content-Disposition": f'attachment; filename="{filename}"'})


@app.get("/admin/dashboard-report.pdf")
def admin_dashboard_report_pdf(
    request: Request,
    dashboard_range: str = "30d",
    dashboard_institution: str | None = None,
    dashboard_radiologist: str | None = None,
    dashboard_date_from: str | None = None,
    dashboard_date_to: str | None = None,
):
    user = require_admin(request)
    org_id = user.get("org_id")
    is_superuser = bool(user.get("is_superuser"))
    institutions = list_institutions(org_id)
    org_name = get_admin_org_name(org_id)

    dashboard_range = (dashboard_range or "30d").strip().lower()
    if dashboard_range not in ("7d", "30d", "90d", "365d", "all"):
        dashboard_range = "30d"
    dashboard_date_from = (dashboard_date_from or "").strip()
    dashboard_date_to = (dashboard_date_to or "").strip()
    if dashboard_date_from and not dashboard_date_to:
        dashboard_date_to = dashboard_date_from
    if dashboard_date_to and not dashboard_date_from:
        dashboard_date_from = dashboard_date_to
    if dashboard_date_from and dashboard_date_to and dashboard_date_from > dashboard_date_to:
        dashboard_date_from, dashboard_date_to = dashboard_date_to, dashboard_date_from

    created_since = None
    if dashboard_range != "all" and not dashboard_date_from and not dashboard_date_to:
        days = int(dashboard_range.replace("d", ""))
        created_since = (datetime.now(timezone.utc) - timedelta(days=days)).isoformat()

    clauses, params = build_admin_case_filters(
        org_id,
        is_superuser,
        institution=dashboard_institution,
        radiologist=dashboard_radiologist,
        created_since=created_since,
        created_from=dashboard_date_from or None,
        created_to=dashboard_date_to or None,
    )

    sql = (
        "SELECT c.*, i.name as institution_name, i.sla_hours as institution_sla_hours "
        "FROM cases c LEFT JOIN institutions i ON c.institution_id = i.id "
        f"WHERE {' AND '.join(clauses)} "
        "ORDER BY c.created_at DESC"
    )
    conn = get_db()
    rows = [dict(r) for r in conn.execute(sql, params).fetchall()]
    conn.close()

    dashboard = build_dashboard_series(rows)
    filter_summary = build_dashboard_filter_summary(
        dashboard_range,
        dashboard_date_from,
        dashboard_date_to,
        dashboard_institution,
        dashboard_radiologist,
        institutions,
    )

    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    accent = colors.HexColor("#1f6feb")
    accent_soft = colors.HexColor("#dbeafe")
    ink = colors.HexColor("#0f172a")
    muted = colors.HexColor("#475569")
    border = colors.HexColor("#cbd5e1")
    card_fill = colors.HexColor("#f8fafc")
    success = colors.HexColor("#dcfce7")
    warning = colors.HexColor("#fef3c7")
    danger = colors.HexColor("#fee2e2")

    left = 36
    right = width - 36
    top = height - 36
    y = top

    def draw_text(x, y_pos, text, font="Helvetica", size=10, color=ink):
        c.setFont(font, size)
        c.setFillColor(color)
        c.drawString(x, y_pos, str(text))

    def draw_card(x, y_top, w, h, label, value, note="", fill=card_fill):
        c.setFillColor(fill)
        c.setStrokeColor(border)
        c.roundRect(x, y_top - h, w, h, 12, stroke=1, fill=1)
        draw_text(x + 10, y_top - 18, label, size=8, color=muted)
        draw_text(x + 10, y_top - 42, value, font="Helvetica-Bold", size=16, color=ink)
        if note:
            draw_text(x + 10, y_top - 56, note, size=7, color=muted)

    def draw_hbar_chart(x, y_top, w, title, items, max_value):
        draw_text(x, y_top, title, font="Helvetica-Bold", size=11)
        chart_y = y_top - 18
        row_h = 18
        for idx, item in enumerate(items[:4]):
            current_y = chart_y - (idx * row_h)
            label = "Completed" if item["label"] == "Vetted" else item["label"]
            value = int(item["value"])
            bar_w = ((w - 110) * value / max_value) if max_value else 0
            draw_text(x, current_y, label, size=8, color=muted)
            c.setFillColor(accent_soft)
            c.roundRect(x + 66, current_y - 8, w - 110, 8, 4, stroke=0, fill=1)
            c.setFillColor(accent)
            c.roundRect(x + 66, current_y - 8, max(bar_w, 2 if value else 0), 8, 4, stroke=0, fill=1)
            draw_text(x + w - 34, current_y, value, size=8, color=ink)

    c.setFillColor(accent)
    c.roundRect(left, y - 54, right - left, 54, 16, stroke=0, fill=1)
    draw_text(left + 16, y - 20, "RadFlow Dashboard Report", font="Helvetica-Bold", size=18, color=colors.white)
    draw_text(left + 16, y - 38, org_name or "Organisation", size=10, color=colors.white)
    draw_text(right - 150, y - 38, datetime.now().strftime("%d %b %Y %H:%M"), size=9, color=colors.white)
    y -= 70

    c.setFillColor(card_fill)
    c.setStrokeColor(border)
    c.roundRect(left, y - 44, right - left, 44, 12, stroke=1, fill=1)
    for idx, line in enumerate(filter_summary):
        draw_text(left + 12 + (idx * 170), y - 26, line, size=8, color=muted)
    y -= 58

    gap = 10
    kpi_w = (right - left - (gap * 4)) / 5
    status_cards = [
        ("Pending", str(dashboard["kpis"]["pending"]), "Awaiting review", warning),
        ("Reopened", str(dashboard["kpis"]["reopened"]), "Returned for action", accent_soft),
        ("Completed", str(dashboard["kpis"]["vetted"]), "Finished cases", success),
        ("Rejected", str(dashboard["kpis"]["rejected"]), "Declined cases", danger),
        ("Total Cases", str(dashboard["kpis"]["total"]), "Filtered scope", card_fill),
    ]
    for idx, (label, value, note, fill) in enumerate(status_cards):
        draw_card(left + idx * (kpi_w + gap), y, kpi_w, 64, label, value, note, fill)
    y -= 78

    op_w = (right - left - (gap * 3)) / 4
    op_cards = [
        ("Completed Avg TAT", dashboard["operational_kpis"]["completed_avg_tat"], "", card_fill),
        ("Unassigned", str(dashboard["operational_kpis"]["unassigned"]), "", card_fill),
        ("SLA Breaches", str(dashboard["operational_kpis"]["sla_breaches"]), "", card_fill),
        ("Completion Rate", dashboard["operational_kpis"]["completion_rate"], "", card_fill),
    ]
    for idx, (label, value, note, fill) in enumerate(op_cards):
        draw_card(left + idx * (op_w + gap), y, op_w, 56, label, value, note, fill)
    y -= 72

    chart_w = (right - left - 16) / 2
    c.setFillColor(card_fill)
    c.setStrokeColor(border)
    c.roundRect(left, y - 118, chart_w, 118, 14, stroke=1, fill=1)
    c.roundRect(left + chart_w + 16, y - 118, chart_w, 118, 14, stroke=1, fill=1)
    draw_hbar_chart(left + 12, y - 12, chart_w - 24, "Cases by Status", dashboard["status_chart"], dashboard["status_max"])
    draw_hbar_chart(left + chart_w + 28, y - 12, chart_w - 24, "Top Institutions", dashboard["institution_chart"], dashboard["institution_max"])
    y -= 134

    c.setFillColor(card_fill)
    c.setStrokeColor(border)
    c.roundRect(left, y - 92, right - left, 92, 14, stroke=1, fill=1)
    draw_text(left + 14, y - 18, "Key Takeaways", font="Helvetica-Bold", size=11)
    takeaway_lines = [
        f"Top institution: {dashboard['insights']['top_institution']}",
        f"Busiest practitioner: {dashboard['insights']['top_practitioner']}",
        f"Largest status group: {'Completed' if dashboard['insights']['largest_status'] == 'Vetted' else dashboard['insights']['largest_status']}",
        f"Cases in selection: {dashboard['kpis']['total']} with {dashboard['operational_kpis']['completion_rate']} completion rate",
    ]
    for idx, line in enumerate(takeaway_lines):
        draw_text(left + 18, y - 38 - (idx * 14), f"- {line}", size=9, color=muted)

    c.showPage()
    c.save()
    buffer.seek(0)

    filename = f"dashboard_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/admin.events.csv")
def admin_events_csv(request: Request):
    user = require_admin(request)
    org_id = user.get("org_id")

    if not table_exists("case_events"):
        raise HTTPException(status_code=404, detail="case_events table not found")

    conn = get_db()
    if org_id and not user.get("is_superuser"):
        rows = conn.execute(
            "SELECT * FROM case_events WHERE org_id = ? ORDER BY created_at",
            (org_id,),
        ).fetchall()
    else:
        rows = conn.execute("SELECT * FROM case_events ORDER BY created_at").fetchall()
    conn.close()

    def iter_csv():
        buf = io.StringIO()
        w = csv.writer(buf)
        w.writerow([
            "Case ID",
            "Org ID",
            "Event",
            "Created At",
            "User ID",
            "Username",
            "Org Role",
            "Decision",
            "Protocol",
            "Comment",
        ])
        yield buf.getvalue()
        buf.seek(0)
        buf.truncate(0)

        for r in rows:
            d = dict(r)
            w.writerow([
                d.get("case_id", ""),
                d.get("org_id", ""),
                d.get("event_type", ""),
                format_csv_timestamp(d.get("created_at", "")),
                d.get("user_id", ""),
                d.get("username", ""),
                d.get("org_role", ""),
                d.get("decision", ""),
                d.get("protocol", ""),
                d.get("comment", ""),
            ])
            yield buf.getvalue()
            buf.seek(0)
            buf.truncate(0)

    filename = f"case_events_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
    return StreamingResponse(iter_csv(), media_type="text/csv", headers={"Content-Disposition": f'attachment; filename="{filename}"'})


@app.get("/admin/notify-radiologist", response_class=HTMLResponse)
def notify_radiologist_page(request: Request, name: str = "", sent: str = "", error: str = ""):
    try:
        user = require_admin(request)
    except HTTPException:
        return redirect_to_login("admin", "/admin/notify-radiologist")

    _uid, _su, org_id, _role = get_current_org_context(request)
    rads = list_radiologists(org_id)

    # Build pending count per radiologist
    conn = get_db()
    pending_counts: dict[str, int] = {}
    rad_emails: dict[str, str] = {}
    for r in rads:
        rname = r["name"]
        if org_id:
            row = conn.execute(
                "SELECT COUNT(*) AS c FROM cases WHERE radiologist = ? AND status IN ('pending','reopened') AND org_id = ?",
                (rname, org_id),
            ).fetchone()
        else:
            row = conn.execute(
                "SELECT COUNT(*) AS c FROM cases WHERE radiologist = ? AND status IN ('pending','reopened')",
                (rname,),
            ).fetchone()
        pending_counts[rname] = row["c"] if row else 0
        rad_emails[rname] = r.get("email") or ""

    notify_history: list[dict[str, str]] = []
    notify_summary: dict[str, dict[str, str | int]] = {}
    summary_since_dt = datetime.now(timezone.utc) - timedelta(hours=24)
    summary_since_iso = summary_since_dt.isoformat()
    since_dt = datetime.now(timezone.utc) - timedelta(days=7)
    since_iso = since_dt.isoformat()
    try:
        if org_id:
            summary_rows = conn.execute(
                """
                SELECT radiologist_name, recipient, created_at, created_by
                FROM notify_events
                WHERE org_id = ? AND created_at >= ?
                ORDER BY created_at DESC
                """,
                (org_id, summary_since_iso),
            ).fetchall()
        else:
            summary_rows = conn.execute(
                """
                SELECT radiologist_name, recipient, created_at, created_by
                FROM notify_events
                WHERE created_at >= ?
                ORDER BY created_at DESC
                """,
                (summary_since_iso,),
            ).fetchall()
        for row in summary_rows or []:
            data = row if isinstance(row, dict) else dict(row)
            rname = data.get("radiologist_name", "") or ""
            if not rname:
                continue
            summary = notify_summary.setdefault(
                rname,
                {
                    "count": 0,
                    "last_created_at": "",
                    "last_created_display": "",
                    "last_created_by": "",
                    "last_recipient": "",
                },
            )
            summary["count"] = int(summary["count"]) + 1
            if not summary["last_created_at"]:
                created_at = data.get("created_at", "") or ""
                summary["last_created_at"] = created_at
                summary["last_created_display"] = format_display_datetime(created_at, created_at)
                summary["last_created_by"] = data.get("created_by", "") or ""
                summary["last_recipient"] = data.get("recipient", "") or ""

        if org_id:
            rows = conn.execute(
                """
                SELECT radiologist_name, channel, recipient, message, created_at, created_by
                FROM notify_events
                WHERE org_id = ? AND created_at >= ?
                ORDER BY created_at DESC
                LIMIT 50
                """,
                (org_id, since_iso),
            ).fetchall()
        else:
            rows = conn.execute(
                """
                SELECT radiologist_name, channel, recipient, message, created_at, created_by
                FROM notify_events
                WHERE created_at >= ?
                ORDER BY created_at DESC
                LIMIT 50
                """,
                (since_iso,),
            ).fetchall()
        for row in rows or []:
            data = row if isinstance(row, dict) else dict(row)
            created_at = data.get("created_at", "")
            dt = parse_iso_dt(created_at)
            created_display = format_display_datetime(created_at, created_at)
            notify_history.append(
                {
                    "radiologist_name": data.get("radiologist_name", ""),
                    "channel": data.get("channel", ""),
                    "recipient": data.get("recipient", ""),
                    "message": data.get("message", ""),
                    "created_at": created_at,
                    "created_display": created_display,
                    "created_by": data.get("created_by", ""),
                }
            )
    except Exception as exc:
        print(f"[NOTIFY] History load failed: {exc}")

    conn.close()

    return templates.TemplateResponse(
        "notify_radiologist.html",
        {
            "request": request,
            "radiologists": rads,
            "pending_counts": pending_counts,
            "rad_emails": rad_emails,
            "app_login_url": APP_BASE_URL,
            "selected_name": name,
            "sent": sent,
            "error": error,
            "smtp_configured": bool(SMTP_HOST),
            "current_user": get_session_user(request),
            "notify_history": notify_history,
            "notify_summary": notify_summary,
        },
    )


@app.post("/admin/notify-radiologist")
def notify_radiologist_send(
    request: Request,
    radiologist_name: str = Form(...),
    channel: str = Form(...),
    recipient: str = Form(""),
    message: str = Form(...),
):
    try:
        require_admin(request)
    except HTTPException:
        return redirect_to_login("admin", "/admin/notify-radiologist")

    if channel == "email":
        if not recipient or "@" not in recipient:
            return RedirectResponse(
                url=f"/admin/notify-radiologist?name={radiologist_name}&error=no_email", status_code=303
            )
        if not SMTP_HOST:
            return RedirectResponse(
                url=f"/admin/notify-radiologist?name={radiologist_name}&error=smtp_not_configured", status_code=303
            )
        import smtplib
        from email.mime.text import MIMEText
        from email.mime.multipart import MIMEMultipart

        msg = MIMEMultipart("alternative")
        msg["Subject"] = "Cases Awaiting Your Review - RadFlow"
        msg["From"] = SMTP_FROM or SMTP_USER
        msg["To"] = recipient.strip()

        html_body = f"""
        <div style="font-family:Arial,sans-serif;max-width:500px;padding:24px;background:#f9f9f9;border-radius:8px;">
          <h2 style="color:#1a1a2e;margin-top:0;">Cases Awaiting Your Review</h2>
          <p style="color:#333;white-space:pre-wrap;">{message}</p>
          <p style="margin:20px 0 0;">
            <a href="{APP_BASE_URL}" style="display:inline-block;padding:10px 16px;background:#1f6feb;color:#ffffff;text-decoration:none;border-radius:8px;">Open RadFlow</a>
          </p>
          <p style="font-size:12px;color:#888;margin-top:24px;">Sent via RadFlow &middot; Healthcare Applications</p>
        </div>
        """
        msg.attach(MIMEText(message, "plain"))
        msg.attach(MIMEText(html_body, "html"))

        try:
            with smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=10) as smtp:
                smtp.ehlo()
                smtp.starttls()
                if SMTP_USER:
                    smtp.login(SMTP_USER, SMTP_PASS)
                smtp.sendmail(msg["From"], [recipient.strip()], msg.as_string())
        except Exception as exc:
            print(f"[NOTIFY] Email send failed: {exc}")
            return RedirectResponse(
                url=f"/admin/notify-radiologist?name={radiologist_name}&error=send_failed", status_code=303
            )

        try:
            _uid, _su, org_id, _role = get_current_org_context(request)
            user = get_session_user(request) or {}
            conn = get_db()
            conn.execute(
                """
                INSERT INTO notify_events (org_id, radiologist_name, channel, recipient, message, created_at, created_by, created_by_id)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    org_id,
                    radiologist_name,
                    channel,
                    recipient.strip(),
                    message,
                    utc_now_iso(),
                    user.get("username"),
                    user.get("id"),
                ),
            )
            conn.commit()
            conn.close()
        except Exception as exc:
            print(f"[NOTIFY] History save failed: {exc}")
        return RedirectResponse(
            url=f"/admin/notify-radiologist?name={radiologist_name}&sent=1", status_code=303
        )

    elif channel == "sms":
        # SMS requires Twilio — not yet configured
        return RedirectResponse(
            url=f"/admin/notify-radiologist?name={radiologist_name}&error=sms_not_configured", status_code=303
        )

    return RedirectResponse(url="/admin/notify-radiologist", status_code=303)


@app.get("/admin/case/{case_id}", response_class=HTMLResponse)
def admin_case_view(request: Request, case_id: str):
    try:
        user = require_admin(request)
    except HTTPException:
        return redirect_to_login("admin", f"/admin/case/{case_id}")

    conn = get_db()
    org_id = user.get("org_id")
    org_name = None
    if org_id and not user.get("is_superuser"):
        row = conn.execute("SELECT * FROM cases WHERE id = ? AND org_id = ?", (case_id, org_id)).fetchone()
    else:
        row = conn.execute("SELECT * FROM cases WHERE id = ?", (case_id,)).fetchone()
    case_dict = None
    if row is not None:
        case_dict = row if isinstance(row, dict) else dict(row)

    if case_dict:
        case_dict = normalize_case_attachment(case_dict)
        case_dict["attachment_previewable"] = is_inline_previewable(case_dict.get("uploaded_filename"))

    if case_dict and case_dict.get("org_id"):
        org_row = conn.execute("SELECT name FROM organisations WHERE id = ?", (case_dict.get("org_id"),)).fetchone()
        if org_row:
            org_name = org_row.get("name") if isinstance(org_row, dict) else org_row[0]

    events = []
    if case_dict and table_exists("case_events"):
        event_rows = conn.execute(
            "SELECT * FROM case_events WHERE case_id = ? ORDER BY created_at ASC",
            (case_id,),
        ).fetchall()
        events = [dict(e) for e in event_rows]

    conn.close()
    if not case_dict:
        raise HTTPException(status_code=404, detail="Case not found")
    # Summary counts for dashboard cards
    conn = get_db()
    if org_id and not user.get("is_superuser"):
        counts_rows = conn.execute(
            "SELECT status, COUNT(*) AS c FROM cases WHERE org_id = ? GROUP BY status",
            (org_id,),
        ).fetchall()
    else:
        counts_rows = conn.execute("SELECT status, COUNT(*) AS c FROM cases GROUP BY status").fetchall()
    conn.close()

    counts = {r["status"]: r["c"] for r in counts_rows}
    pending_count = counts.get("pending", counts.get("PENDING", 0))
    vetted_count = counts.get("vetted", counts.get("VETTED", 0))
    total_count = pending_count + vetted_count

    return templates.TemplateResponse(
        "admin_case.html",
        {"request": request, "case": case_dict, "org_name": org_name, "events": events},
    )


@app.get("/admin/case/{case_id}/timeline.pdf")
def admin_case_timeline_pdf(request: Request, case_id: str):
    try:
        user = require_admin(request)
    except HTTPException:
        return redirect_to_login("admin", f"/admin/case/{case_id}")

    conn = get_db()
    org_id = user.get("org_id")
    if org_id and not user.get("is_superuser"):
        row = conn.execute("SELECT * FROM cases WHERE id = ? AND org_id = ?", (case_id, org_id)).fetchone()
    else:
        row = conn.execute("SELECT * FROM cases WHERE id = ?", (case_id,)).fetchone()

    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Case not found")

    case_dict = row if isinstance(row, dict) else dict(row)
    org_name = ""
    if case_dict.get("org_id"):
        org_row = conn.execute("SELECT name FROM organisations WHERE id = ?", (case_dict.get("org_id"),)).fetchone()
        if org_row:
            org_name = org_row.get("name") if isinstance(org_row, dict) else org_row[0]

    events: list[dict] = []
    if table_exists("case_events"):
        event_rows = conn.execute(
            "SELECT * FROM case_events WHERE case_id = ? ORDER BY created_at ASC",
            (case_id,),
        ).fetchall()
        events = [dict(e) for e in event_rows]
    conn.close()

    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    y = height - 50

    c.setFont("Helvetica-Bold", 14)
    c.drawString(40, y, "Case Timeline Audit Report")
    y -= 22

    c.setFont("Helvetica", 10)
    c.drawString(40, y, f"Case ID: {case_id}")
    y -= 14
    if org_name:
        c.drawString(40, y, f"Organisation: {org_name}")
        y -= 14
    c.drawString(40, y, f"Generated (UTC): {format_display_datetime(utc_now_iso())}")
    y -= 20

    c.setFont("Helvetica-Bold", 10)
    c.drawString(40, y, "Timestamp (UTC)")
    c.drawString(155, y, "Event")
    c.drawString(235, y, "User")
    c.drawString(315, y, "Details")
    y -= 10
    c.line(40, y, width - 40, y)
    y -= 14

    def wrap_text(text_value: str, max_width: int) -> list[str]:
        words = (text_value or "").split()
        if not words:
            return [""]
        lines: list[str] = []
        current = words[0]
        for word in words[1:]:
            candidate = f"{current} {word}"
            if c.stringWidth(candidate, "Helvetica", 9) <= max_width:
                current = candidate
            else:
                lines.append(current)
                current = word
        lines.append(current)
        return lines

    if not events:
        c.setFont("Helvetica", 10)
        c.drawString(40, y, "No timeline events recorded for this case.")
    else:
        for event in events:
            ts = format_display_datetime(event.get("created_at"), event.get("created_at") or "")
            event_type = str(event.get("event_type") or "-")
            username = str(event.get("username") or "-")
            details = str(event.get("comment") or "")
            detail_lines = wrap_text(details, 250)

            required_height = max(14, 12 * len(detail_lines)) + 6
            if y - required_height < 40:
                c.showPage()
                y = height - 50
                c.setFont("Helvetica-Bold", 10)
                c.drawString(40, y, "Timestamp (UTC)")
                c.drawString(155, y, "Event")
                c.drawString(235, y, "User")
                c.drawString(315, y, "Details")
                y -= 10
                c.line(40, y, width - 40, y)
                y -= 14

            c.setFont("Helvetica", 9)
            c.drawString(40, y, ts)
            c.drawString(155, y, event_type)
            c.drawString(235, y, username)
            line_y = y
            for line in detail_lines:
                c.drawString(315, line_y, line)
                line_y -= 12
            y = line_y - 6

    c.save()
    buffer.seek(0)
    filename = f"{case_id}_timeline_audit.pdf"
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )

@app.get("/admin/case/{case_id}/timeline.csv")
def admin_case_timeline_csv(request: Request, case_id: str):
    try:
        user = require_admin(request)
    except HTTPException:
        return redirect_to_login("admin", f"/admin/case/{case_id}")

    conn = get_db()
    org_id = user.get("org_id")
    if org_id and not user.get("is_superuser"):
        row = conn.execute("SELECT * FROM cases WHERE id = ? AND org_id = ?", (case_id, org_id)).fetchone()
    else:
        row = conn.execute("SELECT * FROM cases WHERE id = ?", (case_id,)).fetchone()

    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Case not found")

    case_dict = row if isinstance(row, dict) else dict(row)
    org_name = ""
    if case_dict.get("org_id"):
        org_row = conn.execute("SELECT name FROM organisations WHERE id = ?", (case_dict.get("org_id"),)).fetchone()
        if org_row:
            org_name = org_row.get("name") if isinstance(org_row, dict) else org_row[0]

    events: list[dict] = []
    if table_exists("case_events"):
        event_rows = conn.execute(
            "SELECT * FROM case_events WHERE case_id = ? ORDER BY created_at ASC",
            (case_id,),
        ).fetchall()
        events = [dict(e) for e in event_rows]
    conn.close()

    output = io.StringIO()
    writer = csv.writer(output)
    
    # Write header with case info
    writer.writerow(["Case Timeline Audit Report"])
    writer.writerow([f"Case ID: {case_id}"])
    if org_name:
        writer.writerow([f"Organisation: {org_name}"])
    writer.writerow([f"Generated (UTC): {format_display_datetime(utc_now_iso())}"])
    writer.writerow([])  # Empty row
    
    # Write column headers
    writer.writerow(["Timestamp (UTC)", "Event Type", "User", "Details"])
    
    # Write events
    if not events:
        writer.writerow(["No timeline events recorded for this case."])
    else:
        for event in events:
            ts = format_display_datetime(event.get("created_at"), event.get("created_at") or "")
            event_type = str(event.get("event_type") or "-")
            username = str(event.get("username") or "-")
            details = str(event.get("comment") or "")
            writer.writerow([ts, event_type, username, details])
    
    output.seek(0)
    filename = f"{case_id}_timeline_audit.csv"
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )

# -------------------------
# Case edit
# -------------------------
@app.get("/admin/case/{case_id}/edit", response_class=HTMLResponse)
def admin_case_edit_view(request: Request, case_id: str):
    saved = request.query_params.get("saved", "")
    try:
        user = require_admin(request)
    except HTTPException as e:
        return redirect_to_login("admin", f"/admin/case/{case_id}/edit")

    conn = get_db()
    org_id = user.get("org_id")
    if org_id and not user.get("is_superuser"):
        case = conn.execute("SELECT * FROM cases WHERE id = ? AND org_id = ?", (case_id, org_id)).fetchone()
    else:
        case = conn.execute("SELECT * FROM cases WHERE id = ?", (case_id,)).fetchone()
    institutions = list_institutions(org_id)
    radiologists = list_radiologists(org_id)
    protocols = conn.execute("SELECT DISTINCT protocol FROM cases WHERE protocol IS NOT NULL AND protocol != '' ORDER BY protocol").fetchall()
    conn.close()
    
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    case_dict = dict(case)
    case_dict = normalize_case_attachment(case_dict)
    case_dict["attachment_previewable"] = is_inline_previewable(case_dict.get("uploaded_filename"))
    
    return templates.TemplateResponse(
        "case_edit.html",
        {
            "request": request,
            "case": case_dict,
            "institutions": institutions,
            "radiologists": radiologists,
            "protocols": [p["protocol"] for p in protocols] if protocols else [],
            "user_org_id": org_id,
            "saved": saved == "1",
        }
    )

@app.post("/admin/case/{case_id}/edit")
async def admin_case_edit_save(
    request: Request,
    case_id: str,
    patient_first_name: str = Form(""),
    patient_surname: str = Form(""),
    patient_referral_id: str = Form(""),
    patient_dob: str = Form(""),
    institution_id: str = Form(""),
    study_description: str = Form(""),
    admin_notes: str = Form(""),
    radiologist: str = Form(""),
    modality: str = Form(""),
    protocol: str = Form(""),
    attachment: UploadFile | None = File(None),
):
    try:
        user = require_admin(request)
    except HTTPException:
        raise HTTPException(status_code=403, detail="Not authorized")

    conn = get_db()
    org_id = user.get("org_id")
    if org_id and not user.get("is_superuser"):
        case = conn.execute("SELECT * FROM cases WHERE id = ? AND org_id = ?", (case_id, org_id)).fetchone()
    else:
        case = conn.execute("SELECT * FROM cases WHERE id = ?", (case_id,)).fetchone()
    if not case:
        conn.close()
        raise HTTPException(status_code=404, detail="Case not found")

    case_dict = dict(case)

    cleaned_first_name = patient_first_name.strip()
    cleaned_surname = patient_surname.strip()
    cleaned_referral_id = patient_referral_id.strip()
    cleaned_dob = patient_dob.strip() or None
    cleaned_study_description = study_description.strip()
    cleaned_admin_notes = admin_notes.strip()
    cleaned_protocol = protocol.strip() or None
    cleaned_modality = modality.strip() or None

    institution_raw = institution_id.strip()
    cleaned_institution_id = int(institution_raw) if institution_raw.isdigit() else None

    # Keep existing radiologist if the form submits an empty value.
    # Empty string is allowed by current app flow; avoid writing NULL.
    cleaned_radiologist = radiologist.strip() or (case_dict.get("radiologist") or "")

    replacement_uploaded_filename: str | None = None
    replacement_stored_path: str | None = None
    old_stored_path = case_dict.get("stored_filepath")
    if attachment and attachment.filename:
        replacement_uploaded_filename = attachment.filename
        replacement_bytes = await attachment.read()
        if replacement_bytes:
            if BLOB_STORAGE_ENABLED:
                blob_name = upload_to_blob(case_id, replacement_bytes, replacement_uploaded_filename)
                if blob_name:
                    replacement_stored_path = blob_name

            if not replacement_stored_path:
                safe_name = f"{case_id}_{Path(replacement_uploaded_filename).name}"
                replacement_stored_path = str(UPLOAD_DIR / safe_name)
                with open(replacement_stored_path, "wb") as f:
                    f.write(replacement_bytes)

            if old_stored_path and old_stored_path != replacement_stored_path and str(old_stored_path).startswith(str(UPLOAD_DIR)):
                try:
                    old_path_obj = Path(str(old_stored_path))
                    if old_path_obj.exists():
                        old_path_obj.unlink()
                except Exception:
                    pass

    def _clean(value: str | None) -> str:
        return (value or "").strip()

    update_fields: list[tuple[str, str | int | None]] = []

    def add_field_if_exists(column_name: str, value: str | int | None) -> None:
        if table_has_column("cases", column_name):
            update_fields.append((column_name, value))

    add_field_if_exists("patient_first_name", cleaned_first_name)
    add_field_if_exists("patient_surname", cleaned_surname)
    add_field_if_exists("patient_referral_id", cleaned_referral_id)
    add_field_if_exists("patient_dob", cleaned_dob)
    add_field_if_exists("institution_id", cleaned_institution_id)
    add_field_if_exists("study_description", cleaned_study_description)
    add_field_if_exists("admin_notes", cleaned_admin_notes)
    add_field_if_exists("radiologist", cleaned_radiologist)
    add_field_if_exists("protocol", cleaned_protocol)
    add_field_if_exists("modality", cleaned_modality)
    if replacement_uploaded_filename is not None:
        add_field_if_exists("uploaded_filename", replacement_uploaded_filename)
        add_field_if_exists("stored_filepath", replacement_stored_path)

    if not update_fields:
        conn.close()
        raise HTTPException(status_code=400, detail="No editable fields available for this case")

    changes: list[str] = []
    old_case = dict(case)
    for col, new_val in update_fields:
        if col == "admin_notes":
            continue
        old_val = old_case.get(col)
        old_text = _clean(str(old_val)) if old_val is not None else ""
        new_text = _clean(str(new_val)) if new_val is not None else ""
        if old_text != new_text:
            changes.append(f"{col}: {old_text or '-'} -> {new_text or '-'}")

    update_sql = "UPDATE cases SET " + ", ".join([f"{col} = ?" for col, _ in update_fields]) + " WHERE id = ?"
    update_params = [value for _, value in update_fields] + [case_id]
    try:
        conn.execute(update_sql, update_params)
        conn.commit()
    except Exception as exc:
        conn.close()
        print(f"[ERROR] Failed to save case edits for case {case_id}: {exc}")
        raise HTTPException(status_code=400, detail="Unable to save case changes")
    conn.close()

    if replacement_uploaded_filename is not None:
        changes.append(f"attachment: replaced with {replacement_uploaded_filename}")

    change_summary = "; ".join(changes) if changes else "No field changes"
    note_text = cleaned_admin_notes
    event_comment = change_summary
    if note_text:
        event_comment = f"{change_summary}. Notes: {note_text}"

    event_org_id = org_id or case_dict.get("org_id")
    try:
        insert_case_event(
            case_id=case_id,
            org_id=event_org_id,
            event_type="EDITED",
            user=user,
            comment=event_comment,
        )
    except Exception as exc:
        print(f"[WARN] Saved case {case_id} but failed to write edit event: {exc}")

    return RedirectResponse(url=f"/admin/case/{case_id}/edit?saved=1", status_code=303)


@app.post("/admin/case/{case_id}/assign-radiologist")
def assign_radiologist(request: Request, case_id: str, radiologist: str = Form("")):
    user = require_admin(request)
    org_id = user.get("org_id")

    radiologist = (radiologist or "").strip()

    conn = get_db()
    if org_id and not user.get("is_superuser"):
        case = conn.execute("SELECT id, radiologist, org_id FROM cases WHERE id = ? AND org_id = ?", (case_id, org_id)).fetchone()
    else:
        case = conn.execute("SELECT id, radiologist, org_id FROM cases WHERE id = ?", (case_id,)).fetchone()

    if not case:
        conn.close()
        raise HTTPException(status_code=404, detail="Case not found")

    valid_rads = {r["name"] for r in list_radiologists(org_id)}
    if radiologist and radiologist not in valid_rads:
        conn.close()
        raise HTTPException(status_code=400, detail="Invalid radiologist selection")

    old_rad = case["radiologist"] if isinstance(case, dict) else case[1]
    conn.execute(
        "UPDATE cases SET radiologist = ? WHERE id = ?",
        (radiologist or None, case_id),
    )
    conn.commit()
    conn.close()

    comment = f"{old_rad or 'unassigned'} -> {radiologist or 'unassigned'}"
    insert_case_event(
        case_id=case_id,
        org_id=case["org_id"] if isinstance(case, dict) else case[2],
        event_type="ASSIGNED",
        user=user,
        comment=comment,
    )

    return RedirectResponse(url="/admin", status_code=303)


@app.get("/admin/case/{case_id}/reopen", response_class=HTMLResponse)
def admin_reopen_case_form(request: Request, case_id: str):
    try:
        user = require_admin(request)
    except HTTPException:
        return redirect_to_login("admin", f"/admin/case/{case_id}/reopen")

    org_id = user.get("org_id")
    conn = get_db()
    row = conn.execute("SELECT * FROM cases WHERE id = ?", (case_id,)).fetchone()
    conn.close()

    if not row:
        raise HTTPException(status_code=404, detail="Case not found")

    # Org isolation
    if org_id and row["org_id"] != org_id:
        raise HTTPException(status_code=403, detail="Access denied")

    case = dict(row)
    
    # Only allow reopening vetted or rejected cases
    if case["status"] not in ["vetted", "rejected"]:
        return RedirectResponse(url=f"/admin/case/{case_id}", status_code=303)

    return templates.TemplateResponse(
        "admin_reopen_case.html",
        {"request": request, "case": case, "user": user}
    )


@app.post("/admin/case/{case_id}/reopen")
def admin_reopen_case_submit(
    request: Request,
    case_id: str,
    reopen_notes: str = Form(...),
):
    try:
        user = require_admin(request)
    except HTTPException:
        return RedirectResponse(url="/login", status_code=303)

    org_id = user.get("org_id")
    conn = get_db()
    row = conn.execute("SELECT * FROM cases WHERE id = ?", (case_id,)).fetchone()
    
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Case not found")

    # Org isolation
    if org_id and row["org_id"] != org_id:
        conn.close()
        raise HTTPException(status_code=403, detail="Access denied")

    # Keep reopened notes distinguishable so the UI can highlight them without
    # mixing them into ordinary admin notes.
    current_notes = row["admin_notes"] or ""
    reopen_entry = f"{REOPEN_NOTE_MARKER}\n{reopen_notes.strip()}"
    updated_notes = "\n\n".join(part for part in [current_notes.strip(), reopen_entry] if part).strip()
    
    # Keep the prior decision context visible when a case is reopened.
    # This helps the radiologist see what was decided previously while still
    # making it clear that the case needs a fresh review.
    conn.execute(
        "UPDATE cases SET status = ?, admin_notes = ? WHERE id = ?",
        ("reopened", updated_notes, case_id)
    )
    conn.commit()
    conn.close()
    insert_case_event(
        case_id=case_id,
        org_id=row["org_id"],
        event_type="REOPENED",
        user=user,
        comment=reopen_notes.strip(),
    )
    return RedirectResponse(url=f"/admin/case/{case_id}/edit?saved=1", status_code=303)


def slugify_org_name(name: str) -> str:
    base = re.sub(r"[^a-z0-9]+", "-", str(name or "").strip().lower()).strip("-")
    return base or "organisation"


def ensure_extended_identity_schema() -> None:
    if using_postgres():
        return

    conn = get_db()
    cur = conn.cursor()

    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS organisations (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            slug TEXT NOT NULL UNIQUE,
            is_active INTEGER NOT NULL DEFAULT 1,
            created_at TEXT NOT NULL,
            modified_at TEXT
        )
        """
    )
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS memberships (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            org_id INTEGER NOT NULL,
            user_id INTEGER NOT NULL,
            org_role TEXT NOT NULL DEFAULT 'org_user',
            is_active INTEGER NOT NULL DEFAULT 1,
            created_at TEXT NOT NULL,
            modified_at TEXT,
            UNIQUE(org_id, user_id)
        )
        """
    )
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS user_sessions (
            user_id INTEGER PRIMARY KEY,
            session_id TEXT NOT NULL,
            created_at TEXT NOT NULL
        )
        """
    )
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS radiologist_profiles (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL UNIQUE,
            gmc TEXT,
            specialty TEXT,
            display_name TEXT,
            created_at TEXT NOT NULL,
            modified_at TEXT
        )
        """
    )

    cur.execute("PRAGMA table_info(users)")
    user_cols = {row[1] for row in cur.fetchall()}
    needs_user_upgrade = {"id", "password_hash", "is_superuser", "is_active", "created_at", "modified_at"} - user_cols

    if needs_user_upgrade:
        now = utc_now_iso()
        cur.execute("DROP TABLE IF EXISTS users_extended_new")
        cur.execute(
            """
            CREATE TABLE users_extended_new (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT NOT NULL UNIQUE,
                email TEXT UNIQUE,
                password_hash TEXT NOT NULL,
                salt_hex TEXT NOT NULL,
                is_superuser INTEGER NOT NULL DEFAULT 0,
                is_active INTEGER NOT NULL DEFAULT 1,
                created_at TEXT NOT NULL,
                modified_at TEXT,
                first_name TEXT,
                surname TEXT,
                role TEXT,
                radiologist_name TEXT
            )
            """
        )
        old_rows = cur.execute(
            "SELECT username, first_name, surname, email, role, radiologist_name, salt_hex, pw_hash_hex FROM users ORDER BY username"
        ).fetchall()
        promote_username = None
        for old in old_rows:
            if str(old[4] or "").strip().lower() == "admin" and promote_username is None:
                promote_username = old[0]
        for old in old_rows:
            username = old[0]
            email = (old[3] or "").strip() or None
            role = (old[4] or "user").strip()
            is_superuser = 1 if username == promote_username else 0
            cur.execute(
                """
                INSERT INTO users_extended_new(
                    username, email, password_hash, salt_hex, is_superuser, is_active, created_at, modified_at,
                    first_name, surname, role, radiologist_name
                )
                VALUES (?, ?, ?, ?, ?, 1, ?, ?, ?, ?, ?, ?)
                """,
                (
                    username,
                    email,
                    old[7],
                    old[6],
                    is_superuser,
                    now,
                    now,
                    old[1],
                    old[2],
                    role,
                    old[5],
                ),
            )
        cur.execute("ALTER TABLE users RENAME TO users_legacy_backup")
        cur.execute("ALTER TABLE users_extended_new RENAME TO users")

    cur.execute("PRAGMA table_info(institutions)")
    institution_cols = {row[1] for row in cur.fetchall()}
    if "org_id" not in institution_cols:
        cur.execute("ALTER TABLE institutions ADD COLUMN org_id INTEGER")

    cur.execute("PRAGMA table_info(cases)")
    case_cols = {row[1] for row in cur.fetchall()}
    if "org_id" not in case_cols:
        cur.execute("ALTER TABLE cases ADD COLUMN org_id INTEGER")

    cur.execute("PRAGMA table_info(protocols)")
    protocol_cols = {row[1] for row in cur.fetchall()}
    if "org_id" not in protocol_cols:
        cur.execute("ALTER TABLE protocols ADD COLUMN org_id INTEGER")

    conn.commit()

    org_row = cur.execute("SELECT id FROM organisations ORDER BY id LIMIT 1").fetchone()
    if org_row:
        default_org_id = org_row[0]
    else:
        default_org_name = "Default Organisation"
        default_slug = slugify_org_name(default_org_name)
        suffix = 1
        while cur.execute("SELECT 1 FROM organisations WHERE slug = ?", (default_slug,)).fetchone():
            suffix += 1
            default_slug = f"{slugify_org_name(default_org_name)}-{suffix}"
        cur.execute(
            "INSERT INTO organisations(name, slug, is_active, created_at, modified_at) VALUES (?, ?, 1, ?, ?)",
            (default_org_name, default_slug, utc_now_iso(), utc_now_iso()),
        )
        default_org_id = cur.lastrowid

    cur.execute("UPDATE institutions SET org_id = ? WHERE org_id IS NULL", (default_org_id,))
    cur.execute("UPDATE cases SET org_id = ? WHERE org_id IS NULL", (default_org_id,))
    cur.execute("UPDATE protocols SET org_id = ? WHERE org_id IS NULL", (default_org_id,))

    if table_exists("study_description_presets"):
        cur.execute("UPDATE study_description_presets SET organization_id = ? WHERE organization_id IS NULL OR organization_id = 0", (default_org_id,))

    user_rows = cur.execute("SELECT id, username, role, first_name, surname, radiologist_name FROM users").fetchall()
    for user_row in user_rows:
        user_id = user_row[0]
        username = user_row[1]
        role = str(user_row[2] or "user").strip().lower()
        org_role = "org_admin" if role == "admin" else "radiologist" if role == "radiologist" else "org_user"
        if not cur.execute("SELECT 1 FROM memberships WHERE org_id = ? AND user_id = ?", (default_org_id, user_id)).fetchone():
            cur.execute(
                "INSERT INTO memberships(org_id, user_id, org_role, is_active, created_at, modified_at) VALUES (?, ?, ?, 1, ?, ?)",
                (default_org_id, user_id, org_role, utc_now_iso(), utc_now_iso()),
            )
        if role == "radiologist":
            display_name = (
                str(user_row[5] or "").strip()
                or " ".join(part for part in [str(user_row[3] or "").strip(), str(user_row[4] or "").strip()] if part)
                or username
            )
            if not cur.execute("SELECT 1 FROM radiologist_profiles WHERE user_id = ?", (user_id,)).fetchone():
                cur.execute(
                    "INSERT INTO radiologist_profiles(user_id, gmc, specialty, display_name, created_at, modified_at) VALUES (?, ?, ?, ?, ?, ?)",
                    (user_id, None, None, display_name, utc_now_iso(), utc_now_iso()),
                )

    conn.commit()
    conn.close()


def list_organisations_summary() -> list[dict]:
    if not table_exists("organisations"):
        return []
    conn = get_db()
    rows = conn.execute(
        """
        SELECT
            o.id,
            o.name,
            o.slug,
            o.is_active,
            o.created_at,
            (
                SELECT COUNT(*)
                FROM memberships m
                WHERE m.org_id = o.id AND m.org_role = 'org_admin' AND m.is_active = 1
            ) AS admin_count,
            (
                SELECT COUNT(*)
                FROM memberships m
                WHERE m.org_id = o.id AND m.org_role = 'radiologist' AND m.is_active = 1
            ) AS radiologist_count,
            (
                SELECT COUNT(*)
                FROM memberships m
                WHERE m.org_id = o.id AND m.is_active = 1
            ) AS user_count,
            (
                SELECT COUNT(*)
                FROM institutions i
                WHERE i.org_id = o.id
            ) AS institution_count
        FROM organisations o
        ORDER BY o.name
        """
    ).fetchall()
    conn.close()
    return [dict(row) for row in rows]


def get_organisation_summary(org_id: int) -> dict | None:
    if not table_exists("organisations"):
        return None
    conn = get_db()
    row = conn.execute(
        """
        SELECT
            o.id,
            o.name,
            o.slug,
            o.is_active,
            o.created_at,
            (
                SELECT COUNT(*)
                FROM memberships m
                WHERE m.org_id = o.id AND m.org_role = 'org_admin' AND m.is_active = 1
            ) AS admin_count,
            (
                SELECT COUNT(*)
                FROM memberships m
                WHERE m.org_id = o.id AND m.org_role = 'radiologist' AND m.is_active = 1
            ) AS radiologist_count,
            (
                SELECT COUNT(*)
                FROM memberships m
                WHERE m.org_id = o.id AND m.is_active = 1
            ) AS user_count,
            (
                SELECT COUNT(*)
                FROM institutions i
                WHERE i.org_id = o.id
            ) AS institution_count
        FROM organisations o
        WHERE o.id = ?
        LIMIT 1
        """,
        (org_id,),
    ).fetchone()
    conn.close()
    return dict(row) if row else None


def list_organisation_users(org_id: int) -> list[dict]:
    if not table_exists("users"):
        return []

    conn = get_db()
    profile_join = "LEFT JOIN radiologist_profiles rp ON rp.user_id = u.id" if table_exists("radiologist_profiles") else ""
    profile_gmc = "rp.gmc" if table_exists("radiologist_profiles") else "NULL"
    profile_specialty = "rp.specialty" if table_exists("radiologist_profiles") else "NULL"
    profile_display = "rp.display_name" if table_exists("radiologist_profiles") else "NULL"
    rows = conn.execute(
        f"""
        SELECT
            u.id,
            u.username,
            u.first_name,
            u.surname,
            u.email,
            u.is_active,
            COALESCE(u.mfa_enabled, 0) AS mfa_enabled,
            COALESCE(u.mfa_required, 0) AS mfa_required,
            m.org_role,
            {profile_gmc} AS gmc,
            {profile_specialty} AS specialty,
            COALESCE({profile_display}, NULLIF(TRIM(COALESCE(u.first_name, '') || ' ' || COALESCE(u.surname, '')), ''), u.username) AS display_name
        FROM memberships m
        INNER JOIN users u ON u.id = m.user_id
        {profile_join}
        WHERE m.org_id = ? AND m.is_active = 1
        ORDER BY
            CASE m.org_role
                WHEN 'org_admin' THEN 0
                WHEN 'radiologist' THEN 1
                ELSE 2
            END,
            LOWER(COALESCE(u.surname, '')),
            LOWER(COALESCE(u.first_name, '')),
            LOWER(u.username)
        """,
        (org_id,),
    ).fetchall()
    conn.close()

    users: list[dict] = []
    for row in rows:
        item = dict(row)
        role_value = item.get("org_role") or "org_user"
        item["role_label"] = (
            "Admin" if role_value == "org_admin"
            else "Practitioner" if role_value == "radiologist"
            else "Coordinator"
        )
        users.append(item)
    return users


def list_organisation_institutions(org_id: int) -> list[dict]:
    if not table_exists("institutions"):
        return []
    conn = get_db()
    rows = conn.execute(
        """
        SELECT
            i.id,
            i.name,
            i.sla_hours,
            i.created_at,
            i.modified_at,
            (
                SELECT COUNT(*)
                FROM cases c
                WHERE c.org_id = i.org_id AND c.institution_id = i.id
            ) AS case_count
        FROM institutions i
        WHERE i.org_id = ?
        ORDER BY LOWER(i.name)
        """,
        (org_id,),
    ).fetchall()
    conn.close()
    return [dict(row) for row in rows]

    insert_case_event(
        case_id=case_id,
        org_id=org_id or row["org_id"],
        event_type="REOPENED",
        user=user,
        comment=reopen_notes.strip() or None,
    )

    redirect_target = f"/admin/case/{case_id}"
    referer = request.headers.get("referer", "")
    if referer.endswith(f"/admin/case/{case_id}/edit"):
        redirect_target = f"/admin/case/{case_id}/edit?saved=1"
    return RedirectResponse(url=redirect_target, status_code=303)


# -------------------------
# Radiologist dashboard
# -------------------------
@app.get("/radiologist", response_class=HTMLResponse)
def radiologist_dashboard(request: Request, tab: str = "all"):
    try:
        user = require_radiologist(request)
    except HTTPException:
        return redirect_to_login("radiologist", "/radiologist")

    org_id = user.get("org_id")
    rad_name = user.get("radiologist_name")
    if not rad_name:
        raise HTTPException(status_code=400, detail="Practitioner account not linked to a practitioner name")

    tab = (tab or "all").strip().lower()
    if tab not in ("all", "pending", "vetted", "rejected", "reopened"):
        tab = "all"

    sql = "SELECT c.*, i.sla_hours FROM cases c LEFT JOIN institutions i ON c.institution_id = i.id WHERE c.radiologist = ?"
    params: list[str] = [rad_name]

    if org_id:
        sql += " AND c.org_id = ?"
        params.append(org_id)

    if tab == "pending":
        sql += " AND c.status IN (?, ?)"
        params.extend(["pending", "reopened"])
    elif tab == "vetted":
        sql += " AND c.status = ?"
        params.append("vetted")
    elif tab == "rejected":
        sql += " AND c.status = ?"
        params.append("rejected")
    elif tab == "reopened":
        sql += " AND c.status = ?"
        params.append("reopened")

    sql += " ORDER BY c.created_at DESC"

    conn = get_db()
    rows = conn.execute(sql, params).fetchall()
    conn.close()

    cases: list[dict] = []
    for r in rows:
        d = dict(r)
        created_dt = parse_iso_dt(d.get("created_at"))
        d["created_display"] = format_display_datetime(d.get("created_at"), d.get("created_at") or "")

        secs = tat_seconds(d.get("created_at"), d.get("vetted_at"))
        d["tat_display"] = format_tat(secs)
        
        # Use institution-specific SLA or default to 48 hours
        sla_hours = d.get("sla_hours") or 48
        sla_seconds = sla_hours * 3600
        d["sla_breached"] = (d.get("status") == "pending") and (secs > sla_seconds)

        cases.append(d)

    return templates.TemplateResponse(
        "radiologist_dashboard.html",
        {
            "request": request,
            "cases": cases,
            "tab": tab,
            "current_user": get_session_user(request),
        },
    )


# -------------------------
# Settings (admin)
# -------------------------
@app.get("/settings", response_class=HTMLResponse)
def settings_page(request: Request, error: str = ""):
    try:
        require_admin(request)
    except HTTPException:
        return redirect_to_login("admin", "/settings")

    org_id = get_request_org_id(request)
    org_name = ""
    
    # Ensure default institution exists for this org
    if org_id:
        institutions = list_institutions(org_id)
        conn = get_db()
        org_row = conn.execute("SELECT name FROM organisations WHERE id = ?", (org_id,)).fetchone()
        conn.close()
        if org_row:
            org_name = org_row["name"] or ""
        if not institutions:
            # Create default institution with org name
            conn = get_db()
            org_row = conn.execute("SELECT name FROM organisations WHERE id = ?", (org_id,)).fetchone()
            if org_row:
                org_name = org_row["name"]
                now = utc_now_iso()
                if table_has_column("institutions", "org_id"):
                    conn.execute(
                        """
                        INSERT INTO institutions (name, sla_hours, org_id, created_at, modified_at)
                        VALUES (?, 48, ?, ?, ?)
                        """,
                        (org_name, org_id, now, now)
                    )
                    conn.commit()
            conn.close()
            # Refresh institutions list
            institutions = list_institutions(org_id)
    else:
        institutions = list_institutions(org_id)
    
    rads = list_radiologists(org_id)
    users = list_users(org_id)
    rad_names = [r["name"] for r in rads]
    protocols = list_protocol_rows(org_id)
    study_description_presets = list_study_description_presets(org_id)
    protocol_modalities = sorted({str(p.get("modality") or "").strip() for p in study_description_presets if str(p.get("modality") or "").strip()})
    report_key_scope = org_id or 0
    report_header_text = get_setting(f"report_header:{report_key_scope}", org_name or "")
    report_footer_text = get_setting(f"report_footer:{report_key_scope}", "Confidential workflow document")

    return templates.TemplateResponse(
        "settings.html",
        {
            "request": request,
            "institutions": institutions,
            "radiologists": rads,
            "users": users,
            "rad_names": rad_names,
            "protocols": protocols,
            "study_description_presets": study_description_presets,
            "protocol_modalities": protocol_modalities,
            "org_name": org_name,
            "report_header_text": report_header_text,
            "report_footer_text": report_footer_text,
            "current_user": get_session_user(request),
            "error": error,
        },
    )


@app.post("/settings/report")
def update_report_settings(
    request: Request,
    report_header_text: str = Form(""),
    report_footer_text: str = Form(""),
):
    user = require_admin(request)
    org_id = user.get("org_id") or 0
    set_setting(f"report_header:{org_id}", report_header_text.strip())
    set_setting(f"report_footer:{org_id}", report_footer_text.strip())
    return RedirectResponse(url="/settings?tab=report", status_code=303)


@app.post("/settings/institution/add")
def add_institution(request: Request, name: str = Form(...), sla_hours: str = Form(...)):
    user = require_admin(request)
    org_id = user.get("org_id")
    try:
        sla_val = int(sla_hours)
        if sla_val <= 0 or sla_val > 999:
            raise ValueError()
    except ValueError:
        raise HTTPException(status_code=400, detail="SLA must be a number of hours (1-999)")
    
    upsert_institution(name.strip(), sla_val, org_id)
    return RedirectResponse(url="/settings", status_code=303)


@app.post("/settings/institution/edit/{inst_id}")
def edit_institution(request: Request, inst_id: int, name: str = Form(...), sla_hours: str = Form(...)):
    user = require_admin(request)
    org_id = user.get("org_id")
    try:
        sla_val = int(sla_hours)
        if sla_val <= 0 or sla_val > 999:
            raise ValueError()
    except ValueError:
        raise HTTPException(status_code=400, detail="SLA must be a number of hours (1-999)")
    
    # Verify institution exists
    inst = get_institution(inst_id, org_id)
    if not inst:
        raise HTTPException(status_code=404, detail="Institution not found")
    
    conn = get_db()
    if org_id and table_has_column("institutions", "org_id"):
        conn.execute(
            "UPDATE institutions SET name = ?, sla_hours = ?, modified_at = ? WHERE id = ? AND org_id = ?",
            (name.strip(), sla_val, utc_now_iso(), inst_id, org_id),
        )
    else:
        conn.execute("UPDATE institutions SET name = ?, sla_hours = ?, modified_at = ? WHERE id = ?", (name.strip(), sla_val, utc_now_iso(), inst_id))
    conn.commit()
    conn.close()
    
    return RedirectResponse(url="/settings", status_code=303)


@app.post("/settings/institution/delete/{inst_id}")
def delete_institution_route(request: Request, inst_id: int):
    user = require_admin(request)
    org_id = user.get("org_id")
    inst = get_institution(inst_id, org_id)
    if not inst:
        raise HTTPException(status_code=404, detail="Institution not found")
    
    delete_institution(inst_id, org_id)
    return RedirectResponse(url="/settings", status_code=303)


@app.post("/settings/radiologist/add")
def add_radiologist(request: Request, name: str = Form(...), email: str = Form(""), surname: str = Form(""), gmc: str = Form(""), speciality: str = Form("")):
    require_admin(request)
    conn = get_db()
    # Use upsert to allow both creation and update
    conn.execute(
        "INSERT INTO radiologists (name, email, surname, gmc, speciality) VALUES (?, ?, ?, ?, ?) "
        "ON CONFLICT(name) DO UPDATE SET email=excluded.email, surname=excluded.surname, gmc=excluded.gmc, speciality=excluded.speciality",
        (name.strip(), email.strip(), surname.strip(), gmc.strip(), speciality.strip())
    )
    conn.commit()
    conn.close()
    return RedirectResponse(url="/settings", status_code=303)


@app.post("/settings/radiologist/delete")
def remove_radiologist(request: Request, name: str = Form(...)):
    require_admin(request)
    delete_radiologist(name)
    return RedirectResponse(url="/settings", status_code=303)


@app.get("/settings/radiologist/edit/{name}", response_class=HTMLResponse)
def edit_radiologist_page(request: Request, name: str):
    try:
        require_admin(request)
    except HTTPException:
        return redirect_to_login("admin", "/settings")

    rad = get_radiologist(name)
    if not rad:
        raise HTTPException(status_code=404, detail="Radiologist not found")

    return templates.TemplateResponse("radiologist_edit.html", {"request": request, "rad": rad})


@app.post("/settings/radiologist/update")
def update_radiologist(request: Request, name: str = Form(...), email: str = Form(""), surname: str = Form(""), gmc: str = Form(""), speciality: str = Form("")):
    require_admin(request)
    conn = get_db()
    conn.execute(
        "UPDATE radiologists SET email = ?, surname = ?, gmc = ?, speciality = ? WHERE name = ?",
        (email.strip(), surname.strip(), gmc.strip(), speciality.strip(), name.strip())
    )
    conn.commit()
    conn.close()
    return RedirectResponse(url="/settings", status_code=303)


@app.post("/settings/protocol/add")
def settings_add_protocol(
    request: Request,
    institution_id: str = Form(...),
    study_description_preset_id: str = Form(...),
    instructions: str = Form(""),
):
    user = require_admin(request)
    org_id = user.get("org_id")
    
    if not institution_id or institution_id.strip() == "":
        raise HTTPException(status_code=400, detail="Please select an institution")
    
    try:
        inst_id = int(institution_id)
        inst = get_institution(inst_id, org_id)
        if not inst:
            raise HTTPException(status_code=400, detail="Invalid institution or institution not found")
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid institution ID format")

    try:
        preset_id = int(study_description_preset_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Please select a protocol name")

    preset = get_study_description_preset(preset_id, org_id)
    if not preset:
        raise HTTPException(status_code=400, detail="Invalid study description")
    protocol_name = str(preset.get("description") or "").strip()
    if not protocol_name:
        raise HTTPException(status_code=400, detail="Protocol name is required")
    
    conn = get_db()
    try:
        if org_id and table_has_column("protocols", "org_id"):
            if using_postgres():
                conn.execute(
                    """
                    INSERT INTO protocols (name, institution_id, study_description_preset_id, instructions, last_modified, is_active, org_id)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                    ON CONFLICT (name, institution_id) DO UPDATE SET
                      study_description_preset_id = EXCLUDED.study_description_preset_id,
                      instructions = EXCLUDED.instructions,
                      last_modified = EXCLUDED.last_modified,
                      is_active = EXCLUDED.is_active,
                      org_id = EXCLUDED.org_id
                    """,
                    (protocol_name, inst_id, preset_id, instructions.strip(), datetime.now().isoformat(), 1, org_id)
                )
            else:
                conn.execute(
                    "INSERT OR REPLACE INTO protocols (name, institution_id, study_description_preset_id, instructions, last_modified, is_active, org_id) VALUES (?, ?, ?, ?, ?, ?, ?)",
                    (protocol_name, inst_id, preset_id, instructions.strip(), datetime.now().isoformat(), 1, org_id)
                )
        else:
            if using_postgres():
                conn.execute(
                    """
                    INSERT INTO protocols (name, institution_id, study_description_preset_id, instructions, last_modified, is_active)
                    VALUES (?, ?, ?, ?, ?, ?)
                    ON CONFLICT (name, institution_id) DO UPDATE SET
                      study_description_preset_id = EXCLUDED.study_description_preset_id,
                      instructions = EXCLUDED.instructions,
                      last_modified = EXCLUDED.last_modified,
                      is_active = EXCLUDED.is_active
                    """,
                    (protocol_name, inst_id, preset_id, instructions.strip(), datetime.now().isoformat(), 1)
                )
            else:
                conn.execute(
                    "INSERT OR REPLACE INTO protocols (name, institution_id, study_description_preset_id, instructions, last_modified, is_active) VALUES (?, ?, ?, ?, ?, ?)",
                    (protocol_name, inst_id, preset_id, instructions.strip(), datetime.now().isoformat(), 1)
                )
        conn.commit()
    except Exception as e:
        if hasattr(conn, "rollback"):
            conn.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to add protocol: {str(e)}")
    finally:
        conn.close()
    return RedirectResponse(url="/settings", status_code=303)


@app.post("/settings/protocol/delete")
def settings_delete_protocol(request: Request, name: str = Form(...)):
    user = require_admin(request)
    org_id = user.get("org_id")
    deactivate_protocol(name, org_id)
    return RedirectResponse(url="/settings", status_code=303)


@app.post("/settings/user/add")
def add_user(
    request: Request,
    username: str = Form(...),
    password: str = Form(...),
    role: str = Form(...),
    radiologist_name: str = Form(""),
    first_name: str = Form(""),
    surname: str = Form(""),
    email: str = Form(""),
    gmc: str = Form(""),
    speciality: str = Form(""),
    mfa_required: str = Form("0"),
):
    user = require_admin(request)
    # get_current_org_context is already called inside require_admin; fetch org_id from it
    # to ensure we always have it even if the session dict wasn't updated yet.
    _uid, _su, org_id, _role = get_current_org_context(request)
    if not org_id:
        org_id = user.get("org_id")
    username = username.strip()
    role = role.strip()
    radiologist_name = radiologist_name.strip() or None
    gmc = gmc.strip()
    speciality = speciality.strip()
    mfa_required_value = 1 if str(mfa_required).strip().lower() in {"1", "true", "on", "yes"} else 0

    # For radiologist users, create profile with their name
    if role == "radiologist":
        if not first_name.strip():
            raise HTTPException(status_code=400, detail="Radiologist must have a first name")
        radiologist_name = f"{first_name.strip()} {surname.strip()}".strip() or first_name.strip()
    
    # For admin users, radiologist_name should be None
    if role == "admin" or role == "user":
        radiologist_name = None

    # Extended schema: create user plus membership when org-scoped records exist
    if table_has_column("users", "is_superuser") and org_id:
        salt = secrets.token_bytes(16)
        pw_hash = hash_password(password, salt)
        now = utc_now_iso()

        conn = get_db()
        try:
            email_val = email.strip() or None  # store NULL not '' to avoid UNIQUE constraint clashes
            if using_postgres():
                user_row = conn.execute(
                    """
                    INSERT INTO users(username, email, password_hash, salt_hex, is_superuser, is_active, created_at, modified_at, first_name, surname)
                    VALUES(?, ?, ?, ?, 0, 1, ?, ?, ?, ?)
                    RETURNING id
                    """,
                    (username, email_val, pw_hash.hex(), salt.hex(), now, now, first_name.strip(), surname.strip()),
                ).fetchone()
                user_id = user_row["id"] if isinstance(user_row, dict) else user_row[0]
            else:
                try:
                    conn.execute(
                        """
                        INSERT INTO users(username, email, password_hash, salt_hex, is_superuser, is_active, created_at, modified_at, first_name, surname)
                        VALUES(?, ?, ?, ?, 0, 1, ?, ?, ?, ?)
                        """,
                        (username, email_val, pw_hash.hex(), salt.hex(), now, now, first_name.strip(), surname.strip()),
                    )
                except Exception as _insert_err:
                    _msg = str(_insert_err).lower()
                    if "unique" in _msg and "username" in _msg:
                        return RedirectResponse(url="/settings?error=username_taken", status_code=303)
                    if "unique" in _msg and "email" in _msg:
                        return RedirectResponse(url="/settings?error=email_taken", status_code=303)
                    raise

                user_row = conn.execute("SELECT id FROM users WHERE username = ?", (username,)).fetchone()
                user_id = user_row["id"] if user_row else None
            if not user_id:
                raise HTTPException(status_code=500, detail="Failed to create user")

            conn.execute(
                "UPDATE users SET mfa_required = ?, modified_at = ? WHERE id = ?",
                (mfa_required_value, now, user_id),
            )

            org_role = "org_admin" if role == "admin" else "radiologist" if role == "radiologist" else "org_user"
            conn.execute(
                """
                INSERT INTO memberships (org_id, user_id, org_role, is_active, created_at, modified_at)
                VALUES (?, ?, ?, 1, ?, ?)
                """,
                (org_id, user_id, org_role, now, now),
            )

            if role == "radiologist":
                display = radiologist_name or f"{first_name.strip()} {surname.strip()}".strip() or username
                if using_postgres():
                    conn.execute(
                        """
                        INSERT INTO radiologists(name, first_name, email, surname, gmc, speciality)
                        VALUES(?, ?, ?, ?, ?, ?)
                        ON CONFLICT DO NOTHING
                        """,
                        (display, first_name.strip(), email.strip(), surname.strip(), gmc, speciality),
                    )
                    conn.execute(
                        """
                        INSERT INTO radiologist_profiles(user_id, gmc, specialty, display_name, created_at, modified_at)
                        VALUES(?, ?, ?, ?, ?, ?)
                        ON CONFLICT (user_id) DO NOTHING
                        """,
                        (user_id, gmc or None, speciality or None, display, now, now),
                    )
                else:
                    conn.execute(
                        """
                        INSERT OR IGNORE INTO radiologists(name, first_name, email, surname, gmc, speciality)
                        VALUES(?, ?, ?, ?, ?, ?)
                        """,
                        (display, first_name.strip(), email.strip(), surname.strip(), gmc, speciality),
                    )
                    conn.execute(
                        """
                        INSERT OR IGNORE INTO radiologist_profiles(user_id, gmc, specialty, display_name, created_at, modified_at)
                        VALUES(?, ?, ?, ?, ?, ?)
                        """,
                        (user_id, gmc or None, speciality or None, display, now, now),
                    )

            conn.commit()
        finally:
            conn.close()
    elif table_has_column("users", "is_superuser"):
        # New schema but no org_id context — create user only (no membership row)
        salt = secrets.token_bytes(16)
        pw_hash = hash_password(password, salt)
        now = utc_now_iso()
        conn = get_db()
        email_val = email.strip() or None
        try:
            conn.execute(
                """
                INSERT INTO users(username, email, password_hash, salt_hex, is_superuser, is_active, created_at, modified_at, first_name, surname)
                VALUES(?, ?, ?, ?, 0, 1, ?, ?, ?, ?)
                """,
                (username, email_val, pw_hash.hex(), salt.hex(), now, now, first_name.strip(), surname.strip()),
            )
            conn.execute(
                "UPDATE users SET mfa_required = ?, modified_at = ? WHERE username = ?",
                (mfa_required_value, now, username),
            )
            conn.commit()
        finally:
            conn.close()
    else:
        create_user(username, password, role, radiologist_name, first_name, surname, email)
    return RedirectResponse(url="/settings", status_code=303)


@app.post("/settings/user/delete")
def remove_user(request: Request, username: str = Form(...)):
    require_admin(request)
    if username.strip() == "admin":
        raise HTTPException(status_code=400, detail="Cannot delete admin user")
    delete_user(username.strip())
    return RedirectResponse(url="/settings", status_code=303)


@app.post("/settings/user/edit")
def edit_user(
    request: Request,
    username: str = Form(...),
    first_name: str = Form(""),
    surname: str = Form(""),
    email: str = Form(""),
    role: str = Form(...),
    radiologist_name: str = Form(""),
    password: str = Form(""),
    mfa_required: str = Form("0"),
):
    user = require_admin(request)
    org_id = user.get("org_id")
    username = username.strip()
    role = role.strip()
    radiologist_name = radiologist_name.strip() or None
    mfa_required_value = 1 if str(mfa_required).strip().lower() in {"1", "true", "on", "yes"} else 0
    
    conn = get_db()
    user = conn.execute("SELECT * FROM users WHERE username = ?", (username,)).fetchone()
    conn.close()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Resolve current email to avoid false UniqueViolation when email is unchanged
    current_email = (user["email"] if isinstance(user, dict) else dict(user).get("email", "")) or ""
    new_email = email.strip()
    email_changed = new_email and new_email != current_email

    if email_changed:
        # Check if new email is already taken by another user
        conn2 = get_db()
        conflict = conn2.execute(
            "SELECT id FROM users WHERE email = ? AND username != ?",
            (new_email, username),
        ).fetchone()
        conn2.close()
        if conflict:
            return RedirectResponse(url="/settings?error=email_taken", status_code=303)
    
    if table_has_column("users", "is_superuser"):
        # Extended schema
        conn = get_db()
        if password.strip():
            salt = secrets.token_bytes(16)
            pw_hash = hash_password(password, salt)
            if email_changed:
                conn.execute(
                    "UPDATE users SET first_name = ?, surname = ?, email = ?, password_hash = ?, salt_hex = ? WHERE username = ?",
                    (first_name.strip(), surname.strip(), new_email, pw_hash.hex(), salt.hex(), username)
                )
            else:
                conn.execute(
                    "UPDATE users SET first_name = ?, surname = ?, password_hash = ?, salt_hex = ? WHERE username = ?",
                    (first_name.strip(), surname.strip(), pw_hash.hex(), salt.hex(), username)
                )
        else:
            if email_changed:
                conn.execute(
                    "UPDATE users SET first_name = ?, surname = ?, email = ? WHERE username = ?",
                    (first_name.strip(), surname.strip(), new_email, username)
                )
            else:
                conn.execute(
                    "UPDATE users SET first_name = ?, surname = ? WHERE username = ?",
                    (first_name.strip(), surname.strip(), username)
                )

        if org_id and table_exists("memberships"):
            org_role = "org_admin" if role == "admin" else "radiologist" if role == "radiologist" else "org_user"
            target = conn.execute("SELECT id FROM users WHERE username = ?", (username,)).fetchone()
            target_id = target["id"] if target else None
            if target_id:
                if mfa_required_value:
                    conn.execute(
                        "UPDATE users SET mfa_required = ?, modified_at = ? WHERE id = ?",
                        (1, utc_now_iso(), target_id),
                    )
                else:
                    conn.execute(
                        """
                        UPDATE users
                        SET mfa_required = 0, mfa_enabled = 0, mfa_secret = NULL, mfa_pending_secret = NULL, modified_at = ?
                        WHERE id = ?
                        """,
                        (utc_now_iso(), target_id),
                    )
                conn.execute(
                    "UPDATE memberships SET org_role = ?, modified_at = ? WHERE user_id = ? AND org_id = ? AND is_active = 1",
                    (org_role, utc_now_iso(), target_id, org_id),
                )
        else:
            if mfa_required_value:
                conn.execute(
                    "UPDATE users SET mfa_required = ? WHERE username = ?",
                    (1, username),
                )
            else:
                conn.execute(
                    "UPDATE users SET mfa_required = 0, mfa_enabled = 0, mfa_secret = NULL, mfa_pending_secret = NULL WHERE username = ?",
                    (username,),
                )
    else:
        # Legacy schema
        if password.strip():
            salt = secrets.token_bytes(16)
            pw_hash = hash_password(password, salt)
            conn = get_db()
            if email_changed:
                conn.execute(
                    "UPDATE users SET first_name = ?, surname = ?, email = ?, role = ?, radiologist_name = ?, salt_hex = ?, pw_hash_hex = ? WHERE username = ?",
                    (first_name.strip(), surname.strip(), new_email, role, radiologist_name, salt.hex(), pw_hash.hex(), username)
                )
            else:
                conn.execute(
                    "UPDATE users SET first_name = ?, surname = ?, role = ?, radiologist_name = ?, salt_hex = ?, pw_hash_hex = ? WHERE username = ?",
                    (first_name.strip(), surname.strip(), role, radiologist_name, salt.hex(), pw_hash.hex(), username)
                )
        else:
            conn = get_db()
            if email_changed:
                conn.execute(
                    "UPDATE users SET first_name = ?, surname = ?, email = ?, role = ?, radiologist_name = ? WHERE username = ?",
                    (first_name.strip(), surname.strip(), new_email, role, radiologist_name, username)
                )
            else:
                conn.execute(
                    "UPDATE users SET first_name = ?, surname = ?, role = ?, radiologist_name = ? WHERE username = ?",
                    (first_name.strip(), surname.strip(), role, radiologist_name, username)
                )
        if mfa_required_value:
            conn.execute(
                "UPDATE users SET mfa_required = ? WHERE username = ?",
                (1, username),
            )
        else:
            conn.execute(
                "UPDATE users SET mfa_required = 0, mfa_enabled = 0, mfa_secret = NULL, mfa_pending_secret = NULL WHERE username = ?",
                (username,),
            )
    
    conn.commit()
    conn.close()
    return RedirectResponse(url="/settings", status_code=303)


@app.post("/settings/user/access")
def update_user_access(
    request: Request,
    org_role: str = Form(...),
    user_id: int | None = Form(None),
    username: str = Form(""),
):
    user = require_admin(request)
    org_id = user.get("org_id")

    allowed_roles = {"org_admin", "radiologist", "org_user"}
    if org_role not in allowed_roles:
        raise HTTPException(status_code=400, detail="Invalid access level")

    conn = get_db()
    try:
        if table_exists("memberships") and org_id:
            if not user_id and username:
                row = conn.execute("SELECT id FROM users WHERE username = ?", (username.strip(),)).fetchone()
                user_id = row["id"] if row else None

            if not user_id:
                raise HTTPException(status_code=400, detail="User not found")

            conn.execute(
                "UPDATE memberships SET org_role = ?, modified_at = ? WHERE user_id = ? AND org_id = ? AND is_active = 1",
                (org_role, utc_now_iso(), user_id, org_id),
            )
        else:
            # Legacy fallback
            role = "admin" if org_role == "org_admin" else "radiologist" if org_role == "radiologist" else "user"
            conn.execute("UPDATE users SET role = ? WHERE username = ?", (role, username.strip()))

        conn.commit()
    finally:
        conn.close()

    return RedirectResponse(url="/settings", status_code=303)


@app.post("/settings/protocol/edit/{protocol_id}")
def edit_protocol(
    request: Request,
    protocol_id: int,
    institution_id: str = Form(...),
    study_description_preset_id: str = Form(...),
    instructions: str = Form(""),
):
    user = require_admin(request)
    org_id = user.get("org_id")
    try:
        inst_id = int(institution_id)
        inst = get_institution(inst_id, org_id)
        if not inst:
            raise HTTPException(status_code=400, detail="Invalid institution")
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid institution ID")

    try:
        preset_id = int(study_description_preset_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid protocol name")

    preset = get_study_description_preset(preset_id, org_id)
    if not preset:
        raise HTTPException(status_code=400, detail="Invalid study description")
    protocol_name = str(preset.get("description") or "").strip()
    if not protocol_name:
        raise HTTPException(status_code=400, detail="Protocol name is required")
    
    conn = get_db()
    if org_id and table_has_column("protocols", "org_id"):
        conn.execute(
            "UPDATE protocols SET name = ?, institution_id = ?, study_description_preset_id = ?, instructions = ?, last_modified = ? WHERE id = ? AND org_id = ?",
            (protocol_name, inst_id, preset_id, instructions.strip(), datetime.now().isoformat(), protocol_id, org_id)
        )
    else:
        conn.execute(
            "UPDATE protocols SET name = ?, institution_id = ?, study_description_preset_id = ?, instructions = ?, last_modified = ? WHERE id = ?",
            (protocol_name, inst_id, preset_id, instructions.strip(), datetime.now().isoformat(), protocol_id)
        )
    conn.commit()
    conn.close()
    return RedirectResponse(url="/settings", status_code=303)


@app.post("/settings/protocol/delete/{protocol_id}")
def delete_protocol_route(request: Request, protocol_id: int):
    user = require_admin(request)
    org_id = user.get("org_id")
    conn = get_db()
    if org_id and table_has_column("protocols", "org_id"):
        conn.execute("DELETE FROM protocols WHERE id = ? AND org_id = ?", (protocol_id, org_id))
    else:
        conn.execute("DELETE FROM protocols WHERE id = ?", (protocol_id,))
    conn.commit()
    conn.close()
    return RedirectResponse(url="/settings", status_code=303)

# -------------------------
# Study Description Presets (Multitenant - By Organization)
# -------------------------
@app.get("/api/study-descriptions/by-modality/{modality}")
def get_study_descriptions(modality: str, request: Request, org_id: str = None):
    """Get study description presets by modality for user's organization (searchable via form)"""
    # If org_id is provided as query parameter, use it; otherwise get from session
    if org_id:
        try:
            org_id = int(org_id)
        except (ValueError, TypeError):
            org_id = None
    
    if not org_id:
        # Get user's organization from session
        user = request.session.get("user")
        if not user:
            return []
        
        # Try both org_id and organization_id for backward compatibility
        org_id = user.get("org_id") or user.get("organization_id")
    
    if not org_id:
        return []
    
    conn = get_db()
    modality = modality.upper().strip()
    rows = conn.execute(
        "SELECT id, description FROM study_description_presets WHERE organization_id = ? AND modality = ? AND COALESCE(is_active, 1) = 1 ORDER BY description",
        (org_id, modality)
    ).fetchall()
    if not rows and org_id != 1:
        rows = conn.execute(
            "SELECT id, description FROM study_description_presets WHERE organization_id = 1 AND modality = ? AND COALESCE(is_active, 1) = 1 ORDER BY description",
            (modality,)
        ).fetchall()
    conn.close()
    return [dict(row) for row in rows]


@app.get("/api/protocols/by-study-description/{preset_id}")
def get_protocols_by_study_description(
    preset_id: int,
    request: Request,
    institution_id: str = "",
    org_id: str = "",
):
    target_org_id = None
    if org_id:
        try:
            target_org_id = int(org_id)
        except (ValueError, TypeError):
            target_org_id = None

    if not target_org_id:
        user = request.session.get("user") or {}
        target_org_id = user.get("org_id") or user.get("organization_id")

    target_institution_id = None
    if institution_id:
        try:
            target_institution_id = int(institution_id)
        except (ValueError, TypeError):
            target_institution_id = None

    rows = list_protocol_rows_for_study(
        preset_id,
        institution_id=target_institution_id,
        org_id=target_org_id,
        active_only=True,
    )
    return rows

@app.get("/settings/study-descriptions", response_class=HTMLResponse)
def study_descriptions_page(request: Request):
    """Admin page to manage study description presets for the current organization."""
    user = require_admin(request)
    org_id = user.get("org_id") or user.get("organization_id") or 1
    
    conn = get_db()
    presets = conn.execute(
        "SELECT id, modality, description, COALESCE(is_active, 1) AS is_active, created_at, updated_at FROM study_description_presets WHERE organization_id = ? ORDER BY modality, description",
        (org_id,)
    ).fetchall()
    conn.close()
    return templates.TemplateResponse("superuser_study_descriptions.html", {
        "request": request,
        "current_user": user,
        "presets": [dict(row) for row in presets],
        "modalities": ["MRI", "CT", "XR", "PET", "DEXA"]
    })

@app.post("/settings/study-descriptions/add")
def add_study_description(request: Request, modality: str = Form(...), description: str = Form(...), org_id: str = Form("")):
    """Add new study description preset for the current organization."""
    user = require_admin(request)
    target_org_id = user.get("org_id") or user.get("organization_id") or 1
    if org_id:
        try:
            target_org_id = int(org_id)
        except Exception:
            pass
    modality = modality.upper().strip()
    description = description.strip()
    
    if not modality or not description:
        return RedirectResponse(url="/settings?tab=study-presets&error=empty", status_code=303)
    
    creator_id = user.get("id") or 1
    try:
        creator_id = int(creator_id)
    except Exception:
        creator_id = 1

    conn = get_db()
    try:
        conn.execute(
            "INSERT INTO study_description_presets (organization_id, modality, description, is_active, created_at, updated_at, created_by) VALUES (?, ?, ?, ?, ?, ?, ?)",
            (target_org_id, modality, description, 1, datetime.now(timezone.utc).isoformat(), datetime.now(timezone.utc).isoformat(), creator_id)
        )
        conn.commit()
    except (sqlite3.IntegrityError, SQLAlchemyError):
        conn.close()
        return RedirectResponse(url="/settings?tab=study-presets&error=duplicate", status_code=303)
    finally:
        try:
            conn.close()
        except Exception:
            pass
    
    return RedirectResponse(url="/settings?tab=study-presets", status_code=303)

@app.post("/settings/study-descriptions/archive/{preset_id}")
def archive_study_description(request: Request, preset_id: int):
    """Archive study description preset from the current organization."""
    user = require_admin(request)
    org_id = user.get("org_id") or user.get("organization_id") or 1
    conn = get_db()
    conn.execute(
        "UPDATE study_description_presets SET is_active = 0, updated_at = ? WHERE id = ? AND organization_id = ?",
        (datetime.now(timezone.utc).isoformat(), preset_id, org_id),
    )
    conn.commit()
    conn.close()
    return RedirectResponse(url="/settings?tab=study-presets", status_code=303)


@app.post("/settings/study-descriptions/restore/{preset_id}")
def restore_study_description(request: Request, preset_id: int):
    """Restore archived study description preset for the current organization."""
    user = require_admin(request)
    org_id = user.get("org_id") or user.get("organization_id") or 1
    conn = get_db()
    conn.execute(
        "UPDATE study_description_presets SET is_active = 1, updated_at = ? WHERE id = ? AND organization_id = ?",
        (datetime.now(timezone.utc).isoformat(), preset_id, org_id),
    )
    conn.commit()
    conn.close()
    return RedirectResponse(url="/settings?tab=study-presets", status_code=303)

@app.post("/settings/study-descriptions/edit/{preset_id}")
def edit_study_description(request: Request, preset_id: int, modality: str = Form(...), description: str = Form(...)):
    """Edit study description preset for the current organization."""
    user = require_admin(request)
    org_id = user.get("org_id") or user.get("organization_id") or 1
    modality = modality.upper().strip()
    description = description.strip()
    
    if not modality or not description:
        return RedirectResponse(url="/settings?tab=study-presets&error=empty", status_code=303)
    
    try:
        conn = get_db()
        conn.execute(
            "UPDATE study_description_presets SET modality = ?, description = ?, updated_at = ? WHERE id = ? AND organization_id = ?",
            (modality, description, datetime.now(timezone.utc).isoformat(), preset_id, org_id)
        )
        conn.commit()
        conn.close()
    except sqlite3.IntegrityError:
        return RedirectResponse(url="/settings?tab=study-presets&error=duplicate", status_code=303)
    
    return RedirectResponse(url="/settings?tab=study-presets", status_code=303)

# -------------------------
# Admin submit
# -------------------------
@app.get("/intake/{org_id}", response_class=HTMLResponse)
def intake_form(request: Request, org_id: int, token: str = ""):
    token = (token or "").strip()
    expected = get_setting(f"intake_token:{org_id}", "")
    if not expected or token != expected:
        raise HTTPException(status_code=403, detail="Invalid intake token")

    institutions = list_institutions(org_id)
    return templates.TemplateResponse(
        "intake_submit.html",
        {"request": request, "institutions": institutions},
    )


@app.post("/intake/{org_id}")
async def intake_submit(
    request: Request,
    org_id: int,
    token: str = "",
    patient_first_name: str = Form(...),
    patient_surname: str = Form(...),
    patient_referral_id: str = Form(...),
    patient_dob: str = Form(""),
    institution_id: str = Form(...),
    study_description: str = Form(...),
    admin_notes: str = Form(""),
    attachment: UploadFile | None = File(...),
):
    token = (token or request.query_params.get("token") or "").strip()
    expected = get_setting(f"intake_token:{org_id}", "")
    if not expected or token != expected:
        raise HTTPException(status_code=403, detail="Invalid intake token")

    try:
        inst_id = int(institution_id)
        inst = get_institution(inst_id, org_id)
        if not inst:
            raise HTTPException(status_code=400, detail="Invalid institution selection")
    except (ValueError, TypeError):
        raise HTTPException(status_code=400, detail="Invalid institution ID")

    if not attachment or not attachment.filename:
        raise HTTPException(status_code=400, detail="Attachment is required")

    case_id = generate_case_id(inst_id)
    original_name = attachment.filename
    
    file_bytes = await attachment.read()
    
    # Try blob storage first, fallback to local
    stored_path = None
    if BLOB_STORAGE_ENABLED:
        blob_name = upload_to_blob(case_id, file_bytes, original_name)
        if blob_name:
            stored_path = blob_name
    
    # Fallback: store locally if blob upload failed or disabled
    if not stored_path:
        safe_name = f"{case_id}_{Path(original_name).name}"
        stored_path = str(UPLOAD_DIR / safe_name)
        with open(stored_path, "wb") as f:
            f.write(file_bytes)

    created_at = utc_now_iso()
    conn = get_db()
    
    # Build insert conditionally based on columns that exist
    has_dob_col = table_has_column("cases", "patient_dob")
    if has_dob_col:
        conn.execute(
            "INSERT INTO cases (id, created_at, patient_first_name, patient_surname, patient_referral_id, patient_dob, institution_id, study_description, admin_notes, radiologist, uploaded_filename, stored_filepath, status, vetted_at, org_id) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (case_id, created_at, patient_first_name.strip(), patient_surname.strip(), patient_referral_id.strip(), patient_dob.strip() or None, inst_id, study_description.strip(), admin_notes.strip(), "", original_name, stored_path, "pending", None, org_id),
        )
    else:
        # PostgreSQL doesn't have patient_dob column, exclude it
        conn.execute(
            "INSERT INTO cases (id, created_at, patient_first_name, patient_surname, patient_referral_id, institution_id, study_description, admin_notes, radiologist, uploaded_filename, stored_filepath, status, vetted_at, org_id) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (case_id, created_at, patient_first_name.strip(), patient_surname.strip(), patient_referral_id.strip(), inst_id, study_description.strip(), admin_notes.strip(), "", original_name, stored_path, "pending", None, org_id),
        )
    conn.commit()
    conn.close()

    insert_case_event(
        case_id=case_id,
        org_id=org_id,
        event_type="SUBMITTED",
        user={"username": "external"},
        comment=admin_notes.strip() or None,
    )

    return RedirectResponse(url="/", status_code=303)


@app.get("/submit/referral-trial", response_class=HTMLResponse)
def referral_trial_form(request: Request):
    user = require_admin(request)
    org_id = user.get("org_id")
    institutions = list_institutions(org_id)
    radiologists = list_radiologists(org_id)

    return templates.TemplateResponse(
        "referral_trial.html",
        {
            "request": request,
            "institutions": institutions,
            "radiologists": radiologists,
            "user_org_id": org_id,
            "draft": {
                "patient_first_name": "",
                "patient_surname": "",
                "patient_referral_id": "",
                "patient_dob": "",
                "study_description": "",
                "modality": "",
                "admin_notes": "",
                "radiologist": "",
                "institution_id": "",
                "attachment_token": "",
                "attachment_original_name": "",
            },
            "parse_warnings": [],
            "parse_confidence": None,
            "parse_preview": "",
            "error": "",
        },
    )


@app.post("/submit/referral-trial/parse", response_class=HTMLResponse)
async def referral_trial_parse(
    request: Request,
    institution_id: str = Form(""),
    attachment: UploadFile | None = File(...),
):
    user = require_admin(request)
    org_id = user.get("org_id")
    institutions = list_institutions(org_id)
    radiologists = list_radiologists(org_id)

    if not attachment or not attachment.filename:
        return templates.TemplateResponse(
            "referral_trial.html",
            {
                "request": request,
                "institutions": institutions,
                "radiologists": radiologists,
                "draft": {"institution_id": institution_id or ""},
                "parse_warnings": [],
                "parse_confidence": None,
                "parse_preview": "",
                "error": "Please select a referral file to parse.",
            },
        )

    file_bytes = await attachment.read()
    parsed = parse_referral_attachment(attachment.filename, file_bytes)

    temp_name = f"trial_{uuid4().hex}_{Path(attachment.filename).name}"
    temp_path = UPLOAD_DIR / temp_name
    with open(temp_path, "wb") as temp_file:
        temp_file.write(file_bytes)

    draft = parsed.get("fields", {})
    draft["institution_id"] = (institution_id or "").strip()
    draft["radiologist"] = ""
    draft["attachment_token"] = temp_name
    draft["attachment_original_name"] = attachment.filename

    return templates.TemplateResponse(
        "referral_trial.html",
        {
            "request": request,
            "institutions": institutions,
            "radiologists": radiologists,
            "user_org_id": org_id,
            "draft": draft,
            "parse_warnings": parsed.get("warnings", []),
            "parse_confidence": parsed.get("confidence"),
            "parse_preview": parsed.get("text_preview", ""),
            "error": "",
        },
    )


@app.post("/submit/referral-trial/create")
def referral_trial_create(
    request: Request,
    patient_first_name: str = Form(""),
    patient_surname: str = Form(""),
    patient_referral_id: str = Form(""),
    patient_dob: str = Form(""),
    institution_id: str = Form(...),
    modality: str = Form(""),
    study_description: str = Form(...),
    admin_notes: str = Form(""),
    radiologist: str = Form(""),
    attachment_token: str = Form(...),
    attachment_original_name: str = Form("referral_upload"),
):
    user = require_admin(request)
    org_id = user.get("org_id")

    if not patient_first_name.strip() or not patient_surname.strip() or not patient_referral_id.strip() or not study_description.strip():
        raise HTTPException(status_code=400, detail="Patient name, referral ID, and study description are required")

    try:
        inst_id = int(institution_id)
        inst = get_institution(inst_id, org_id)
        if not inst:
            raise HTTPException(status_code=400, detail="Invalid institution selection")
    except (ValueError, TypeError):
        raise HTTPException(status_code=400, detail="Invalid institution ID")

    radiologist = radiologist.strip() if radiologist else ""
    if radiologist:
        valid_rads = {r["name"] for r in list_radiologists(org_id)}
        if radiologist not in valid_rads:
            raise HTTPException(status_code=400, detail="Invalid radiologist selection")

    token_name = Path(attachment_token).name
    if not token_name.startswith("trial_"):
        raise HTTPException(status_code=400, detail="Invalid attachment token")

    temp_path = UPLOAD_DIR / token_name
    if not temp_path.exists():
        raise HTTPException(status_code=400, detail="Trial attachment not found. Please parse again.")

    with open(temp_path, "rb") as temp_file:
        file_bytes = temp_file.read()

    case_id = generate_case_id()
    original_name = (attachment_original_name or "referral_upload").strip() or "referral_upload"

    stored_path = None
    if BLOB_STORAGE_ENABLED:
        blob_name = upload_to_blob(case_id, file_bytes, original_name)
        if blob_name:
            stored_path = blob_name

    if not stored_path:
        safe_name = f"{case_id}_{Path(original_name).name}"
        stored_path = str(UPLOAD_DIR / safe_name)
        with open(stored_path, "wb") as f:
            f.write(file_bytes)

    try:
        temp_path.unlink(missing_ok=True)
    except Exception:
        pass

    created_at = utc_now_iso()
    conn = get_db()

    has_dob_col = table_has_column("cases", "patient_dob")
    has_modality_col = table_has_column("cases", "modality")
    case_modality = modality.strip().upper() if modality else None

    if has_dob_col and has_modality_col:
        conn.execute(
            "INSERT INTO cases (id, created_at, patient_first_name, patient_surname, patient_referral_id, patient_dob, institution_id, study_description, modality, admin_notes, radiologist, uploaded_filename, stored_filepath, status, vetted_at, org_id) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (case_id, created_at, patient_first_name.strip(), patient_surname.strip(), patient_referral_id.strip(), patient_dob.strip() or None, inst_id, study_description.strip(), case_modality, admin_notes.strip(), radiologist, original_name, stored_path, "pending", None, org_id),
        )
    elif has_dob_col:
        conn.execute(
            "INSERT INTO cases (id, created_at, patient_first_name, patient_surname, patient_referral_id, patient_dob, institution_id, study_description, admin_notes, radiologist, uploaded_filename, stored_filepath, status, vetted_at, org_id) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (case_id, created_at, patient_first_name.strip(), patient_surname.strip(), patient_referral_id.strip(), patient_dob.strip() or None, inst_id, study_description.strip(), admin_notes.strip(), radiologist, original_name, stored_path, "pending", None, org_id),
        )
    elif has_modality_col:
        conn.execute(
            "INSERT INTO cases (id, created_at, patient_first_name, patient_surname, patient_referral_id, institution_id, study_description, modality, admin_notes, radiologist, uploaded_filename, stored_filepath, status, vetted_at, org_id) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (case_id, created_at, patient_first_name.strip(), patient_surname.strip(), patient_referral_id.strip(), inst_id, study_description.strip(), case_modality, admin_notes.strip(), radiologist, original_name, stored_path, "pending", None, org_id),
        )
    else:
        conn.execute(
            "INSERT INTO cases (id, created_at, patient_first_name, patient_surname, patient_referral_id, institution_id, study_description, admin_notes, radiologist, uploaded_filename, stored_filepath, status, vetted_at, org_id) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (case_id, created_at, patient_first_name.strip(), patient_surname.strip(), patient_referral_id.strip(), inst_id, study_description.strip(), admin_notes.strip(), radiologist, original_name, stored_path, "pending", None, org_id),
        )

    conn.commit()
    conn.close()

    insert_case_event(
        case_id=case_id,
        org_id=org_id,
        event_type="SUBMITTED",
        user={"username": user.get("username") or "admin"},
        comment="Created via referral trial parser",
    )

    return RedirectResponse(url=f"/submitted/{case_id}", status_code=303)


@app.get("/submit", response_class=HTMLResponse)
def submit_form(request: Request):
    try:
        user = require_admin(request)
    except HTTPException:
        return redirect_to_login("admin", "/submit")

    org_id = user.get("org_id")
    institutions = list_institutions(org_id)
    radiologists = list_radiologists(org_id)
    
    # Keep institution org_id available for existing study description filtering
    for inst in institutions:
        if "org_id" not in inst and org_id:
            inst["org_id"] = org_id
    
    return templates.TemplateResponse(
        "submit.html",
        {
            "request": request,
            "institutions": institutions,
            "radiologists": radiologists,
            "user_org_id": org_id,
        },
    )


@app.post("/submit")
async def submit_case(
    request: Request,
    patient_first_name: str = Form(...),
    patient_surname: str = Form(...),
    patient_referral_id: str = Form(...),
    patient_dob: str = Form(""),
    institution_id: str = Form(...),
    org_id_form: str = Form(""),
    modality: str = Form(""),
    study_description: str = Form(...),
    admin_notes: str = Form(""),
    radiologist: str = Form(""),
    attachment: UploadFile | None = File(...),
    action: str = Form("submit"),
    extra_study_description: list[str] = Form([]),
    extra_modality: list[str] = Form([]),
    extra_radiologist: list[str] = Form([]),
):
    user = require_admin(request)
    org_id = user.get("org_id")
    form_org_id = (org_id_form or "").strip()
    if table_exists("memberships") and not user.get("is_superuser") and not org_id:
        raise HTTPException(status_code=403, detail="Organisation access required")

    # Validate institution
    try:
        inst_id = int(institution_id)
        inst = get_institution(inst_id, org_id)
        if not inst:
            raise HTTPException(status_code=400, detail="Invalid institution selection")
    except (ValueError, TypeError):
        raise HTTPException(status_code=400, detail="Invalid institution ID")

    if user.get("is_superuser"):
        inst_org_id = inst.get("org_id") if isinstance(inst, dict) else None
        if inst_org_id:
            org_id = inst_org_id
        elif form_org_id:
            try:
                org_id = int(form_org_id)
            except ValueError:
                pass

    # Validate radiologist (optional - can assign later via bulk assignment)
    radiologist = radiologist.strip() if radiologist else ""
    if radiologist:
        valid_rads = {r["name"] for r in list_radiologists(org_id)}
        if radiologist not in valid_rads:
            raise HTTPException(status_code=400, detail="Invalid radiologist selection")

    # Validate attachment is provided
    if not attachment or not attachment.filename:
        raise HTTPException(status_code=400, detail="Attachment is required")

    cleaned_extra_cases: list[tuple[str, str | None, str]] = []
    if extra_study_description:
        valid_rads = {r["name"] for r in list_radiologists(org_id)}
        for i, extra_desc in enumerate(extra_study_description):
            normalized_desc = (extra_desc or "").strip()
            if not normalized_desc:
                continue
            normalized_rad = extra_radiologist[i].strip() if i < len(extra_radiologist) else ""
            if normalized_rad and normalized_rad not in valid_rads:
                normalized_rad = ""
            normalized_modality = (
                extra_modality[i].strip().upper()
                if i < len(extra_modality) and extra_modality[i].strip()
                else None
            )
            cleaned_extra_cases.append((normalized_desc, normalized_modality, normalized_rad))

    generated_case_ids = generate_case_ids(1 + len(cleaned_extra_cases), institution_id=inst_id)
    case_id = generated_case_ids[0]
    original_name = attachment.filename
    
    file_bytes = await attachment.read()
    
    # Try blob storage first, fallback to local
    stored_path = None
    if BLOB_STORAGE_ENABLED:
        blob_name = upload_to_blob(case_id, file_bytes, original_name)
        if blob_name:
            stored_path = blob_name
    
    # Fallback: store locally if blob upload failed or disabled
    if not stored_path:
        safe_name = f"{case_id}_{Path(original_name).name}"
        stored_path = str(UPLOAD_DIR / safe_name)
        with open(stored_path, "wb") as f:
            f.write(file_bytes)

    created_at = utc_now_iso()

    conn = get_db()
    
    # Build insert conditionally based on columns that exist
    has_dob_col = table_has_column("cases", "patient_dob")
    has_modality_col = table_has_column("cases", "modality")
    case_modality = modality.strip().upper() if modality else None
    if has_dob_col and has_modality_col:
        conn.execute(
            "INSERT INTO cases (id, created_at, patient_first_name, patient_surname, patient_referral_id, patient_dob, institution_id, study_description, modality, admin_notes, radiologist, uploaded_filename, stored_filepath, status, vetted_at, org_id) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (case_id, created_at, patient_first_name.strip(), patient_surname.strip(), patient_referral_id.strip(), patient_dob.strip() or None, inst_id, study_description.strip(), case_modality, admin_notes.strip(), radiologist, original_name, stored_path, "pending", None, org_id),
        )
    elif has_dob_col:
        conn.execute(
            "INSERT INTO cases (id, created_at, patient_first_name, patient_surname, patient_referral_id, patient_dob, institution_id, study_description, admin_notes, radiologist, uploaded_filename, stored_filepath, status, vetted_at, org_id) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (case_id, created_at, patient_first_name.strip(), patient_surname.strip(), patient_referral_id.strip(), patient_dob.strip() or None, inst_id, study_description.strip(), admin_notes.strip(), radiologist, original_name, stored_path, "pending", None, org_id),
        )
    elif has_modality_col:
        conn.execute(
            "INSERT INTO cases (id, created_at, patient_first_name, patient_surname, patient_referral_id, institution_id, study_description, modality, admin_notes, radiologist, uploaded_filename, stored_filepath, status, vetted_at, org_id) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (case_id, created_at, patient_first_name.strip(), patient_surname.strip(), patient_referral_id.strip(), inst_id, study_description.strip(), case_modality, admin_notes.strip(), radiologist, original_name, stored_path, "pending", None, org_id),
        )
    else:
        conn.execute(
            "INSERT INTO cases (id, created_at, patient_first_name, patient_surname, patient_referral_id, institution_id, study_description, admin_notes, radiologist, uploaded_filename, stored_filepath, status, vetted_at, org_id) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (case_id, created_at, patient_first_name.strip(), patient_surname.strip(), patient_referral_id.strip(), inst_id, study_description.strip(), admin_notes.strip(), radiologist, original_name, stored_path, "pending", None, org_id),
        )
    conn.commit()
    conn.close()

    # Create additional cases for extra studies (same patient/institution/attachment)
    if cleaned_extra_cases:
        for idx, (extra_desc, extra_modality_value, extra_rad) in enumerate(cleaned_extra_cases, start=1):
            extra_case_id = generated_case_ids[idx]
            # Copy attachment for the extra case
            extra_stored_path = stored_path
            if stored_path and original_name:
                if BLOB_STORAGE_ENABLED and not stored_path.startswith("/"):
                    # Blob storage: copy blob
                    blob_bytes = download_from_blob(stored_path)
                    if blob_bytes:
                        extra_blob_name = upload_to_blob(extra_case_id, blob_bytes, original_name)
                        if extra_blob_name:
                            extra_stored_path = extra_blob_name
                else:
                    # Local fallback: copy file
                    extra_safe_name = f"{extra_case_id}_{Path(original_name).name}"
                    extra_stored_path = str(UPLOAD_DIR / extra_safe_name)
                    shutil.copy2(stored_path, extra_stored_path)
            conn2 = get_db()
            
            # Build insert conditionally based on columns that exist
            if has_dob_col and has_modality_col:
                conn2.execute(
                    "INSERT INTO cases (id, created_at, patient_first_name, patient_surname, patient_referral_id, patient_dob, institution_id, study_description, modality, admin_notes, radiologist, uploaded_filename, stored_filepath, status, vetted_at, org_id) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                    (extra_case_id, utc_now_iso(), patient_first_name.strip(), patient_surname.strip(), patient_referral_id.strip(), patient_dob.strip() or None, inst_id, extra_desc, extra_modality_value, admin_notes.strip(), extra_rad, original_name, extra_stored_path, "pending", None, org_id),
                )
            elif has_dob_col:
                conn2.execute(
                    "INSERT INTO cases (id, created_at, patient_first_name, patient_surname, patient_referral_id, patient_dob, institution_id, study_description, admin_notes, radiologist, uploaded_filename, stored_filepath, status, vetted_at, org_id) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                    (extra_case_id, utc_now_iso(), patient_first_name.strip(), patient_surname.strip(), patient_referral_id.strip(), patient_dob.strip() or None, inst_id, extra_desc, admin_notes.strip(), extra_rad, original_name, extra_stored_path, "pending", None, org_id),
                )
            elif has_modality_col:
                conn2.execute(
                    "INSERT INTO cases (id, created_at, patient_first_name, patient_surname, patient_referral_id, institution_id, study_description, modality, admin_notes, radiologist, uploaded_filename, stored_filepath, status, vetted_at, org_id) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                    (extra_case_id, utc_now_iso(), patient_first_name.strip(), patient_surname.strip(), patient_referral_id.strip(), inst_id, extra_desc, extra_modality_value, admin_notes.strip(), extra_rad, original_name, extra_stored_path, "pending", None, org_id),
                )
            else:
                conn2.execute(
                    "INSERT INTO cases (id, created_at, patient_first_name, patient_surname, patient_referral_id, institution_id, study_description, admin_notes, radiologist, uploaded_filename, stored_filepath, status, vetted_at, org_id) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                    (extra_case_id, utc_now_iso(), patient_first_name.strip(), patient_surname.strip(), patient_referral_id.strip(), inst_id, extra_desc, admin_notes.strip(), extra_rad, original_name, extra_stored_path, "pending", None, org_id),
                )
            conn2.commit()
            conn2.close()

    # Redirect to admin dashboard
    return RedirectResponse(url="/admin", status_code=303)


@app.get("/submitted/{case_id}", response_class=HTMLResponse)
def submitted(request: Request, case_id: str):
    user = require_admin(request)
    conn = get_db()
    org_id = user.get("org_id")
    if org_id and not user.get("is_superuser"):
        row = conn.execute("SELECT * FROM cases WHERE id = ? AND org_id = ?", (case_id, org_id)).fetchone()
    else:
        row = conn.execute("SELECT * FROM cases WHERE id = ?", (case_id,)).fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="Case not found")
    return templates.TemplateResponse("submitted.html", {"request": request, "case": row})


# -------------------------
# iRefer Guidelines lookup (radiologist only)
# -------------------------

@app.get("/irefer/search")
def irefer_search(request: Request, q: str = ""):
    """Proxy iRefer guidelines search. Results are cached in memory after first fetch."""
    require_radiologist(request)  # radiologist-only endpoint

    global _irefer_guidelines_cache

    # Fetch and cache on first call
    if not _irefer_guidelines_cache and IREFER_API_KEY:
        try:
            url = "https://api.irefer.org.uk/prod-irefer/guidelines?language=en"
            req = urllib.request.Request(url, headers={"Ocp-Apim-Subscription-Key": IREFER_API_KEY})
            with urllib.request.urlopen(req, timeout=10) as resp:
                data = _json.loads(resp.read().decode())
                _irefer_guidelines_cache = data.get("value", [])
        except Exception as exc:
            return JSONResponse({"error": f"iRefer API error: {exc}", "results": []})

    if not _irefer_guidelines_cache:
        return JSONResponse({
            "error": "iRefer API key not configured. Set IREFER_API_KEY environment variable.",
            "results": []
        })

    # Filter by query
    q_lower = q.lower().strip()
    matches = []
    for g in _irefer_guidelines_cache:
        haystack = " ".join(filter(None, [
            g.get("ClinicalDiagnosticIssue", ""),
            g.get("SearchTerms", ""),
            " ".join(g.get("Section", []))
        ])).lower()
        if not q_lower or q_lower in haystack:
            investigations = [
                {
                    "investigation": inv.get("Investigation", ""),
                    "recommendation": inv.get("Recommendation", ""),
                    "grade": inv.get("Grade", ""),
                    "min_dose": inv.get("MinDose", ""),
                    "max_dose": inv.get("MaxDose", ""),
                    "comment": inv.get("Comment", ""),
                }
                for inv in g.get("Investigations", [])
            ]
            matches.append({
                "id": g.get("Id", ""),
                "code": g.get("Code", ""),
                "title": g.get("ClinicalDiagnosticIssue", ""),
                "section": g.get("Section", []),
                "body": g.get("Body", ""),
                "last_updated": g.get("LastUpdated", ""),
                "investigations": investigations,
            })
            if len(matches) >= 10:  # cap results for performance
                break

    return JSONResponse({"results": matches})


# -------------------------
# Vet (radiologist)
# -------------------------
@app.get("/vet/{case_id}", response_class=HTMLResponse)
def vet_form(request: Request, case_id: str):
    user = require_radiologist(request)
    rad_name = user.get("radiologist_name")
    org_id = user.get("org_id")
    org_name = None

    conn = get_db()
    if org_id:
        row = conn.execute(
            "SELECT c.*, i.name as institution_name FROM cases c LEFT JOIN institutions i ON c.institution_id = i.id WHERE c.id = ? AND c.org_id = ?",
            (case_id, org_id),
        ).fetchone()
    else:
        row = conn.execute("SELECT c.*, i.name as institution_name FROM cases c LEFT JOIN institutions i ON c.institution_id = i.id WHERE c.id = ?", (case_id,)).fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="Case not found")

    case = dict(row)
    case = normalize_case_attachment(case)
    base_admin_notes, reopened_admin_notes = partition_admin_notes(case.get("admin_notes"))
    case["admin_note_blocks"] = base_admin_notes
    case["reopened_note_blocks"] = reopened_admin_notes
    if case["radiologist"] != rad_name:
        raise HTTPException(status_code=403, detail="Not your case")

    if table_exists("case_events"):
        conn = get_db()
        last_event = conn.execute(
            "SELECT event_type, username, created_at FROM case_events WHERE case_id = ? ORDER BY created_at DESC LIMIT 1",
            (case_id,),
        ).fetchone()
        conn.close()
        should_log_open = True
        if last_event:
            last_type = str(last_event["event_type"] if isinstance(last_event, dict) else last_event[0] or "").upper()
            last_user = str(last_event["username"] if isinstance(last_event, dict) else last_event[1] or "")
            last_at = parse_iso_dt(last_event["created_at"] if isinstance(last_event, dict) else last_event[2])
            if last_type == "OPENED" and last_user == (user.get("username") or "") and last_at:
                should_log_open = (datetime.now(timezone.utc) - last_at) > timedelta(minutes=2)
        if should_log_open:
            insert_case_event(
                case_id=case_id,
                org_id=org_id,
                event_type="OPENED",
                user=user,
                comment="Case opened by practitioner",
            )

    protocols = []
    preset_id = None
    if case.get("study_description") and case.get("modality"):
        conn = get_db()
        preset_row = conn.execute(
            """
            SELECT id
            FROM study_description_presets
            WHERE organization_id = ? AND modality = ? AND description = ?
            LIMIT 1
            """,
            ((org_id or 1), str(case.get("modality") or "").strip().upper(), str(case.get("study_description") or "").strip()),
        ).fetchone()
        if not preset_row and (org_id or 1) != 1:
            preset_row = conn.execute(
                """
                SELECT id
                FROM study_description_presets
                WHERE organization_id = 1 AND modality = ? AND description = ?
                LIMIT 1
                """,
                (str(case.get("modality") or "").strip().upper(), str(case.get("study_description") or "").strip()),
            ).fetchone()
        conn.close()
        if preset_row:
            preset_id = preset_row["id"] if isinstance(preset_row, dict) else preset_row[0]

    if preset_id:
        protocols = list_protocol_rows_for_study(
            preset_id,
            institution_id=case.get("institution_id"),
            org_id=org_id,
            active_only=True,
        )

    if not protocols and case.get("institution_id"):
        conn = get_db()
        if org_id and table_has_column("protocols", "org_id"):
            proto_rows = conn.execute(
                "SELECT name, instructions FROM protocols WHERE institution_id = ? AND is_active = 1 AND org_id = ? ORDER BY name",
                (case.get("institution_id"), org_id)
            ).fetchall()
        else:
            proto_rows = conn.execute(
                "SELECT name, instructions FROM protocols WHERE institution_id = ? AND is_active = 1 ORDER BY name",
                (case.get("institution_id"),)
            ).fetchall()
        conn.close()
        protocols = [dict(p) for p in proto_rows]
    elif not protocols:
        protocols = [{"name": p, "instructions": ""} for p in list_protocols(active_only=True, org_id=org_id)]

    if org_id:
        conn = get_db()
        org_row = conn.execute("SELECT name FROM organisations WHERE id = ?", (org_id,)).fetchone()
        conn.close()
        if org_row:
            org_name = org_row.get("name") if isinstance(org_row, dict) else org_row[0]

    return templates.TemplateResponse(
        "vet.html",
        {
            "request": request,
            "case": case,
            "decisions": DECISIONS,
            "protocols": protocols,
            "org_name": org_name,
        },
    )


@app.post("/vet/{case_id}")
def vet_submit(
    request: Request,
    case_id: str,
    protocol: str = Form(""),
    decision: str = Form(...),
    decision_comment: str = Form(""),
    contrast_required: str = Form(""),
    contrast_details: str = Form(""),
):
    user = require_radiologist(request)
    rad_name = user.get("radiologist_name")
    org_id = user.get("org_id")

    if decision not in DECISIONS:
        raise HTTPException(status_code=400, detail="Invalid decision")

    # If decision is "Reject", comment is mandatory and protocol is not required
    if decision == "Reject":
        if not decision_comment.strip():
            raise HTTPException(status_code=400, detail="Comment is required when rejecting a case")
        protocol = ""  # Clear protocol for rejected cases
    else:
        # For Approve/Approve with comment, protocol is required
        if not protocol.strip():
            raise HTTPException(status_code=400, detail="Protocol is required for approved cases")

    conn = get_db()
    if org_id:
        row = conn.execute("SELECT radiologist FROM cases WHERE id = ? AND org_id = ?", (case_id, org_id)).fetchone()
    else:
        row = conn.execute("SELECT radiologist FROM cases WHERE id = ?", (case_id,)).fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Case not found")
    if row["radiologist"] != rad_name:
        conn.close()
        raise HTTPException(status_code=403, detail="Not your case")

    # Determine status based on decision
    if decision == "Reject":
        case_status = "rejected"
    else:
        case_status = "vetted"

    conn.execute(
        """
        UPDATE cases
        SET status = ?,
            protocol = ?,
            decision = ?,
            decision_comment = ?,
            vetted_at = ?,
            contrast_required = ?,
            contrast_details = ?
        WHERE id = ?
        """,
        (case_status, protocol.strip(), decision, decision_comment.strip(), utc_now_iso(),
         contrast_required.strip() or None, contrast_details.strip() or None, case_id),
    )
    conn.commit()
    conn.close()

    insert_case_event(
        case_id=case_id,
        org_id=org_id,
        event_type="VETTED",
        user=user,
        decision=decision,
        protocol=protocol.strip() or None,
        comment=decision_comment.strip() or None,
    )

    return RedirectResponse(url="/radiologist", status_code=303)


# -------------------------
# Attachments + PDF
# -------------------------
@app.get("/case/{case_id}/attachment")
def download_attachment(request: Request, case_id: str):
    user = require_login(request)

    if user["role"] == "radiologist":
        raise HTTPException(status_code=403, detail="Radiologists are not allowed to download attachments")

    conn = get_db()
    row = conn.execute(
        "SELECT stored_filepath, uploaded_filename, radiologist, org_id FROM cases WHERE id = ?",
        (case_id,),
    ).fetchone()
    conn.close()

    if not row:
        raise HTTPException(status_code=404, detail="No attachment found")

    case_data = row if isinstance(row, dict) else dict(row)

    if not case_data.get("stored_filepath"):
        raise HTTPException(status_code=410, detail="Referral file has expired and is no longer available (7-day retention policy).")

    if user.get("role") == "radiologist" and case_data.get("radiologist") != user.get("radiologist_name"):
        raise HTTPException(status_code=403, detail="Not your case")

    org_id = user.get("org_id")
    if org_id and not user.get("is_superuser") and case_data.get("org_id") and case_data.get("org_id") != org_id:
        raise HTTPException(status_code=403, detail="Access denied")

    stored_path = case_data.get("stored_filepath")
    
    # Try blob storage first
    file_bytes = None
    if BLOB_STORAGE_ENABLED and stored_path and not stored_path.startswith("/"):
        file_bytes = download_from_blob(stored_path)
        if file_bytes:
            return FileResponse(
                io.BytesIO(file_bytes),
                filename=case_data.get("uploaded_filename") or Path(stored_path).name
            )
    
    # Fallback to local filesystem
    if os.path.exists(stored_path):
        return FileResponse(stored_path, filename=case_data.get("uploaded_filename") or Path(stored_path).name)
    
    # File not found
    clear_case_stored_filepath(case_id)
    raise HTTPException(status_code=410, detail="Referral file missing or expired")


@app.get("/case/{case_id}/attachment/inline")
def view_attachment_inline(request: Request, case_id: str):
    user = require_login(request)

    conn = get_db()
    row = conn.execute(
        "SELECT stored_filepath, uploaded_filename, radiologist, org_id FROM cases WHERE id = ?",
        (case_id,),
    ).fetchone()
    conn.close()

    if not row:
        raise HTTPException(status_code=404, detail="No attachment found")

    case_data = row if isinstance(row, dict) else dict(row)

    if not case_data.get("stored_filepath"):
        raise HTTPException(status_code=410, detail="Referral file has expired and is no longer available (7-day retention policy).")

    if user.get("role") == "radiologist" and case_data.get("radiologist") != user.get("radiologist_name"):
        raise HTTPException(status_code=403, detail="Not your case")

    org_id = user.get("org_id")
    if org_id and not user.get("is_superuser") and case_data.get("org_id") and case_data.get("org_id") != org_id:
        raise HTTPException(status_code=403, detail="Access denied")

    stored_path = case_data.get("stored_filepath")
    filename = case_data.get("uploaded_filename") or Path(stored_path).name
    media_type, _ = mimetypes.guess_type(filename)
    headers = {"Content-Disposition": f'inline; filename="{filename}"'}
    
    # Try blob storage first
    file_bytes = None
    if BLOB_STORAGE_ENABLED and stored_path and not stored_path.startswith("/"):
        file_bytes = download_from_blob(stored_path)
        if file_bytes:
            return FileResponse(
                io.BytesIO(file_bytes),
                media_type=media_type or "application/octet-stream",
                headers=headers
            )
    
    # Fallback to local filesystem
    if os.path.exists(stored_path):
        return FileResponse(stored_path, media_type=media_type or "application/octet-stream", headers=headers)
    
    # File not found
    clear_case_stored_filepath(case_id)
    raise HTTPException(status_code=410, detail="Referral file missing or expired")


@app.get("/case/{case_id}/attachment/preview", response_class=HTMLResponse)
def view_attachment_preview(request: Request, case_id: str):
    user = require_login(request)

    conn = get_db()
    row = conn.execute(
        "SELECT stored_filepath, uploaded_filename, radiologist, org_id FROM cases WHERE id = ?",
        (case_id,),
    ).fetchone()
    conn.close()

    if not row:
        raise HTTPException(status_code=404, detail="No attachment found")

    case_data = row if isinstance(row, dict) else dict(row)

    if not case_data.get("stored_filepath"):
        raise HTTPException(status_code=410, detail="Referral file has expired and is no longer available (7-day retention policy).")

    if user.get("role") == "radiologist" and case_data.get("radiologist") != user.get("radiologist_name"):
        raise HTTPException(status_code=403, detail="Not your case")

    org_id = user.get("org_id")
    if org_id and not user.get("is_superuser") and case_data.get("org_id") and case_data.get("org_id") != org_id:
        raise HTTPException(status_code=403, detail="Access denied")

    stored_path = case_data.get("stored_filepath")
    filename = case_data.get("uploaded_filename") or Path(stored_path).name
    lower_name = str(filename).lower()
    media_type, _ = mimetypes.guess_type(filename)

    if lower_name.endswith(".pdf"):
        return HTMLResponse(
            f"""<!doctype html>
<html><head><meta charset="utf-8"><style>html,body{{height:100%;margin:0;background:#0b1220}}iframe{{width:100%;height:100%;border:0;background:#fff}}</style></head>
<body><iframe src="/case/{case_id}/attachment/inline#view=FitH"></iframe></body></html>"""
        )

    if lower_name.endswith((".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp")):
        return HTMLResponse(
            f"""<!doctype html>
<html><head><meta charset="utf-8"><style>html,body{{height:100%;margin:0;background:#0b1220}}body{{display:flex;align-items:center;justify-content:center;padding:12px;box-sizing:border-box}}img{{max-width:100%;max-height:100%;object-fit:contain;background:#fff;border-radius:8px}}</style></head>
<body><img src="/case/{case_id}/attachment/inline" alt="{html.escape(filename)}"></body></html>"""
        )

    file_bytes = load_case_attachment_bytes(stored_path)
    if file_bytes is None:
        clear_case_stored_filepath(case_id)
        raise HTTPException(status_code=410, detail="Referral file missing or expired")

    if lower_name.endswith((".txt", ".csv", ".json", ".xml", ".html", ".htm")):
        try:
            text_content = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            text_content = file_bytes.decode("latin-1", errors="replace")
        safe_text = html.escape(text_content)
        return HTMLResponse(
            f"""<!doctype html>
<html><head><meta charset="utf-8"><style>html,body{{margin:0;background:#fff;color:#0f172a;font-family:Segoe UI,Arial,sans-serif}}pre{{margin:0;padding:18px;white-space:pre-wrap;word-break:break-word;font-size:14px;line-height:1.5}}</style></head>
<body><pre>{safe_text}</pre></body></html>"""
        )

    if lower_name.endswith(".docx"):
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            blocks: list[str] = []
            for para in doc.paragraphs:
                text_value = (para.text or "").strip()
                if text_value:
                    blocks.append(f"<p>{html.escape(text_value)}</p>")
            for table in doc.tables:
                rows_html: list[str] = []
                for row in table.rows:
                    cells = "".join(f"<td>{html.escape((cell.text or '').strip())}</td>" for cell in row.cells)
                    rows_html.append(f"<tr>{cells}</tr>")
                if rows_html:
                    blocks.append(f"<table>{''.join(rows_html)}</table>")
            if not blocks:
                blocks.append("<p>No previewable text found in this document.</p>")
            return HTMLResponse(
                """<!doctype html>
<html><head><meta charset="utf-8"><style>
html,body{margin:0;background:#fff;color:#0f172a;font-family:Segoe UI,Arial,sans-serif}
.docx-wrap{padding:20px 22px;font-size:14px;line-height:1.6}
p{margin:0 0 12px}
table{border-collapse:collapse;width:100%;margin:10px 0 16px}
td{border:1px solid #cbd5e1;padding:8px 10px;vertical-align:top}
</style></head><body><div class="docx-wrap">"""
                + "".join(blocks) +
                "</div></body></html>"
            )
        except Exception as exc:
            print(f"[attachment-preview] docx preview failed for {case_id}: {exc}")

    fallback_message = (
        "Preview is not available for this file type in the current environment. "
        "Use the download link below to open the original attachment."
    )
    return HTMLResponse(
        f"""<!doctype html>
<html><head><meta charset="utf-8"><style>
html,body{{height:100%;margin:0;background:#0b1220;color:#e2e8f0;font-family:Segoe UI,Arial,sans-serif}}
body{{display:flex;align-items:center;justify-content:center;padding:24px;box-sizing:border-box}}
.card{{max-width:520px;background:rgba(15,23,42,0.9);border:1px solid rgba(148,163,184,0.2);border-radius:14px;padding:24px;text-align:center}}
.name{{font-weight:600;color:#fff;margin-bottom:10px}}
.msg{{color:#cbd5e1;line-height:1.6;margin-bottom:16px}}
.btn{{display:inline-block;padding:10px 14px;border-radius:8px;background:#1f6feb;color:#fff;text-decoration:none}}
</style></head>
<body><div class="card"><div class="name">{html.escape(filename)}</div><div class="msg">{html.escape(fallback_message)}</div><a class="btn" href="/case/{case_id}/attachment">Download attachment</a></div></body></html>"""
    )


@app.get("/case/{case_id}/pdf")
def case_pdf(request: Request, case_id: str, inline: bool = False):
    try:
        user = require_login(request)

        conn = get_db()
        row = conn.execute("SELECT * FROM cases WHERE id = ?", (case_id,)).fetchone()
        conn.close()
        if not row:
            raise HTTPException(status_code=404, detail="Case not found")
        if user.get("role") == "radiologist":
            raise HTTPException(status_code=403, detail="Radiologists are not allowed to download PDFs")

        org_id = user.get("org_id")
        case_data = row if isinstance(row, dict) else dict(row)

        if org_id and not user.get("is_superuser") and case_data.get("org_id") and case_data.get("org_id") != org_id:
            raise HTTPException(status_code=403, detail="Access denied")

        pdf_path = UPLOAD_DIR / f"{case_id}_vetting.pdf"

        # Bug 8: Convert row to dict to avoid sqlite3.Row.get() issues
        if isinstance(row, dict):
            case_data = row
        else:
            case_data = dict(row)

        # Organisation name
        org_name = ""
        if case_data.get("org_id"):
            conn = get_db()
            org_row = conn.execute("SELECT name FROM organisations WHERE id = ?", (case_data.get("org_id"),)).fetchone()
            conn.close()
            if org_row:
                org_name = org_row.get("name") if isinstance(org_row, dict) else org_row[0]
        report_setting_scope = case_data.get("org_id") or org_id or 0
        report_header_text = get_setting(f"report_header:{report_setting_scope}", org_name or "").strip() or (org_name or "Organisation")
        report_footer_text = get_setting(f"report_footer:{report_setting_scope}", "Confidential workflow document").strip() or "Confidential workflow document"

        events: list[dict] = []
        if table_exists("case_events"):
            conn = get_db()
            event_rows = conn.execute(
                "SELECT * FROM case_events WHERE case_id = ? ORDER BY created_at ASC",
                (case_id,),
            ).fetchall()
            conn.close()
            events = [dict(e) for e in event_rows]

        # Radiologist details (profile + GMC)
        rad_name = case_data.get("radiologist", "")
        rad_display = rad_name
        rad_gmc = ""
        if rad_name and table_exists("radiologist_profiles") and table_exists("users"):
            conn = get_db()
            params = [rad_name, rad_name]
            sql = (
                "SELECT rp.display_name, rp.gmc, u.username "
                "FROM radiologist_profiles rp "
                "JOIN users u ON u.id = rp.user_id "
            )
            if table_exists("memberships") and case_data.get("org_id"):
                sql += "LEFT JOIN memberships m ON m.user_id = u.id "
                sql += "WHERE (rp.display_name = ? OR u.username = ?) AND m.org_id = ? "
                params.append(case_data.get("org_id"))
            else:
                sql += "WHERE rp.display_name = ? OR u.username = ? "
            sql += "LIMIT 1"
            prof = conn.execute(sql, params).fetchone()
            conn.close()
            if prof:
                prof = dict(prof)
                rad_display = prof.get("display_name") or rad_display
                rad_gmc = prof.get("gmc") or ""
        elif rad_name:
            rad = get_radiologist(rad_name)
            if rad:
                rad_gmc = rad.get("gmc", "")

        institution_name = ""
        if case_data.get("institution_id"):
            inst = get_institution(case_data.get("institution_id"))
            if inst:
                institution_name = inst["name"]

        def format_datetime(iso_string: str | None) -> str:
            return format_display_datetime(iso_string)

        def decision_label(case_row: dict) -> str:
            status = str(case_row.get("status") or "").lower()
            decision = str(case_row.get("decision") or "").strip()
            if status == "rejected" or decision == "Reject":
                return "Rejected"
            if decision == "Approve with comment":
                return "Justified with comment"
            if decision == "Approve":
                return "Justified"
            if decision:
                return decision
            if status == "vetted":
                return "Justified"
            return status.title() if status else "Pending"

        def status_label(status_value: str | None) -> str:
            status = str(status_value or "").lower()
            if status == "vetted":
                return "JUSTIFIED"
            return status.upper() if status else "PENDING"

        def event_label(event: dict) -> str:
            event_type = str(event.get("event_type") or "").upper()
            if event_type == "SUBMITTED":
                return "Case submitted"
            if event_type == "ASSIGNED":
                return "Assigned to practitioner"
            if event_type == "OPENED":
                return "Case opened by practitioner"
            if event_type == "REOPENED":
                return "Case reopened by admin"
            if event_type == "VETTED":
                return "Justification recorded"
            if event_type == "EDITED":
                return "Case edited"
            return event_type.title()

        def note_entries_from_events(all_events: list[dict]) -> list[dict]:
            entries: list[dict] = []
            for event in all_events:
                event_type = str(event.get("event_type") or "").upper()
                comment = str(event.get("comment") or "").strip()
                if not comment:
                    continue
                if event_type == "SUBMITTED":
                    entries.append({
                        "kind": "Admin note",
                        "created_at": event.get("created_at"),
                        "text": comment,
                    })
                elif event_type == "REOPENED":
                    entries.append({
                        "kind": "Reopened note",
                        "created_at": event.get("created_at"),
                        "text": comment,
                    })
                elif event_type == "EDITED" and "Notes:" in comment:
                    note_text = comment.split("Notes:", 1)[1].strip()
                    if note_text:
                        entries.append({
                            "kind": "Admin note",
                            "created_at": event.get("created_at"),
                            "text": note_text,
                        })
            return entries

        protocol_notes = ""
        protocol_name = case_data.get("protocol")
        if case_data.get("decision") != "Reject" and protocol_name:
            try:
                conn = get_db()
                if case_data.get("org_id"):
                    protocol_row = conn.execute(
                        "SELECT instructions FROM protocols WHERE name = ? AND org_id = ? LIMIT 1",
                        (protocol_name, case_data.get("org_id")),
                    ).fetchone()
                else:
                    protocol_row = conn.execute(
                        "SELECT instructions FROM protocols WHERE name = ? LIMIT 1",
                        (protocol_name,),
                    ).fetchone()
                conn.close()
                if protocol_row:
                    protocol_notes = protocol_row.get("instructions") if isinstance(protocol_row, dict) else protocol_row["instructions"]
            except Exception as exc:
                print(f"Error fetching protocol instructions: {exc}")

        c = canvas.Canvas(str(pdf_path), pagesize=A4)
        width, height = A4
        left = 42
        right = width - 42
        y = height - 54

        accent = colors.HexColor("#1f6feb")
        accent_soft = colors.HexColor("#dbeafe")
        ink = colors.HexColor("#0f172a")
        muted = colors.HexColor("#475569")
        border = colors.HexColor("#cbd5e1")
        note_fill = colors.HexColor("#f8fafc")
        reopened_fill = colors.HexColor("#ecfeff")
        reopened_border = colors.HexColor("#67e8f9")
        if str(case_data.get("status") or "").lower() == "rejected":
            status_fill = colors.HexColor("#fee2e2")
            status_text = colors.HexColor("#b91c1c")
        elif str(case_data.get("status") or "").lower() == "reopened":
            status_fill = colors.HexColor("#cffafe")
            status_text = colors.HexColor("#0f766e")
        elif str(case_data.get("status") or "").lower() == "vetted":
            status_fill = colors.HexColor("#dcfce7")
            status_text = colors.HexColor("#166534")
        else:
            status_fill = colors.HexColor("#fef3c7")
            status_text = colors.HexColor("#92400e")

        def new_page():
            nonlocal y
            c.showPage()
            y = height - 54

        def ensure_space(required: int):
            nonlocal y
            if y - required < 50:
                new_page()

        def draw_wrapped(text_value: str, x: int, y_top: int, max_width: int, *, font_name: str = "Helvetica", font_size: int = 10, color=ink, leading: int = 13) -> int:
            text = str(text_value or "").strip()
            if not text:
                return y_top
            c.setFillColor(color)
            c.setFont(font_name, font_size)
            paragraphs = text.splitlines() or [text]
            y_cursor = y_top
            for paragraph in paragraphs:
                words = paragraph.split() or [""]
                line_buf = words[0]
                for word in words[1:]:
                    candidate = f"{line_buf} {word}".strip()
                    if c.stringWidth(candidate, font_name, font_size) <= max_width:
                        line_buf = candidate
                    else:
                        c.drawString(x, y_cursor, line_buf)
                        y_cursor -= leading
                        line_buf = word
                c.drawString(x, y_cursor, line_buf)
                y_cursor -= leading
            return y_cursor

        def wrap_lines(text_value: str, max_width: int, *, font_name: str = "Helvetica", font_size: int = 10) -> list[str]:
            text = str(text_value or "").strip()
            if not text:
                return []
            output: list[str] = []
            paragraphs = text.splitlines() or [text]
            for paragraph in paragraphs:
                words = paragraph.split() or [""]
                line_buf = words[0]
                for word in words[1:]:
                    candidate = f"{line_buf} {word}".strip()
                    if c.stringWidth(candidate, font_name, font_size) <= max_width:
                        line_buf = candidate
                    else:
                        output.append(line_buf)
                        line_buf = word
                output.append(line_buf)
            return output

        def wrapped_height(text_value: str, max_width: int, *, font_name: str = "Helvetica", font_size: int = 10, leading: int = 13) -> int:
            text = str(text_value or "").strip()
            if not text:
                return 0
            paragraphs = text.splitlines() or [text]
            total_lines = 0
            for paragraph in paragraphs:
                words = paragraph.split() or [""]
                line_buf = words[0]
                line_count = 1
                for word in words[1:]:
                    candidate = f"{line_buf} {word}".strip()
                    if c.stringWidth(candidate, font_name, font_size) <= max_width:
                        line_buf = candidate
                    else:
                        line_count += 1
                        line_buf = word
                total_lines += line_count
            return total_lines * leading

        def section_title(title: str):
            nonlocal y
            ensure_space(52)
            c.setFillColor(accent)
            c.setFont("Helvetica-Bold", 13)
            c.drawString(left, y, title)
            y -= 12
            c.setStrokeColor(accent_soft)
            c.setLineWidth(1)
            c.line(left, y, right, y)
            y -= 20

        def draw_info_grid(rows: list[tuple[str, str]]):
            nonlocal y
            col_gap = 24
            col_width = (right - left - col_gap) / 2
            grid_rows = [rows[i:i + 2] for i in range(0, len(rows), 2)]
            row_heights: list[int] = []
            for pair in grid_rows:
                max_height = 34
                for _, value in pair:
                    value_height = wrapped_height(value or "-", int(col_width - 24), font_size=10, leading=12)
                    max_height = max(max_height, 22 + value_height)
                row_heights.append(max_height)
            box_height = sum(row_heights) + 18
            ensure_space(int(box_height + 24))
            top = y
            c.setFillColor(colors.white)
            c.setStrokeColor(border)
            c.roundRect(left, top - box_height + 8, right - left, box_height, 10, stroke=1, fill=1)
            current_y = top - 18
            for pair, this_row_height in zip(grid_rows, row_heights):
                for idx, (label, value) in enumerate(pair):
                    x = left + (col_width + col_gap) * idx + 14
                    c.setFillColor(muted)
                    c.setFont("Helvetica-Bold", 8)
                    c.drawString(x, current_y, label.upper())
                    c.setFillColor(ink)
                    c.setFont("Helvetica", 10)
                    draw_wrapped(value or "-", x, current_y - 13, int(col_width - 24), font_size=10, color=ink, leading=12)
                current_y -= this_row_height
            y = top - box_height - 14

        def draw_text_card(text_value: str, *, min_height: int = 92, font_size: int = 10, leading: int = 13):
            nonlocal y
            content = str(text_value or "").strip() or "Not recorded"
            text_height = wrapped_height(content, int(right - left - 28), font_size=font_size, leading=leading)
            box_height = max(min_height, text_height + 36)
            ensure_space(box_height + 18)
            top = y
            c.setFillColor(colors.white)
            c.setStrokeColor(accent_soft)
            c.roundRect(left, top - box_height, right - left, box_height, 12, stroke=1, fill=1)
            draw_wrapped(content, left + 14, top - 22, int(right - left - 28), font_size=font_size, color=ink, leading=leading)
            y = top - box_height - 16

        def draw_note_card(kind: str, created_at: str | None, text_value: str):
            nonlocal y
            lines = str(text_value or "").strip().splitlines() or [""]
            estimated_height = 52 + max(0, len(lines) - 1) * 12
            ensure_space(estimated_height + 12)
            fill_color = reopened_fill if "Reopened" in kind else note_fill
            stroke_color = reopened_border if "Reopened" in kind else border
            c.setFillColor(fill_color)
            c.setStrokeColor(stroke_color)
            c.roundRect(left, y - estimated_height, right - left, estimated_height, 10, stroke=1, fill=1)
            c.setFillColor(accent if "Reopened" in kind else muted)
            c.setFont("Helvetica-Bold", 9)
            c.drawString(left + 14, y - 16, kind.upper())
            c.setFillColor(muted)
            c.setFont("Helvetica", 9)
            c.drawRightString(right - 14, y - 16, format_datetime(created_at))
            y_text = y - 34
            y_after = draw_wrapped(text_value, left + 14, y_text, int(right - left - 28), font_size=10, color=ink, leading=13)
            y = min(y - estimated_height - 10, y_after - 8)

        def draw_timeline_row(timestamp: str, label: str, details: str = ""):
            nonlocal y
            detail_lines = details.splitlines() if details else []
            row_height = 22 + len(detail_lines) * 11
            ensure_space(row_height + 10)
            c.setStrokeColor(border)
            c.line(left, y, right, y)
            y -= 14
            c.setFillColor(muted)
            c.setFont("Helvetica-Bold", 9)
            c.drawString(left, y, timestamp)
            c.setFillColor(ink)
            c.setFont("Helvetica-Bold", 10)
            c.drawString(left + 118, y, label)
            y -= 13
            if details:
                y = draw_wrapped(details, left + 118, y, int(right - left - 118), font_size=9, color=muted, leading=11)
            y -= 6

        # Header
        c.setFillColor(accent)
        c.roundRect(left, y, right - left, 8, 4, stroke=0, fill=1)
        y -= 28
        c.setFillColor(muted)
        header_bottom_y = draw_wrapped(
            report_header_text,
            left,
            y,
            int(right - left - 180),
            font_name="Helvetica-Bold",
            font_size=13,
            color=muted,
            leading=15,
        )
        c.setFont("Helvetica", 10)
        c.drawRightString(right, y + 3, f"Generated: {format_datetime(utc_now_iso())}")
        y = header_bottom_y - 14
        c.setFillColor(ink)
        c.setFont("Helvetica-Bold", 20)
        c.drawString(left, y, "Justification Decision Report")
        y -= 18
        c.setFillColor(muted)
        c.setFont("Helvetica", 10)
        c.drawRightString(right, y, f"Case ID: {case_data.get('id', '')}")
        y -= 28

        badge_width = 110
        c.setFillColor(status_fill)
        c.setStrokeColor(status_fill)
        c.roundRect(left, y - 6, badge_width, 24, 12, stroke=1, fill=1)
        c.setFillColor(status_text)
        c.setFont("Helvetica-Bold", 10)
        c.drawCentredString(left + badge_width / 2, y + 2, status_label(case_data.get("status")))
        y -= 28

        section_title("Case Summary")
        patient_name = f"{case_data.get('patient_first_name') or ''} {case_data.get('patient_surname') or ''}".strip() or "Not recorded"
        draw_info_grid([
            ("Patient Name", patient_name),
            ("Patient ID / NHS Number", case_data.get("patient_referral_id") or "Not recorded"),
            ("Date of Birth", case_data.get("patient_dob") or "Not recorded"),
            ("Institution", institution_name or "Not recorded"),
            ("Modality", case_data.get("modality") or "Not recorded"),
            ("Study Description", case_data.get("study_description") or "Not recorded"),
            ("Radiologist", rad_display or "Not assigned"),
            ("Radiologist GMC", rad_gmc or "Not recorded"),
        ])

        section_title("Justification Decision")
        decision_value_width = int(right - left - 175)
        decision_rows = [
            ("Decision", decision_label(case_data)),
            ("Protocol", case_data.get("protocol") or "Not recorded"),
            ("Justified By", rad_display or "Not assigned"),
            ("Radiologist GMC", rad_gmc or "Not recorded"),
            ("Justified At", format_datetime(case_data.get("vetted_at")) or "Not recorded"),
        ]
        row_heights: list[int] = []
        for _, value in decision_rows:
            row_heights.append(max(18, wrapped_height(str(value), decision_value_width, font_size=10, leading=12)))
        comment_height = wrapped_height(case_data.get("decision_comment") or "", int(right - left - 28), font_size=10, leading=12)
        decision_box_height = 22 + sum(row_heights) + (len(row_heights) - 1) * 10
        if case_data.get("decision_comment"):
            decision_box_height += 30 + comment_height
        decision_box_height += 20
        ensure_space(decision_box_height + 18)
        c.setFillColor(colors.white)
        c.setStrokeColor(accent_soft)
        c.roundRect(left, y - decision_box_height, right - left, decision_box_height, 12, stroke=1, fill=1)
        top_y = y - 18
        row_y = top_y
        for idx, (label, value) in enumerate(decision_rows):
            c.setFillColor(muted)
            c.setFont("Helvetica-Bold", 8)
            c.drawString(left + 14, row_y, label.upper())
            c.setFillColor(ink)
            c.setFont("Helvetica", 10)
            value_y = draw_wrapped(str(value), left + 145, row_y, decision_value_width, font_size=10, color=ink, leading=12)
            row_y -= row_heights[idx] + 10
        if case_data.get("decision_comment"):
            row_y -= 2
            c.setFillColor(muted)
            c.setFont("Helvetica-Bold", 8)
            c.drawString(left + 14, row_y, "DECISION COMMENT")
            row_y -= 16
            row_y = draw_wrapped(case_data.get("decision_comment") or "", left + 14, row_y, int(right - left - 28), font_size=10, color=ink, leading=12)
        y = y - decision_box_height - 16

        if protocol_notes:
            section_title("Protocol Notes")
            draw_text_card(protocol_notes, min_height=92, font_size=10, leading=13)

        section_title("Timeline")
        if events:
            for event in events:
                event_details = ""
                if str(event.get("event_type") or "").upper() == "SUBMITTED":
                    event_details = ""
                elif event.get("comment"):
                    event_details = str(event.get("comment") or "").strip()
                elif event.get("decision"):
                    event_details = str(event.get("decision") or "").strip()
                draw_timeline_row(
                    format_datetime(event.get("created_at")),
                    event_label(event),
                    event_details,
                )
        else:
            draw_timeline_row(format_datetime(case_data.get("created_at")), "Case submitted", "")

        footer_lines = wrap_lines(report_footer_text, int((right - left) * 0.52), font_name="Helvetica", font_size=8) or ["Confidential workflow document"]
        footer_height = max(16, len(footer_lines) * 10)
        ensure_space(22 + footer_height)
        c.setStrokeColor(border)
        c.line(left, y, right, y)
        y -= 16
        c.setFillColor(muted)
        c.setFont("Helvetica", 8)
        c.drawString(left, y, "Generated by RadFlow")
        footer_y = y
        for idx, line in enumerate(footer_lines):
            c.drawRightString(right, footer_y - (idx * 10), line)

        c.showPage()
        c.save()

        if inline:
            headers = {"Content-Disposition": f'inline; filename="vetting_{case_id}.pdf"'}
            return FileResponse(str(pdf_path), media_type="application/pdf", headers=headers)

        return FileResponse(str(pdf_path), filename=f"vetting_{case_id}.pdf")
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {e}")


# -------------------------
# SUPERUSER ROUTES - Multi-Tenant Management
# -------------------------

@app.get("/account", response_class=HTMLResponse)
def account_page(request: Request, msg: str = "", error: str = ""):
    """Any authenticated user can view/edit their own profile (name, email, password)."""
    user = get_session_user(request)
    if not user:
        return RedirectResponse(url="/login?expired=1", status_code=303)

    conn = get_db()
    db_user = conn.execute(
        "SELECT id, username, first_name, surname, email, COALESCE(mfa_enabled, 0) AS mfa_enabled, COALESCE(mfa_required, 0) AS mfa_required, mfa_pending_secret FROM users WHERE username = ?",
        (user["username"],)
    ).fetchone()
    conn.close()

    if not db_user:
        return RedirectResponse(url="/login?expired=1", status_code=303)

    db_user = dict(db_user)
    if user.get("org_role") in ("org_admin",) or user.get("role") in ("admin",):
        back_url = "/admin"
    else:
        back_url = "/radiologist"

    msg_html = ""
    if msg == "saved":
        msg_html = '<div style="background:rgba(74,222,128,0.12);border:1px solid rgba(74,222,128,0.3);color:#4ade80;padding:12px 16px;border-radius:8px;margin-bottom:16px;">✅ Your profile has been updated.</div>'
    elif msg == "pw_changed":
        msg_html = '<div style="background:rgba(74,222,128,0.12);border:1px solid rgba(74,222,128,0.3);color:#4ade80;padding:12px 16px;border-radius:8px;margin-bottom:16px;">✅ Password changed successfully.</div>'

    error_html = ""
    if error == "email_taken":
        error_html = '<div style="background:rgba(239,68,68,0.12);border:1px solid rgba(239,68,68,0.4);color:#fca5a5;padding:12px 16px;border-radius:8px;margin-bottom:16px;">⚠️ That email address is already in use by another account.</div>'
    elif error == "pw_mismatch":
        error_html = '<div style="background:rgba(239,68,68,0.12);border:1px solid rgba(239,68,68,0.4);color:#fca5a5;padding:12px 16px;border-radius:8px;margin-bottom:16px;">⚠️ New passwords do not match. Please try again.</div>'
    elif error == "pw_wrong":
        error_html = '<div style="background:rgba(239,68,68,0.12);border:1px solid rgba(239,68,68,0.4);color:#fca5a5;padding:12px 16px;border-radius:8px;margin-bottom:16px;">⚠️ Current password is incorrect.</div>'
    elif error == "pw_short":
        error_html = '<div style="background:rgba(239,68,68,0.12);border:1px solid rgba(239,68,68,0.4);color:#fca5a5;padding:12px 16px;border-radius:8px;margin-bottom:16px;">⚠️ New password must be at least 8 characters.</div>'

    msg_map = {
        "saved": "Your profile has been updated.",
        "pw_changed": "Password changed successfully.",
        "mfa_required": "Authenticator-based MFA is required for admin access. Complete setup below to continue.",
        "mfa_started": "Authenticator setup started. Add the account in your app, then enter the 6-digit code below to finish.",
        "mfa_enabled": "Authenticator-based MFA is now enabled for your account.",
        "mfa_disabled": "Authenticator-based MFA has been disabled.",
        "mfa_managed": "MFA is currently managed by your administrator and cannot be disabled here.",
    }
    msg_text = msg_map.get(msg, "")
    msg_html = (
        f'<div style="background:rgba(74,222,128,0.12);border:1px solid rgba(74,222,128,0.3);color:#4ade80;padding:12px 16px;border-radius:8px;margin-bottom:16px;">{html.escape(msg_text)}</div>'
        if msg_text else ""
    )
    if msg == "mfa_started":
        msg_html = f'<div style="background:rgba(96,165,250,0.12);border:1px solid rgba(96,165,250,0.3);color:#93c5fd;padding:12px 16px;border-radius:8px;margin-bottom:16px;">{html.escape(msg_text)}</div>'
    elif msg == "mfa_required":
        msg_html = f'<div style="background:rgba(234,179,8,0.12);border:1px solid rgba(234,179,8,0.35);color:#fde68a;padding:12px 16px;border-radius:8px;margin-bottom:16px;">{html.escape(msg_text)}</div>'

    error_map = {
        "email_taken": "That email address is already in use by another account.",
        "pw_mismatch": "New passwords do not match. Please try again.",
        "pw_wrong": "Current password is incorrect.",
        "pw_short": "New password must be at least 8 characters.",
        "mfa_invalid": "The authentication code was not valid. Please try again.",
        "mfa_pw_wrong": "Your current password was incorrect.",
        "mfa_managed": "MFA is currently managed by your administrator and cannot be disabled here.",
    }
    error_text = error_map.get(error, "")
    error_html = (
        f'<div style="background:rgba(239,68,68,0.12);border:1px solid rgba(239,68,68,0.4);color:#fca5a5;padding:12px 16px;border-radius:8px;margin-bottom:16px;">{html.escape(error_text)}</div>'
        if error_text else ""
    )

    mfa_enabled = bool(db_user.get("mfa_enabled"))
    mfa_required = bool(db_user.get("mfa_required"))
    mfa_pending_secret = db_user.get("mfa_pending_secret") or ""
    mfa_uri = build_totp_uri(mfa_pending_secret, db_user["username"]) if mfa_pending_secret else ""
    mfa_qr_data_uri = build_totp_qr_data_uri(mfa_uri) if mfa_uri else ""

    page_html = f"""<!DOCTYPE html>
<html>
<head>
    <title>My Account</title>
    <link rel="stylesheet" href="/static/css/site.css">
    <style>
        .account-wrap {{ max-width: 1400px; width: 95%; margin: 0 auto; padding: 14px 20px 32px; }}
        .account-shell {{ max-width: 1280px; margin: 0 auto; }}
        .page-title {{ font-size: 2em; color: white; margin: 0 0 6px 0; }}
        .page-sub {{ color: var(--muted); margin: 0 0 28px 0; }}
        .card {{
            background: var(--card-bg);
            border: 1px solid var(--card-border);
            border-radius: 10px;
            padding: 24px;
            margin-bottom: 24px;
        }}
        .cards-grid {{
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 24px;
            align-items: start;
            margin-bottom: 24px;
        }}
        .cards-grid .card {{
            margin-bottom: 0;
        }}
        .card h3 {{ margin-top: 0; color: rgba(255,255,255,0.9); font-size: 1.15em; }}
        .form-group {{ display: flex; flex-direction: column; gap: 5px; margin-bottom: 16px; }}
        .form-group label {{ font-size: 0.88em; color: rgba(255,255,255,0.65); font-weight: 500; }}
        .form-group input {{
            background: rgba(255,255,255,0.06);
            border: 1px solid rgba(255,255,255,0.15);
            border-radius: 6px;
            color: #fff;
            padding: 9px 12px;
            font-size: 0.95em;
        }}
        .form-group input:focus {{ outline: none; border-color: rgba(31,111,235,0.6); }}
        .topbar {{ display: flex; justify-content: space-between; align-items: center; gap: 10px; margin-bottom: 24px; }}
        .topbar-actions {{ display: flex; gap: 10px; align-items: center; }}
        .profile-grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 16px; }}
        .read-only {{ color: rgba(255,255,255,0.5); font-size: 0.92em; padding: 8px 0; }}
        @media (max-width: 900px) {{
            .account-wrap {{ width: 100%; padding: 14px 12px 24px; box-sizing: border-box; }}
            .profile-grid {{ grid-template-columns: 1fr; }}
            .cards-grid {{ grid-template-columns: 1fr; }}
        }}
        @media (max-width: 640px) {{
            .topbar {{ flex-direction: column; align-items: stretch; }}
            .topbar-actions {{ width: 100%; justify-content: space-between; }}
        }}
    </style>
</head>
<body>
<div id="session-expiry-warning" style="display:none;position:fixed;top:0;left:0;right:0;z-index:9999;
    background:rgba(234,179,8,0.95);color:#1a1200;text-align:center;padding:10px 16px;font-weight:600;font-size:14px;">
    ⚠️ Your session will expire soon due to inactivity.
    <button onclick="document.getElementById('session-expiry-warning').style.display='none'"
        style="margin-left:16px;background:rgba(0,0,0,0.15);border:none;border-radius:4px;padding:4px 10px;cursor:pointer;font-weight:600;">
        Dismiss
    </button>
</div>
<div class="account-wrap">
    <div class="account-shell">
    <div class="topbar">
        <a href="{back_url}" class="btn secondary">&larr; Back</a>
        <div class="topbar-actions">
            <a href="/logout" class="btn secondary">Logout</a>
        </div>
    </div>
    <h1 class="page-title">My Account</h1>
    <p class="page-sub">Edit your personal details. Role and permissions are managed by your administrator.</p>

    {msg_html}{error_html}

    <div class="cards-grid">
    <!-- Profile Details -->
    <div class="card">
        <h3>Personal Details</h3>
        <form method="POST" action="/account/edit">
            <div class="profile-grid">
                <div class="form-group" style="flex:1;">
                    <label>First Name</label>
                    <input type="text" name="first_name" value="{db_user.get('first_name') or ''}">
                </div>
                <div class="form-group" style="flex:1;">
                    <label>Surname</label>
                    <input type="text" name="surname" value="{db_user.get('surname') or ''}">
                </div>
            </div>
            <div class="form-group">
                <label>Email Address</label>
                <input type="email" name="email" value="{db_user.get('email') or ''}">
            </div>
            <div class="form-group">
                <label>Username <span style="color:var(--muted);font-weight:400;">(cannot be changed)</span></label>
                <div class="read-only">{db_user['username']}</div>
            </div>
            <button type="submit" class="btn btn-primary">Save Details</button>
        </form>
    </div>

    <!-- Change Password -->
    <div class="card">
        <h3>Change Password</h3>
        <form method="POST" action="/account/change-password">
            <div class="form-group">
                <label>Current Password</label>
                <input type="password" name="current_password" required autocomplete="current-password">
            </div>
            <div class="form-group">
                <label>New Password</label>
                <input type="password" name="new_password" required autocomplete="new-password"
                       minlength="8" placeholder="At least 8 characters">
            </div>
            <div class="form-group">
                <label>Confirm New Password</label>
                <input type="password" name="confirm_password" required autocomplete="new-password">
            </div>
            <button type="submit" class="btn btn-primary">Change Password</button>
        </form>
    </div>
    </div>

    <div class="card">
        <h3>Authenticator App MFA</h3>
        <p class="page-sub" style="margin-bottom:16px;">Use Microsoft Authenticator, Google Authenticator, or another TOTP app for a second sign-in step.</p>
        {('<div class="read-only" style="margin-bottom:12px;color:#fde68a;">Your organisation currently requires MFA for this account.</div>' if mfa_required else '')}
        {(
            f'''
        <div class="read-only" style="margin-bottom:12px;color:#4ade80;">Authenticator-based MFA is enabled for this account.</div>
        <form method="POST" action="/account/mfa/disable">
            <div class="form-group">
                <label>Current Password</label>
                <input type="password" name="current_password" required autocomplete="current-password">
            </div>
            <div class="form-group">
                <label>Authenticator Code</label>
                <input type="text" name="code" inputmode="numeric" pattern="[0-9]{{6}}" maxlength="6" placeholder="123456" required>
            </div>
            {"<button type=\"submit\" class=\"btn secondary\">Disable MFA</button>" if not mfa_required else "<div class=\"read-only\" style=\"color:#93c5fd;\">MFA is required for this account and can only be turned off by an administrator.</div>"}
        </form>
            '''
            if mfa_enabled else
            (
                f'''
        <div class="read-only" style="margin-bottom:12px;">Setup is in progress. Add this secret to your authenticator app and then verify with a current 6-digit code.</div>
        <div style="display:flex;justify-content:center;margin:8px 0 18px;">
            <img src="{mfa_qr_data_uri}" alt="MFA QR code" style="background:#fff;padding:10px;border-radius:10px;max-width:220px;width:100%;height:auto;">
        </div>
        <div class="form-group">
            <label>Manual Setup Secret</label>
            <input type="text" value="{mfa_pending_secret}" readonly>
        </div>
        <div class="form-group">
            <label>Setup URI</label>
            <input type="text" value="{html.escape(mfa_uri, quote=True)}" readonly>
        </div>
        <form method="POST" action="/account/mfa/enable">
            <div class="form-group">
                <label>6-Digit Code From App</label>
                <input type="text" name="code" inputmode="numeric" pattern="[0-9]{{6}}" maxlength="6" placeholder="123456" required>
            </div>
            <button type="submit" class="btn btn-primary">Enable MFA</button>
        </form>
                '''
                if mfa_pending_secret else
                '''
        <div class="read-only" style="margin-bottom:12px;">Authenticator-based MFA is currently disabled.</div>
        <form method="POST" action="/account/mfa/begin">
            <button type="submit" class="btn btn-primary">Set Up Authenticator App</button>
        </form>
                '''
            )
        )}
    </div>
    </div>
</div>
<script src="/static/js/session.js"></script>
</body>
</html>"""

    return HTMLResponse(content=page_html)


@app.post("/account/edit")
def account_edit(
    request: Request,
    first_name: str = Form(""),
    surname: str = Form(""),
    email: str = Form(""),
):
    """Any authenticated user: update own name/email (not role)."""
    user = get_session_user(request)
    if not user:
        return RedirectResponse(url="/login?expired=1", status_code=303)

    username = user["username"]
    new_email = email.strip()

    conn = get_db()
    db_user = conn.execute("SELECT email FROM users WHERE username = ?", (username,)).fetchone()
    conn.close()
    if not db_user:
        return RedirectResponse(url="/login?expired=1", status_code=303)

    current_email = db_user["email"] or ""
    email_changed = new_email and new_email != current_email

    if email_changed:
        conn2 = get_db()
        conflict = conn2.execute(
            "SELECT id FROM users WHERE email = ? AND username != ?", (new_email, username)
        ).fetchone()
        conn2.close()
        if conflict:
            return RedirectResponse(url="/account?error=email_taken", status_code=303)

    conn = get_db()
    if email_changed:
        conn.execute(
            "UPDATE users SET first_name = ?, surname = ?, email = ? WHERE username = ?",
            (first_name.strip(), surname.strip(), new_email, username),
        )
    else:
        conn.execute(
            "UPDATE users SET first_name = ?, surname = ? WHERE username = ?",
            (first_name.strip(), surname.strip(), username),
        )
    conn.commit()
    conn.close()

    # Update session display name if changed
    if first_name.strip() or surname.strip():
        user["first_name"] = first_name.strip()
        user["surname"] = surname.strip()
        request.session["user"] = user

    return RedirectResponse(url="/account?msg=saved", status_code=303)


@app.post("/account/change-password")
def account_change_password(
    request: Request,
    current_password: str = Form(...),
    new_password: str = Form(...),
    confirm_password: str = Form(...),
):
    """Any authenticated user: change own password with current-password verification."""
    user = get_session_user(request)
    if not user:
        return RedirectResponse(url="/login?expired=1", status_code=303)

    username = user["username"]

    if new_password != confirm_password:
        return RedirectResponse(url="/account?error=pw_mismatch", status_code=303)

    if len(new_password) < 8:
        return RedirectResponse(url="/account?error=pw_short", status_code=303)

    if not verify_current_password(username, current_password):
        return RedirectResponse(url="/account?error=pw_wrong", status_code=303)

    # Set new password
    new_salt = secrets.token_bytes(16)
    new_hash = hash_password(new_password, new_salt)

    conn = get_db()
    conn.execute(
        "UPDATE users SET salt_hex = ?, password_hash = ? WHERE username = ?",
        (new_salt.hex(), new_hash.hex(), username),
    )
    conn.commit()
    conn.close()

    return RedirectResponse(url="/account?msg=pw_changed", status_code=303)


@app.post("/account/mfa/begin")
def account_mfa_begin(request: Request):
    user = get_session_user(request)
    if not user:
        return RedirectResponse(url="/login?expired=1", status_code=303)

    secret = generate_totp_secret()
    conn = get_db()
    conn.execute(
        "UPDATE users SET mfa_pending_secret = ?, modified_at = ? WHERE username = ?",
        (secret, utc_now_iso(), user["username"]),
    )
    conn.commit()
    conn.close()
    return RedirectResponse(url="/account?msg=mfa_started", status_code=303)


@app.post("/account/mfa/enable")
def account_mfa_enable(request: Request, code: str = Form(...)):
    user = get_session_user(request)
    if not user:
        return RedirectResponse(url="/login?expired=1", status_code=303)

    conn = get_db()
    row = conn.execute(
        "SELECT mfa_pending_secret FROM users WHERE username = ?",
        (user["username"],),
    ).fetchone()
    if not row:
        conn.close()
        return RedirectResponse(url="/login?expired=1", status_code=303)

    pending_secret = row["mfa_pending_secret"] if isinstance(row, dict) else row[0]
    if not verify_totp_code(pending_secret, code):
        conn.close()
        return RedirectResponse(url="/account?error=mfa_invalid", status_code=303)

    conn.execute(
        "UPDATE users SET mfa_secret = ?, mfa_pending_secret = NULL, mfa_enabled = 1, modified_at = ? WHERE username = ?",
        (pending_secret, utc_now_iso(), user["username"]),
    )
    conn.commit()
    conn.close()
    user["mfa_enabled"] = 1
    request.session["user"] = user
    return RedirectResponse(url="/account?msg=mfa_enabled", status_code=303)


@app.post("/account/mfa/disable")
def account_mfa_disable(
    request: Request,
    current_password: str = Form(...),
    code: str = Form(...),
):
    user = get_session_user(request)
    if not user:
        return RedirectResponse(url="/login?expired=1", status_code=303)

    if not verify_current_password(user["username"], current_password):
        return RedirectResponse(url="/account?error=mfa_pw_wrong", status_code=303)

    conn = get_db()
    row = conn.execute(
        "SELECT mfa_secret, COALESCE(mfa_required, 0) AS mfa_required FROM users WHERE username = ?",
        (user["username"],),
    ).fetchone()
    if not row:
        conn.close()
        return RedirectResponse(url="/login?expired=1", status_code=303)

    mfa_secret = row["mfa_secret"] if isinstance(row, dict) else row[0]
    mfa_required = int(row["mfa_required"] if isinstance(row, dict) else row[1] or 0)
    if mfa_required:
        conn.close()
        return RedirectResponse(url="/account?error=mfa_managed", status_code=303)
    if not verify_totp_code(mfa_secret, code):
        conn.close()
        return RedirectResponse(url="/account?error=mfa_invalid", status_code=303)

    conn.execute(
        "UPDATE users SET mfa_secret = NULL, mfa_pending_secret = NULL, mfa_enabled = 0, modified_at = ? WHERE username = ?",
        (utc_now_iso(), user["username"]),
    )
    conn.commit()
    conn.close()
    user["mfa_enabled"] = 0
    request.session["user"] = user
    return RedirectResponse(url="/account?msg=mfa_disabled", status_code=303)


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000, log_level="info")




